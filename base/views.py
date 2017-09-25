# -*- coding: utf-8 -*-
from django.http import HttpResponseRedirect
from django.shortcuts import get_object_or_404, render
from base.models import *
from base.forms import *
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
import xlrd
from xlrd.biffh import XLRDError
import datetime
import os
from django.template import Context
from django.template.loader import get_template
from django.template import RequestContext
from xhtml2pdf import pisa
from django.db.models import Q, F, Count
from dal import autocomplete
from django.contrib.auth.models import Group
from templatetags.app_filters import format_money, format_quantidade, format_numero_extenso, format_numero
from django.db.transaction import atomic
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.graphics.barcode.common import I2of5
from reportlab.lib.utils import simpleSplit
import collections
from django.conf import settings
from django.core.mail import send_mail
from xlutils.copy import copy # http://pypi.python.org/pypi/xlutils
from xlrd import open_workbook # http://pypi.python.org/pypi/xlrd
import xlwt
from django.core.exceptions import PermissionDenied
from licita import settings
LARGURA = 210*mm
ALTURA = 297*mm
from django.utils.formats import localize
from datetime import timedelta
from django.db import transaction
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx import Document
from docx.shared import Inches, Pt


def get_config(secretaria=None):
    if secretaria and secretaria.logo:
        return secretaria
    if Configuracao.objects.exists():
        return Configuracao.objects.latest('id')
    return False

def get_config_geral():
    if Configuracao.objects.exists():
        return Configuracao.objects.latest('id')
    return False

class SecretariaAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):

        # Don't forget to filter out results depending on the visitor !
        if not self.request.user.is_authenticated():
            return Secretaria.objects.none()

        qs = Secretaria.objects.exclude(id=self.request.user.pessoafisica.setor.secretaria.id)

        if self.q:
            qs = qs.filter(nome__istartswith=self.q)

        return qs

class ParticipantePregaoAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):

        # Don't forget to filter out results depending on the visitor !
        if not self.request.user.is_authenticated():
            return Fornecedor.objects.none()

        qs = Fornecedor.objects.all()

        if self.q:
            qs = qs.filter(Q(cnpj__istartswith=self.q) | Q(razao_social__istartswith=self.q))

        return qs

class PessoaFisicaAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):

        # Don't forget to filter out results depending on the visitor !
        if not self.request.user.is_authenticated():
            return PessoaFisica.objects.none()

        qs = PessoaFisica.objects.all()

        if self.q:
            qs = qs.filter(Q(nome__istartswith=self.q) | Q(cpf__istartswith=self.q))

        return qs


class MaterialConsumoAutocomplete(autocomplete.Select2QuerySetView):
    def get_queryset(self):

        # Don't forget to filter out results depending on the visitor !
        if not self.request.user.is_authenticated():
            return MaterialConsumo.objects.none()

        qs = MaterialConsumo.objects.all()

        if self.q:
            qs = qs.filter(Q(codigo__icontains=self.q) | Q(nome__icontains=self.q))

        return qs

def imprimir_cabecalho(document, configuracao, logo, municipio):

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    # hdr_cells2 = table.rows[1].cells
    # hdr_cells3 = table.rows[2].cells


    # style2 = document.styles['Normal']
    # font = style2.font
    # font.name = 'Arial'
    # font.size = Pt(6)
    #
    # style = document.styles['Normal']
    # font = style.font
    # font.name = 'Arial'
    # font.size = Pt(11)



    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(logo, width=Inches(1.75))
    a, b = hdr_cells[:2]
    a.merge(b)

    #document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'%s' % (configuracao.nome))

    #document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'Endereço: %s, %s' % (configuracao.endereco, municipio))

    # paragraph2 = hdr_cells2[0].paragraphs[0]
    # paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # paragraph2.style = document.styles['Normal']
    # hdr_cells2[0].text =  u'%s' % (configuracao.nome)
    # a, b = hdr_cells2[:2]
    # a.merge(b)
    # a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #
    # paragraph3 = hdr_cells2[1].paragraphs[0]
    # paragraph3.style2 = document.styles['Normal']
    # paragraph3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    #
    #
    # #hdr_cells2[0].text =  u'Sistema Orçamentário, Financeiro e Contábil'
    # hdr_cells3[0].text =  u'Endereço: %s, %s' % (configuracao.endereco, municipio)
    #
    # a, b = hdr_cells3[:2]
    # a.merge(b)

def logout(request):
    messages.error(request, u'Usuário não vinculado à um setor. Procure o administrador do sistema.')
    return HttpResponseRedirect(u'/admin/logout/?next=/')

@login_required()
def index(request):

    if not hasattr(request.user, 'pessoafisica') or not hasattr(request.user.pessoafisica, 'setor'):
        return HttpResponseRedirect(u'/base/logout/')
    eh_ordenador_despesa = False
    if get_config():
        eh_ordenador_despesa = request.user.pessoafisica == get_config().ordenador_despesa
    tem_solicitacao = False

    if MovimentoSolicitacao.objects.filter(setor_destino=request.user.pessoafisica.setor, recebido_por__isnull=True).exists():
        tem_solicitacao = MovimentoSolicitacao.objects.filter(setor_destino=request.user.pessoafisica.setor, recebido_por__isnull=True)[0]
    tem_preencher_itens = list()
    if SolicitacaoLicitacao.objects.filter(liberada_para_pedido=True, interessados=request.user.pessoafisica.setor.secretaria, prazo_resposta_interessados__gte=datetime.datetime.now().date(), itemsolicitacaolicitacao__isnull=False, setor_atual=F('setor_origem')).exists():
        for item in SolicitacaoLicitacao.objects.filter(interessados=request.user.pessoafisica.setor.secretaria, prazo_resposta_interessados__gte=datetime.datetime.now().date(), itemsolicitacaolicitacao__isnull=False):
            if not ItemQuantidadeSecretaria.objects.filter(solicitacao=item, secretaria=request.user.pessoafisica.setor.secretaria).exists():
                if not item in tem_preencher_itens:
                    tem_preencher_itens.append(item)
                continue


    return render(request, 'index.html', locals(), RequestContext(request))

@login_required()
def solicitacoes(request):
    title = u'Solicitações'
    return render(request, 'solicitacoes.html', locals(), RequestContext(request))


@login_required()
def administracao(request):
    if request.user.is_superuser:
        title = u'Administração'
        return render(request, 'administracao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def auditoria(request):
    if request.user.is_superuser:
        title = u'Auditoria'
        return render(request, 'auditoria.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastros(request):
    title = u'Cadastros'
    return render(request, 'cadastros.html', locals(), RequestContext(request))

@login_required()
def fornecedor(request, fornecedor_id):

    fornecedor = get_object_or_404(Fornecedor, pk= fornecedor_id)
    title = u'Dados do Fornecedor: %s' % fornecedor.razao_social
    exibe_popup = True
    return render(request, 'ver_fornecedores.html', locals(), RequestContext(request))

@login_required()
def pregao(request, pregao_id):

    pregao = get_object_or_404(Pregao, pk= pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        recebida_setor = pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor)
        # if not (pregao.solicitacao.setor_atual == request.user.pessoafisica.setor) and pregao.eh_ativo():
        #     pregao.situacao = Pregao.CONCLUIDO
        #     pregao.save()
        #
        # elif pregao.solicitacao.setor_atual == request.user.pessoafisica.setor and pregao.situacao == Pregao.CONCLUIDO:
        #     pregao.situacao = Pregao.CADASTRADO
        #     pregao.save()

        eh_credenciamento = pregao.solicitacao.eh_credenciamento()
        eh_lote = pregao.criterio.id == CriterioPregao.LOTE

        eh_maior_desconto = pregao.eh_maior_desconto()
        title = u'%s ' % pregao
        if eh_lote:
            lotes = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True)
            itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao=ItemSolicitacaoLicitacao.CADASTRADO)
        else:
            itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao=ItemSolicitacaoLicitacao.CADASTRADO)
        #title = u'Pregão: %s (Processo: %s) - Situação: %s' % (pregao.num_pregao, pregao.num_processo, pregao.situacao)
        itens_pregao_unidades = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False)
        participantes = ParticipantePregao.objects.filter(pregao=pregao, desclassificado=False)
        resultados = ResultadoItemPregao.objects.filter(item__in=itens_pregao.values_list('id',flat=True))
        buscou = False
        ids_ganhador = list()

        participante = u'0'
        if request.GET.get('participante'):
            buscou = True
            participante = get_object_or_404(ParticipantePregao, pk=request.GET.get('participante'))

            for opcao in itens_pregao:
                resultados = ResultadoItemPregao.objects.filter(item=opcao, situacao=ResultadoItemPregao.CLASSIFICADO).order_by('ordem')
                if resultados.exists() and resultados[0].participante == participante:
                    ids_ganhador.append(resultados[0].item.id)

            if not pregao.solicitacao.eh_credenciamento():
                itens_pregao = itens_pregao.filter(id__in=ids_ganhador)

            form = GanhadoresForm(request.POST or None, participantes = participantes, initial=dict(ganhador=participante))
        else:
            form = GanhadoresForm(request.POST or None, participantes = participantes)


        return render(request, 'pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastra_proposta_pregao(request, pregao_id):

    title=u'Cadastrar Proposta'
    pregao = get_object_or_404(Pregao, pk= pregao_id)
    if not pregao.tem_resultado() and request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        itens = pregao.solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=False).order_by('item')
        edicao=False
        participante = None
        selecionou = False
        total = ParticipantePregao.objects.filter(pregao=pregao).count()
        informados = PropostaItemPregao.objects.filter(pregao=pregao).values('participante').distinct().count()
        if request.GET.get('participante'):
            selecionou = True
            participante = get_object_or_404(ParticipantePregao, pk=request.GET.get('participante'))
            itens = PropostaItemPregao.objects.filter(pregao=pregao, participante=participante, item__eh_lote=False).order_by('item')
            if itens.exists():
                edicao=True
            else:
                itens = pregao.solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=False).order_by('item')
        if edicao or selecionou or request.POST:
            form = CadastraPrecoParticipantePregaoForm(request.POST or None, request.FILES, pregao=pregao, initial=dict(fornecedor=participante))
        else:
            form = CadastraPrecoParticipantePregaoForm(request.POST or None, request.FILES, pregao=pregao)
        if form.is_valid():

            arquivo_up = form.cleaned_data.get('arquivo')
            fornecedor = form.cleaned_data.get('fornecedor')
            if arquivo_up:
                sheet = None
                try:
                    with transaction.atomic():
                        workbook = xlrd.open_workbook(file_contents=arquivo_up.read())
                        sheet = workbook.sheet_by_index(0)
                except XLRDError:
                    raise Exception(u'Não foi possível processar a planilha. Verfique se o formato do arquivo é .xls ou .xlsx.')
                PropostaItemPregao.objects.filter(pregao=pregao, participante=fornecedor).delete()
                for row in range(9, sheet.nrows):
                    try:
                        with transaction.atomic():
                            item = unicode(sheet.cell_value(row, 0)).strip()
                            marca = unicode(sheet.cell_value(row, 5)).strip() or None
                            valor = unicode(sheet.cell_value(row, 6)).strip()
                            if row == 0:
                                if item != u'Item' or valor != u'VALOR UNITÁRIO':
                                    raise Exception(u'Não foi possível processar a planilha. As colunas devem ter Item e Valor.')
                            else:
                                if item and valor:
                                    item_do_pregao = ItemSolicitacaoLicitacao.objects.get(eh_lote=False, solicitacao=pregao.solicitacao,item=int(sheet.cell_value(row, 0)))
                                    if PropostaItemPregao.objects.filter(item=item_do_pregao, pregao=pregao, participante=fornecedor).exists():
                                        PropostaItemPregao.objects.filter(item=item_do_pregao, pregao=pregao, participante=fornecedor).update(marca=marca, valor=valor)
                                    else:
                                        novo_preco = PropostaItemPregao()
                                        novo_preco.item = item_do_pregao
                                        novo_preco.pregao = pregao
                                        novo_preco.participante = fornecedor
                                        try:
                                            Decimal(valor)
                                        except:
                                            messages.error(request, u'o valor %s do %s é inválido.' % (valor, item_do_pregao))
                                            return HttpResponseRedirect(u'/base/cadastra_proposta_pregao/%s/?participante=%s' % (pregao.id, fornecedor.id))

                                        if Decimal(valor) <= 0:
                                            messages.error(request, u'o valor %s do %s é inválido.' % (valor, item_do_pregao))
                                            return HttpResponseRedirect(u'/base/cadastra_proposta_pregao/%s/?participante=%s' % (pregao.id, fornecedor.id))


                                        novo_preco.valor = valor
                                        novo_preco.marca = marca
                                        novo_preco.save()
                                    if not ParticipanteItemPregao.objects.filter(participante=fornecedor, item=item_do_pregao).exists():
                                        participante_item = ParticipanteItemPregao()
                                        participante_item.participante = fornecedor
                                        participante_item.item = item_do_pregao
                                        participante_item.save()


                    except ValueError:
                        raise Exception(u'Alguma coluna da planilha possui o valor incorreto.')
            if not edicao and not arquivo_up:
                for idx, item in enumerate(request.POST.getlist('itens'), 1):
                    if item:
                        item_do_pregao = ItemSolicitacaoLicitacao.objects.get(eh_lote=False, solicitacao=pregao.solicitacao, id=request.POST.getlist('id_item')[idx-1])
                        novo_preco = PropostaItemPregao()
                        novo_preco.item = item_do_pregao
                        novo_preco.pregao = pregao
                        novo_preco.participante = fornecedor
                        novo_preco.valor = item.replace('.','').replace(',','.')
                        novo_preco.marca = request.POST.getlist('marcas')[idx-1]
                        novo_preco.save()


            lotes = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True)
            if lotes.exists():
                for lote in lotes:
                    itens = ItemLote.objects.filter(lote=lote)
                    if itens.exists():
                        propostas = PropostaItemPregao.objects.filter(item__in=itens.values_list('item', flat=True), participante=fornecedor, pregao=pregao)
                        if propostas.exists():
                            total_propostas = 0
                            for proposta in propostas:
                                total_propostas = total_propostas + proposta.valor * proposta.item.quantidade
                            if PropostaItemPregao.objects.filter(item=lote, pregao=pregao, participante=fornecedor).exists():
                                PropostaItemPregao.objects.filter(item=lote, pregao=pregao, participante=fornecedor).update(valor=total_propostas)
                            else:
                                nova_proposta_para_lote = PropostaItemPregao()
                                nova_proposta_para_lote.pregao = pregao
                                nova_proposta_para_lote.item = lote
                                nova_proposta_para_lote.participante = fornecedor
                                nova_proposta_para_lote.valor = total_propostas
                                nova_proposta_para_lote.save()

            messages.success(request, u'Proposta cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/cadastra_proposta_pregao/%s/?participante=%s&preencheu=True' % (pregao.id, fornecedor.id))
        else:
            if edicao or selecionou:
                if request.GET.get('preencheu'):
                    form = CadastraPrecoParticipantePregaoForm(request.POST or None, pregao=pregao, preencher_box=True, initial=dict(fornecedor=participante))
                else:
                    form = CadastraPrecoParticipantePregaoForm(request.POST or None, pregao=pregao, initial=dict(fornecedor=participante))

            else:
                form = CadastraPrecoParticipantePregaoForm(request.POST or None, pregao=pregao)

        return render(request, 'cadastra_proposta_pregao.html', locals(), RequestContext(request))

    else:
        raise PermissionDenied


@login_required()
def propostas_item(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk= item_id)
    itens = PropostaItemPregao.objects.filter(item=item)
    titulo = u'Valores - Item %s - %s' % (item.item, item.solicitacao.pregao_set.all()[0])
    eh_modalidade_desconto = item.solicitacao.eh_maior_desconto()

    return render(request, 'propostas_item.html', locals(), RequestContext(request))

@login_required()
def propostas_item_credenciamento(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk= item_id)
    pesquisas = ItemPesquisaMercadologica.objects.filter(item=item)

    return render(request, 'propostas_item_credenciamento.html', locals(), RequestContext(request))

@login_required()
def cadastra_participante_pregao(request, pregao_id):

    title=u'Cadastrar Participante do Pregão'
    pregao = get_object_or_404(Pregao, pk= pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        id_sessao = "%s_fornecedor" % (request.user.pessoafisica.id)
        request.session[id_sessao] = pregao.id
        form = CadastraParticipantePregaoForm(request.POST or None, pregao=pregao)
        if form.is_valid():
            if ParticipantePregao.objects.filter(pregao=pregao,fornecedor=form.cleaned_data['fornecedor']).exists():
                messages.error(request, u'Este fornecedor já foi cadastrado.')
                return HttpResponseRedirect(u'/base/pregao/%s/' % pregao.id)

            o = form.save(False)
            o.pregao = pregao
            if form.cleaned_data.get('sem_representante'):
                o.pode_dar_lance = False
                historico = HistoricoPregao()
                historico.pregao = pregao
                historico.data = datetime.datetime.now()
                historico.obs = u'Ausência do participante: %s. Motivo: %s' % (form.cleaned_data['fornecedor'], form.cleaned_data.get('obs_ausencia_participante'))
                historico.save()
            o.save()

            messages.success(request, u'Participante cadastrado com sucesso')
            return HttpResponseRedirect(u'/base/pregao/%s/#fornecedores' % pregao.id)

        return render(request, 'cadastra_participante_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastra_visitante_pregao(request, pregao_id):
    title=u'Cadastrar Visitante do Pregão'
    pregao = get_object_or_404(Pregao, pk= pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = VisitantePregaoForm(request.POST or None)
        if form.is_valid():
            o = form.save(False)
            o.pregao = pregao
            o.save()
            messages.success(request, u'Visitante cadastrado com sucesso')
            return HttpResponseRedirect(u'/base/gerenciar_visitantes/%s/' % pregao.id)

        return render(request, 'cadastra_visitante_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def rodada_pregao(request, item_id):

    item = get_object_or_404(ItemSolicitacaoLicitacao, pk= item_id)
    pregao = get_object_or_404(Pregao, solicitacao=item.solicitacao)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        rodadas = RodadaPregao.objects.filter(item=item).order_by('-rodada')
        if rodadas.exists():
            num_rodada = rodadas[0].rodada + 1

        else:
            num_rodada = 1

        RodadaPregao.objects.filter(item=item).update(atual=False)
        nova_rodada = RodadaPregao()
        nova_rodada.rodada = num_rodada
        nova_rodada.pregao = pregao
        nova_rodada.item = item
        nova_rodada.atual = True
        nova_rodada.save()
        messages.success(request, u'Nova rodada cadastrada.')
        return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
    else:
        raise PermissionDenied

@login_required()
def lances_rodada_pregao(request, rodada_id, item_id):

    rodada = get_object_or_404(RodadaPregao, pk= rodada_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and rodada.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        title=u'Rodada %s do %s' % (rodada.rodada, rodada.pregao)
        lances_rodadas = LanceItemRodadaPregao.objects.filter(rodada=rodada).order_by('item')
        return render(request, 'lances_rodada_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def declinar_lance(request, rodada_id, item_id, participante_id):

    rodada = get_object_or_404(RodadaPregao, pk= rodada_id)
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        participante = get_object_or_404(ParticipantePregao, pk=participante_id)

        novo_lance = LanceItemRodadaPregao()
        novo_lance.item = item
        novo_lance.rodada = rodada
        novo_lance.participante = participante
        novo_lance.declinio = True
        novo_lance.ordem_lance = rodada.get_ordem_lance()
        novo_lance.save()

        messages.success(request, u'Lance declinado com sucesso.')
        return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
    else:
        raise PermissionDenied


@login_required()
def lances_item(request, item_id):

    item = get_object_or_404(ItemSolicitacaoLicitacao, pk= item_id)
    pregao = item.solicitacao.get_pregao()
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        desempatar = False
        botao_incluir = False
        eh_modalidade_desconto = item.solicitacao.eh_maior_desconto()

        fornecedores_lance = PropostaItemPregao.objects.filter(item=item, concorre=True).order_by('-concorre', 'desclassificado','desistencia', 'valor')
        if request.GET.get('empate'):
            desempatar = True
        # if not PropostaItemPregao.objects.filter(item=item).exists():
        #     messages.error(request, u'Este item não possui nenhuma proposta cadastrada.')
        #    return HttpResponseRedirect(u'/base/pregao/%s/#propostas' % item.get_licitacao().id)
        rodadas = RodadaPregao.objects.filter(item=item)
        itens_do_lote = False
        if item.eh_lote:
            itens_do_lote = ItemSolicitacaoLicitacao.objects.filter(id__in=ItemLote.objects.filter(lote=item).values_list('item', flat=True))


        if request.GET and request.GET.get('filtrar') == u'1':
            if item.ja_recebeu_lance():
                messages.info(request,u'Não é possível aplicar filtro após o início dos lances.')
                if desempatar:
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
                else:
                    return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            else:
                if not request.GET.get('incluido') and not request.GET.get('incluir'):
                    item.filtrar_por_10_porcento()
                botao_incluir = True
                if request.GET.get('incluir') == u'1':
                    fornecedores_lance = item.get_lista_participantes(inativos=True)
                else:
                    fornecedores_lance = item.get_lista_participantes(ativos=True)
        elif  request.GET and request.GET.get('filtrar') == u'2':
            if item.ja_recebeu_lance():
                messages.info(request,u'Não é possível aplicar filtro após o início dos lances.')
                if desempatar:
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
                else:
                    return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            else:
                item.filtrar_todos_ativos()
        sorteio = False
        if request.GET and request.GET.get('sorteio') == u'1':
            sorteio = True
            item.gerar_resultado(apaga=False)



        form = LanceForm(request.POST or None)
        if form.is_valid():
            rodada_atual = item.get_rodada_atual()
            tem_empate_beneficio = item.tem_empate_beneficio()
            if desempatar and tem_empate_beneficio:
                participante = item.get_participante_desempate()
            else:
                participante = item.get_proximo_lance() or item.get_participante_desempate()
            rodada_anterior = int(rodada_atual.rodada) - 1
            valor_anterior_registrado = 0
            if PropostaItemPregao.objects.filter(item=item, participante=participante).exists():
                valor_anterior_registrado = PropostaItemPregao.objects.filter(item=item, participante=participante)[0].valor
            if LanceItemRodadaPregao.objects.filter(item=item, participante=participante, rodada__rodada__lt=rodada_atual.rodada, valor__isnull=False).order_by('-rodada__rodada').exists():

                valor_anterior_registrado = LanceItemRodadaPregao.objects.filter(item=item, participante=participante, valor__isnull=False).order_by('-rodada__rodada')[0].valor


            #if not eh_modalidade_desconto:
            if int(rodada_atual.rodada) == 1 and form.cleaned_data.get('lance') >= PropostaItemPregao.objects.filter(item=item, participante=participante)[0].valor:
                messages.error(request, u'Você não pode dar um lance maior do que sua proposta.')

                if tem_empate_beneficio:
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
                else:
                    return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)



            if int(rodada_atual.rodada) > 1 and form.cleaned_data.get('lance') >= valor_anterior_registrado:
                messages.error(request, u'Você não pode dar um lance maior do que o seu último lance registrado: <b>R$: %s</b>.' % format_money(valor_anterior_registrado))

                if tem_empate_beneficio:
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
                else:
                    return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)

            if form.cleaned_data.get('lance') > item.valor_medio:
                messages.error(request, u'Você não pode dar um lance maior do que o valor máximo do item.')
                if tem_empate_beneficio:
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
                else:
                    return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)

            if LanceItemRodadaPregao.objects.filter(item=item, valor=form.cleaned_data.get('lance')):
                messages.error(request, u'Este lance já foi dado.')
                if tem_empate_beneficio:
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
                else:
                    return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)

            if desempatar and tem_empate_beneficio:
                if LanceItemRodadaPregao.objects.filter(item=item, valor__lt=form.cleaned_data.get('lance')).exists():
                    messages.error(request, u'Você não pode dar um lance maior que o menor lance atual.')
                    return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)

            # else:
            #     if int(rodada_atual.rodada) == 1 and form.cleaned_data.get('lance') <= PropostaItemPregao.objects.get(item=item, participante=participante).valor:
            #         messages.error(request, u'Você não pode dar um lance menor do que sua proposta.')
            #         return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            #
            #     if int(rodada_atual.rodada) > 1 and form.cleaned_data.get('lance') >= valor_anterior_registrado:
            #         messages.error(request, u'Você não pode dar um lance maior do que o seu último lance registrado: <b>R$: %s</b>.' % format_money(valor_anterior_registrado))
            #         return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            #
            #     # if form.cleaned_data.get('lance') >= item.valor_medio:
            #     #     messages.error(request, u'Você não pode dar uma lance maior do que o valor máximo do item.')
            #     #     return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            #
            #     if LanceItemRodadaPregao.objects.filter(item=item, valor=form.cleaned_data.get('lance')):
            #         messages.error(request, u'Este lance já foi dado.')
            #         return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            #
            #     if desempatar and item.tem_empate_beneficio():
            #         if LanceItemRodadaPregao.objects.filter(item=item, valor__gt=form.cleaned_data.get('lance')).exists():
            #             messages.error(request, u'Você não pode dar um lance menor que o maior lance atual.')
            #             return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)



            novo_lance = LanceItemRodadaPregao()
            novo_lance.item = item
            novo_lance.rodada = rodada_atual
            novo_lance.participante = participante
            novo_lance.valor = form.cleaned_data.get('lance')
            novo_lance.ordem_lance = rodada_atual.get_ordem_lance()
            novo_lance.save()
            messages.success(request, u'Novo lance cadastrado com sucesso.')

        if item.eh_lote:
            title=u'Lances - Lote: %s' % item.item
        else:
            title=u'Lances - Item: %s' % item
        tabela = {}
        resultado = {}
        lances = LanceItemRodadaPregao.objects.filter(item=item)


        num_rodadas =  rodadas.values('rodada').order_by('rodada').distinct('rodada')
        for num in num_rodadas:
            chave = '%s' % num['rodada']
            tabela[chave] = []
        for lance in lances.order_by('id'):
            chave = '%s' % str(lance.rodada.rodada)
            tabela[chave].append(lance)

        resultado = collections.OrderedDict(sorted(tabela.items(), reverse=True,  key=lambda x: int(x[0])))
        return render(request, 'lances_item.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def ver_fornecedores(request, fornecedor_id=None):
    title=u'Lista de Fornecedores'
    fornecedores = Fornecedor.objects.all().order_by('razao_social')
    form = BuscaFornecedorForm(request.POST or None)
    if form.is_valid():
        fornecedores = fornecedores.filter(Q(razao_social__icontains=form.cleaned_data.get('nome')) | Q(cnpj__icontains=form.cleaned_data.get('nome')))
    exibe_popup = False

    if fornecedor_id:
        fornecedor = get_object_or_404(Fornecedor, pk= fornecedor_id)
        exibe_popup = True

    return render(request, 'ver_fornecedores.html', locals(), RequestContext(request))

@login_required()
def ver_pregoes(request):
    title=u'Licitações'
    pregoes = Pregao.objects.all().order_by('-id')
    eh_ordenador_despesa = False
    if get_config():
        eh_ordenador_despesa = request.user.pessoafisica == get_config().ordenador_despesa

    form = BuscarLicitacaoForm(request.GET or None)

    if form.is_valid():
        if form.cleaned_data.get('info'):
            pregoes = pregoes.filter(Q(solicitacao__processo__numero__icontains=form.cleaned_data.get('info')) | Q(solicitacao__num_memorando__icontains=form.cleaned_data.get('info')) | Q(num_pregao__icontains=form.cleaned_data.get('info')) )

        if form.cleaned_data.get('modalidade'):
            if form.cleaned_data.get('modalidade').id == ModalidadePregao.PREGAO:
                pregoes = pregoes.filter(modalidade__in=[ModalidadePregao.PREGAO, ModalidadePregao.PREGAO_SRP])
            else:
                pregoes = pregoes.filter(modalidade=form.cleaned_data.get('modalidade'))

        if form.cleaned_data.get('data_inicial'):
            pregoes = pregoes.filter(data_abertura__gte=form.cleaned_data.get('data_inicial'))

        if form.cleaned_data.get('data_final'):
            pregoes = pregoes.filter(data_abertura__lte=form.cleaned_data.get('data_final'))

        if form.cleaned_data.get('situacao'):
            pregoes = pregoes.filter(situacao=form.cleaned_data.get('situacao'))

        if form.cleaned_data.get('secretaria'):
            pregoes = pregoes.filter(solicitacao__setor_origem__secretaria=form.cleaned_data.get('secretaria'))

    return render(request, 'ver_pregoes.html', locals(), RequestContext(request))

@login_required()
def itens_solicitacao(request, solicitacao_id):
    config = get_config_geral()
    if config:
        url = config.url

    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    title=u'%s' % (solicitacao)
    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=False).order_by('item')
    eh_lote = False
    contrato = False
    if PedidoContrato.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoContrato.objects.filter(solicitacao=solicitacao).order_by('item')
        contrato = True
    elif PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).order_by('item')
    elif solicitacao.credenciamento_origem:
        pedidos = PedidoCredenciamento.objects.filter(solicitacao=solicitacao).order_by('item')

    eh_lote = solicitacao.eh_lote()
    ja_registrou_preco = ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao, secretaria=request.user.pessoafisica.setor.secretaria)
    ja_aprovou = ja_registrou_preco.filter(aprovado=True).exists()
    setor_do_usuario = request.user.pessoafisica.setor
    recebida_no_setor = solicitacao.recebida_setor(setor_do_usuario)
    pode_gerenciar = solicitacao.recebida_setor(request.user.pessoafisica.setor)
    eh_gerente = request.user.groups.filter(name='Gerente') and pode_gerenciar



    return render(request, 'itens_solicitacao.html', locals(), RequestContext(request))


@login_required()
def cadastrar_item_solicitacao(request, solicitacao_id):
    title=u'Cadastrar Item'
    id_user = '%s' % request.user.pessoafisica.id
    request.session[id_user] = solicitacao_id
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if solicitacao.setor_origem == request.user.pessoafisica.setor and not solicitacao.prazo_aberto:
        form = CadastrarItemSolicitacaoForm(request.POST or None, initial=dict(solicitacao=solicitacao), solicitacao=solicitacao)
        if form.is_valid():
            o = form.save(False)
            o.solicitacao = solicitacao
            o.item = solicitacao.get_proximo_item()
            if o.valor_medio:
                o.total = o.valor_medio * o.quantidade
            o.save()
            novo_item = ItemQuantidadeSecretaria()
            novo_item.solicitacao = solicitacao
            novo_item.item = o
            novo_item.secretaria = request.user.pessoafisica.setor.secretaria
            novo_item.quantidade = o.quantidade
            novo_item.aprovado = True
            novo_item.avaliado_por = request.user
            novo_item.avaliado_em = datetime.datetime.now()
            novo_item.save()

            # send_mail('Subject here', 'Here is the message.', settings.EMAIL_HOST_USER,
            #  ['walkyso@gmail.com'], fail_silently=False)


            messages.success(request, u'Item cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)

        return render(request, 'cadastrar_item_solicitacao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

def baixar_editais(request):
    hoje = datetime.date.today()
    pregoes = Pregao.objects.all().order_by('-id')
    form = BaixarEditaisForm(request.GET or None)
    if form.is_valid():
        if form.cleaned_data.get('modalidade'):
            pregoes = pregoes.filter(modalidade=form.cleaned_data.get('modalidade'))
        if form.cleaned_data.get('numero'):
            pregoes = pregoes.filter(num_pregao__icontains=form.cleaned_data.get('numero'))

    return render(request, 'baixar_editais.html', locals(), RequestContext(request))

def baixar_atas(request):
    hoje = datetime.date.today()
    atas = AtaRegistroPreco.objects.all().order_by('-numero')
    form = BaixarAtasForm(request.GET or None)
    if form.is_valid():
        if form.cleaned_data.get('numero'):
            atas = atas.filter(numero__icontains=form.cleaned_data.get('numero'))

    return render(request, 'baixar_atas.html', locals(), RequestContext(request))


def baixar_contratos(request):
    hoje = datetime.date.today()
    contratos = Contrato.objects.all().order_by('-numero')
    form = BaixarContratoForm(request.POST or None)
    if form.is_valid():
        if form.cleaned_data.get('numero'):
            contratos = contratos.filter(numero__icontains=form.cleaned_data.get('numero'))

    return render(request, 'baixar_contratos.html', locals(), RequestContext(request))

@login_required()
def ver_solicitacoes(request):
    title=u'Lista de Solicitações'
    setor = request.user.pessoafisica.setor
    movimentacoes_setor = MovimentoSolicitacao.objects.filter(Q(setor_origem=setor) | Q(setor_destino=setor))
    solicitacoes = SolicitacaoLicitacao.objects.filter(Q(setor_origem=setor, situacao=SolicitacaoLicitacao.CADASTRADO)  | Q(setor_atual=setor, situacao__in=[SolicitacaoLicitacao.RECEBIDO, SolicitacaoLicitacao.EM_LICITACAO])).order_by('-data_cadastro')
    outras = SolicitacaoLicitacao.objects.filter(Q(id__in=movimentacoes_setor.values_list('solicitacao', flat=True)) | Q(interessados=setor.secretaria)).distinct().order_by('-data_cadastro')
    aba1 = u''
    aba2 = u'in active'
    class_aba1 = u''
    class_aba2 = u'active'
    form = BuscarSolicitacaoForm(request.GET or None)

    if form.is_valid():
        aba1 = u'in active'
        aba2 = u''
        class_aba1 = u'active'
        class_aba2 = u''
        outras = SolicitacaoLicitacao.objects.all()
        if form.cleaned_data.get('info'):
            outras = outras.filter(Q(processo__numero__icontains=form.cleaned_data.get('info')) | Q(num_memorando__icontains=form.cleaned_data.get('info')) | Q(pregao__num_pregao__icontains=form.cleaned_data.get('info')))
        if form.cleaned_data.get('ano'):
           outras = outras.filter(data_cadastro__year=form.cleaned_data.get('ano'))

        if form.cleaned_data.get('secretaria'):
            outras = outras.filter(setor_origem__secretaria=form.cleaned_data.get('secretaria'))

    return render(request, 'ver_solicitacoes.html', locals(), RequestContext(request))

@login_required()
def rejeitar_solicitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and solicitacao.eh_apta() and solicitacao.minuta_aprovada and not solicitacao.tem_pregao_cadastrado():
        title=u'Negar Solicitação %s' % solicitacao.num_memorando
        form = RejeitarSolicitacaoForm(request.POST or None, instance=solicitacao)
        if form.is_valid():
            o = form.save(False)
            o.situacao = SolicitacaoLicitacao.NEGADA
            o.save()
            messages.success(request, u'Solicitação rejeitada com sucesso.')
            return HttpResponseRedirect(u'/base/ver_solicitacoes/')

        return render(request, 'rejeitar_solicitacao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastrar_pregao(request, solicitacao_id):

    title=u'Cadastrar Licitação'
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = PregaoForm(request.POST or None, solicitacao=solicitacao, request=request)
        if form.is_valid():
            form.save()
            solicitacao.situacao = SolicitacaoLicitacao.EM_LICITACAO
            if not solicitacao.processo and form.cleaned_data.get('num_processo'):
                novo_processo = Processo()
                novo_processo.pessoa_cadastro = request.user
                novo_processo.numero = form.cleaned_data.get('num_processo')
                novo_processo.objeto = form.cleaned_data.get('objeto')
                novo_processo.tipo = Processo.TIPO_MEMORANDO
                novo_processo.setor_origem = request.user.pessoafisica.setor
                novo_processo.save()
                solicitacao.processo = novo_processo
            solicitacao.save()
            messages.success(request, u'Licitação cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/' % form.instance.id)

        return render(request, 'cadastrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastrar_solicitacao(request):
    if request.user.has_perm('base.pode_cadastrar_solicitacao'):
        title=u'Cadastrar Solicitação'
        form = SolicitacaoForm(request.POST or None, request=request)
        if form.is_valid():
            o = form.save(False)
            o.setor_origem = request.user.pessoafisica.setor
            o.setor_atual = request.user.pessoafisica.setor
            o.data_cadastro = datetime.datetime.now()

            o.cadastrado_por = request.user
            if not form.cleaned_data['interessados'] and not form.cleaned_data['todos_interessados']:
                o.prazo_resposta_interessados = None
            o.save()

            if form.cleaned_data['todos_interessados']:
                for item in Secretaria.objects.all():
                    o.interessados.add(item)

            elif form.cleaned_data['outros_interessados']:
                form.save_m2m()
            messages.success(request, u'Solicitação cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % form.instance.id)

        return render(request, 'cadastrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def editar_solicitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Editar Solicitação'
        form = SolicitacaoForm(request.POST or None, instance=solicitacao, request=request)
        if form.is_valid():
            o = form.save(False)
            o.setor_origem = request.user.pessoafisica.setor
            o.setor_atual = request.user.pessoafisica.setor
            o.data_cadastro = datetime.datetime.now()
            o.cadastrado_por = request.user
            if not form.cleaned_data.get('interessados'):
                o.prazo_resposta_interessados = None
            o.save()
            if form.cleaned_data.get('interessados') and form.cleaned_data['outros_interessados']:
                form.save_m2m()
            messages.success(request, u'Solicitação cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % form.instance.id)

        return render(request, 'cadastrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def cadastrar_material(request, solicitacao_id):
    title=u'Cadastrar Material'
    form = MaterialConsumoForm(request.POST or None)
    if form.is_valid():
        form.save()
        messages.success(request, u'Material %s cadastrado com sucesso.' % form.instance)
        return HttpResponseRedirect(u'/base/cadastrar_item_solicitacao/%s/' % solicitacao_id)

    return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))

@login_required()
def cadastrar_documento(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    title=u'Cadastrar Documento'
    form = DocumentoSolicitacaoForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        o = form.save(False)
        o.solicitacao = solicitacao
        o.cadastrado_por = request.user
        o.cadastrado_em = datetime.datetime.now()
        o.save()
        messages.success(request, u'Documento cadastrado com sucesso.')
        return HttpResponseRedirect(u'/base/lista_documentos/%s/' % solicitacao_id)

    return render(request, 'cadastrar_documento.html', locals(), RequestContext(request))

@login_required()
def enviar_para_licitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    solicitacao.situacao = SolicitacaoLicitacao.ENVIADO
    solicitacao.save()
    messages.success(request, u'Solicitação enviada com sucesso.')
    return HttpResponseRedirect(u'/base/ver_solicitacoes/')

@login_required()
def registrar_preco_item(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)

    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor) and item.solicitacao.prazo_aberto:
        title = u'Registrar Valor - %s' % item
        form = RegistrarPrecoItemForm(request.POST or None, request.FILES or None, instance=item)
        if form.is_valid():
            o = form.save(False)
            o.total = o.quantidade*o.valor_medio
            o.save()
            messages.success(request, u'Valor registrado com sucesso.')
            return HttpResponseRedirect(u'/itens_solicitacao/%s/' % item.solicitacao.id)
        return render(request, 'registrar_preco_item.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def pesquisa_mercadologica(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    return HttpResponseRedirect(u'/base/planilha_pesquisa_mercadologica/%s/' % solicitacao.id)


def planilha_pesquisa_mercadologica(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao).order_by('item')


    nome = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/modelo_proposta_pesquisa_mercadologica')
    file_path = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/modelo_proposta_pesquisa_mercadologica.xls')
    rb = open_workbook(file_path,formatting_info=True)

    wb = copy(rb) # a writable copy (I can't read values out of this, only write to it)
    w_sheet = wb.get_sheet(0) # the sheet to write to within the writable copy

    sheet = rb.sheet_by_name("Sheet1")
    for idx, item in enumerate(itens, 0):
        row_index = idx + 1
        style = xlwt.XFStyle()
        style.alignment.wrap = 1

        w_sheet.write(row_index, 0, item.item)
        w_sheet.write(row_index, 1, item.material.nome, style)
        w_sheet.write(row_index, 2, item.unidade.nome)


    salvou = nome + u'_%s' % solicitacao.id + '.xls'
    wb.save(salvou)

    arquivo = open(salvou, "rb")


    content_type = 'application/vnd.ms-excel'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    nome_arquivo = salvou.split('/')[-1]
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(salvou)
    return response


def preencher_pesquisa_mercadologica(request, solicitacao_id):
    title = u'Preencher Pesquisa Mercadológica'
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if not solicitacao.prazo_aberto:
        messages.error(request, u'Prazo de envio encerrado.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s' % solicitacao.id)
    if not solicitacao.pode_cadastrar_pesquisa():
        messages.error(request, u'Apenas uma proposta pode ser cadastrada.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s' % solicitacao.id)
    if not request.user.is_authenticated():
        return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/2/' % solicitacao.id)



    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao).order_by('item')

    form = PesquisaMercadologicaForm(request.POST or None, request=request, solicitacao=solicitacao)
    if form.is_valid():

        if form.cleaned_data.get('origem_opcao') is False:
            #o.origem = PesquisaMercadologica.ATA_PRECO
            return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/1/' % solicitacao.id)
        elif form.cleaned_data.get('origem_opcao') is True:
            #o.origem = PesquisaMercadologica.FORNECEDOR
            return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/2/' % solicitacao.id)
        else:
            messages.error(request, u'Selecione a origem.')
            return HttpResponseRedirect(u'/base/preencher_pesquisa_mercadologica/%s/' % solicitacao.id)



        # if form.cleaned_data.get('origem_opcao') is False:
        #     messages.success(request, u'Cadastro realizado com sucesso.')
        #     return HttpResponseRedirect(u'/base/upload_itens_pesquisa_mercadologica/%s/' % o.id)
        #
        # else:
        #     messages.success(request, u'Cadastro realizado com sucesso. Envie a planilha abaixo com os valores dos itens.')
        #     return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/' % o.id)

    return render(request, 'preencher_pesquisa_mercadologica.html', locals(), RequestContext(request))

def preencher_itens_pesquisa_mercadologica(request, solicitacao_id, origem):
    title=u'Preencher Itens da Pesquisa Mercadológica'
    #pesquisa = get_object_or_404(PesquisaMercadologica, pk=pesquisa_id)
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao)
    if origem == u'1':
        texto = u'Ata de Registro de Preço'
    else:
        texto = u'Fornecedor'
    form = ContinuaPesquisaMercadologicaForm(request.POST or None, request.FILES or None, origem=origem)
    form2 = UploadPesquisaForm(request.POST or None, request.FILES or None)
    if request.POST:

        if form.is_valid() and form2.is_valid():
            tem_algum_valor = False
            pesquisa = form.save(False)
            if origem == u'1':
                pesquisa.origem = PesquisaMercadologica.ATA_PRECO
            else:
                pesquisa.origem = PesquisaMercadologica.FORNECEDOR
            pesquisa.solicitacao = solicitacao
            pesquisa.save()
            arquivo_up = form2.cleaned_data.get('arquivo')
            if arquivo_up:
                sheet = None
                try:
                    with transaction.atomic():
                        workbook = xlrd.open_workbook(file_contents=arquivo_up.read())
                        sheet = workbook.sheet_by_index(0)
                except XLRDError:
                    raise Exception(u'Não foi possível processar a planilha. Verfique se o formato do arquivo é .xls ou .xlsx.')

                validade = unicode(sheet.cell_value(7, 1)).strip()
                if not validade:
                    messages.error(request, u'Preencha a validade da proposta na planilha.')
                    return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/%s/' % (solicitacao_id, origem))

                for row in range(10, sheet.nrows):

                    item = unicode(sheet.cell_value(row, 0)).strip()
                    marca = unicode(sheet.cell_value(row, 4)).strip() or None
                    valor = unicode(sheet.cell_value(row, 5)).strip()

                    if item and valor:
                        tem_algum_valor = True
                        item_do_pregao = ItemSolicitacaoLicitacao.objects.get(eh_lote=False, solicitacao=solicitacao,item=int(sheet.cell_value(row, 0)))
                        novo_preco = ItemPesquisaMercadologica()
                        novo_preco.pesquisa = pesquisa
                        novo_preco.item = item_do_pregao
                        try:
                            with transaction.atomic():
                                Decimal(valor)
                        except:
                            messages.error(request, u'o valor %s do %s é inválido.' % (valor, item_do_pregao))
                            return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/%s/' % (solicitacao_id, origem))
                        novo_preco.valor_maximo = valor
                        novo_preco.marca = marca
                        novo_preco.save()
                        pesquisa.validade_proposta = validade
                        pesquisa.save()
                        messages.success(request, u'Valores cadastrados com sucesso.')
                if not tem_algum_valor:
                    pesquisa.delete()
                    messages.error(request, u'Nenhum item possui valor informado.')
                    return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)


                return HttpResponseRedirect(u'/base/upload_pesquisa_mercadologica/%s/' % pesquisa.id)


            if request.POST.get('validade') in [''] and pesquisa.origem == PesquisaMercadologica.FORNECEDOR:
                messages.error(request, u'Preencha a validade da proposta.')
                return HttpResponseRedirect(u'/base/preencher_itens_pesquisa_mercadologica/%s/' % pesquisa.id)
            for idx, item in enumerate(request.POST.getlist('itens'), 1):
                if item:
                    tem_algum_valor = True
                    item_do_pregao = ItemSolicitacaoLicitacao.objects.get(solicitacao=pesquisa.solicitacao, id=request.POST.getlist('id_item')[idx-1])
                    novo_preco = ItemPesquisaMercadologica()
                    novo_preco.pesquisa = pesquisa
                    novo_preco.item = item_do_pregao
                    novo_preco.valor_maximo = item.replace('.','').replace(',','.')
                    novo_preco.marca = request.POST.getlist('marcas')[idx-1]
                    novo_preco.save()
            if request.POST.get('validade'):
                pesquisa.validade_proposta = request.POST.get('validade')
                pesquisa.save()
                if not tem_algum_valor:
                    pesquisa.delete()
                    messages.error(request, u'Nenhum item possui valor informado.')
                    return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)

                messages.success(request, u'Valores cadastrados com sucesso.')
                return HttpResponseRedirect(u'/base/upload_pesquisa_mercadologica/%s/' % pesquisa.id)
            else:
                if not tem_algum_valor:
                    pesquisa.delete()
                    messages.error(request, u'Nenhum item possui valor informado.')
                    return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)
                messages.success(request, u'Valores cadastrados com sucesso.')
                return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % pesquisa.solicitacao.id)


    return render(request, 'preencher_itens_pesquisa_mercadologica.html', locals(), RequestContext(request))


def upload_pesquisa_mercadologica(request, pesquisa_id):
    title=u'Enviar Orçamento da Pesquisa Mercadológica'
    pesquisa = get_object_or_404(PesquisaMercadologica, pk=pesquisa_id)
    form = UploadPesquisaForm(request.POST or None, request.FILES or None, instance=pesquisa)
    if form.is_valid():
        o = form.save(False)
        o.cadastrada_em = datetime.datetime.now()
        o.save()
        messages.success(request, u'Pesquisa cadastrada com sucesso.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % pesquisa.solicitacao.id)

    return render(request, 'upload_pesquisa_mercadologica.html', locals(), RequestContext(request))


def imprimir_pesquisa(request, pesquisa_id):
    pesquisa = get_object_or_404(PesquisaMercadologica, pk=pesquisa_id)

    destino_arquivo = u'upload/pesquisas/rascunhos/%s.pdf' % pesquisa_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'pesquisa': pesquisa, 'data_emissao':data_emissao}

    template = get_template('imprimir_pesquisa.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def ver_pesquisa_mercadologica(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    title=u'Pesquisa Mercadológica'
    pesquisas = ItemPesquisaMercadologica.objects.filter(item=item)
    return render(request, 'ver_pesquisa_mercadologica.html', locals(), RequestContext(request))


@login_required()
def resultado_classificacao(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title = u'Classificação - %s' % item
        lances = ResultadoItemPregao.objects.filter(item=item).order_by('ordem')
        pregao = item.get_licitacao()
        return render(request, 'resultado_classificacao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def desclassificar_do_pregao(request, participante_id):
    title=u'Desclassificar Participante'
    participante = get_object_or_404(ParticipantePregao, pk=participante_id)
    pregao = participante.pregao
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = DesclassificaParticipantePregao(request.POST or None, instance = participante)
        if form.is_valid():
            o = form.save(False)
            o.desclassificado = True
            o.save()
            if PropostaItemPregao.objects.filter(participante=participante).exists():
                PropostaItemPregao.objects.filter(participante=participante).update(concorre=False, desclassificado=True, motivo_desclassificacao=o.motivo_desclassificacao)
            historico = HistoricoPregao()
            historico.pregao = participante.pregao
            historico.data = datetime.datetime.now()
            historico.obs = u'Desclassificação do participante: %s. Motivo: %s' % (participante, form.cleaned_data.get('motivo_desclassificacao'))
            historico.save()
            messages.success(request, u'Desclassificação cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/' % participante.pregao.id)
        return render(request, 'desclassificar_do_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def planilha_propostas(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    pregao = get_object_or_404(Pregao, solicitacao=solicitacao)
    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=False).order_by('item')

    nome = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/modelo_proposta_fornecedor')
    file_path = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/modelo_proposta_fornecedor.xls')
    rb = open_workbook(file_path,formatting_info=True)

    wb = copy(rb) # a writable copy (I can't read values out of this, only write to it)
    w_sheet = wb.get_sheet(0) # the sheet to write to within the writable copy

    sheet = rb.sheet_by_name("Sheet1")
    w_sheet.write(1, 1, pregao.get_titulo())
    w_sheet.write(2, 1, pregao.objeto)
    w_sheet.write(3, 1, pregao.get_local())

    for idx, item in enumerate(itens, 0):
        row_index = idx + 9
        style = xlwt.XFStyle()
        style.alignment.wrap = 1

        w_sheet.write(row_index, 0, item.__unicode__()[5:])
        w_sheet.write(row_index, 1, item.material.nome, style)
        w_sheet.write(row_index, 2, item.unidade.nome)
        w_sheet.write(row_index, 3, item.quantidade)
        w_sheet.write(row_index, 4, item.valor_medio)

    salvou = nome + u'_%s' % pregao.id + '.xls'
    wb.save(salvou)

    arquivo = open(salvou, "rb")


    content_type = 'application/vnd.ms-excel'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    nome_arquivo = salvou.split('/')[-1]
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(salvou)
    return response

def planilha_propostas_solicitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)

    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=False).order_by('item')

    nome = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/planilha_propostas_solicitacao')
    file_path = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/planilha_propostas_solicitacao.xls')
    rb = open_workbook(file_path,formatting_info=True)

    wb = copy(rb) # a writable copy (I can't read values out of this, only write to it)
    w_sheet = wb.get_sheet(0) # the sheet to write to within the writable copy

    sheet = rb.sheet_by_name("Sheet1")
    w_sheet.write(1, 1, u'Solicitação %s - %s' % (solicitacao.num_memorando, solicitacao.setor_origem))
    w_sheet.write(2, 1, solicitacao.objeto)



    for idx, item in enumerate(itens, 0):
        row_index = idx + 10
        style = xlwt.XFStyle()
        style.alignment.wrap = 1

        w_sheet.write(row_index, 0, item.item)
        w_sheet.write(row_index, 1, item.material.nome, style)
        w_sheet.write(row_index, 2, item.unidade.nome)
        w_sheet.write(row_index, 3, item.quantidade)

    salvou = nome + u'_%s' % solicitacao.id + '.xls'
    wb.save(salvou)

    arquivo = open(salvou, "rb")


    content_type = 'application/vnd.ms-excel'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    nome_arquivo = salvou.split('/')[-1]
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(salvou)
    return response


@login_required()
def remover_participante(request, proposta_id, situacao):
    proposta = get_object_or_404(PropostaItemPregao, pk=proposta_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and proposta.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        if situacao == u'1':
            title=u'Registrar Desclassificação'
        elif situacao == u'2':
            title=u'Registrar Desistências'

        form = RemoverParticipanteForm(request.POST or None)
        if form.is_valid():
            if situacao in [u'1', u'3']:
                proposta.desclassificado = True
                proposta.motivo_desclassificacao = form.cleaned_data.get('motivo')

            elif situacao == u'2':
                proposta.desistencia = True
                proposta.motivo_desistencia = form.cleaned_data.get('motivo')

            proposta.concorre = False
            proposta.save()

            historico = HistoricoPregao()
            historico.pregao = proposta.pregao
            historico.data = datetime.datetime.now()
            if situacao in [u'1', u'3']:
                historico.obs = u'Desclassificação do participante: %s do Item: %s. Motivo: %s' % (proposta.participante, proposta.item.item, form.cleaned_data.get('motivo'))
            elif situacao == u'2':
                historico.obs = u'Desistência do participante: %s do Item: %s. Motivo: %s' % (proposta.participante, proposta.item.item, form.cleaned_data.get('motivo'))
            historico.save()
            messages.success(request, u'Desistência/Desclassificação registrada com sucesso.')
            if situacao == u'3':
                return HttpResponseRedirect(u'/base/cadastra_proposta_pregao/%s/?participante=%s' % (proposta.pregao.id, proposta.participante.id))
            else:
                return HttpResponseRedirect(u'/base/lances_item/%s/' % proposta.item.id)
        return render(request, 'remover_participante.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def adicionar_por_discricionaridade(request, proposta_id):
    proposta = get_object_or_404(PropostaItemPregao, pk=proposta_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and proposta.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        proposta.concorre = True
        proposta.save()
        messages.success(request, u'Participante adicionado.')
        return HttpResponseRedirect(u'/base/lances_item/%s/?filtrar=1&incluido=1' % proposta.item.id)
    else:
        raise PermissionDenied

@login_required()
def gerar_resultado(request, item_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
        item.gerar_resultado()
        messages.success(request, u'Resultados gerados com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % item.get_licitacao().id)
    else:
        raise PermissionDenied


@login_required()
def resultado_alterar(request, resultado_id, situacao):
    title=u'Alterar Situação de Fornecedor'
    resultado = get_object_or_404(ResultadoItemPregao, pk=resultado_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and resultado.item.get_licitacao().solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = ResultadoObsForm(request.POST or None, instance=resultado)
        if form.is_valid():
            if situacao ==u'1':
                resultado.situacao = ResultadoItemPregao.INABILITADO

            elif situacao == u'2':
                resultado.situacao = ResultadoItemPregao.DESCLASSIFICADO
            elif situacao == u'3':
                resultado.situacao = ResultadoItemPregao.CLASSIFICADO
            form.save()

            historico = HistoricoPregao()
            historico.pregao = resultado.item.get_licitacao()
            historico.data = datetime.datetime.now()
            if situacao == u'1':
                historico.obs = u'Inabilitação do participante: %s do Item: %s. Motivo: %s' % (resultado.participante, resultado.item.item, form.cleaned_data.get('observacoes'))
            elif situacao == u'2':
                historico.obs = u'Desclassificação do participante: %s do Item: %s. Motivo: %s' % (resultado.participante, resultado.item.item, form.cleaned_data.get('observacoes'))
            elif situacao == u'3':
                historico.obs = u'Classificação do participante: %s do Item: %s. Motivo: %s' % (resultado.participante, resultado.item.item, form.cleaned_data.get('observacoes'))
            historico.save()
            messages.success(request, u'Situação alterada com sucesso.')
            return HttpResponseRedirect(u'/base/resultado_classificacao/%s/' % resultado.item.id)
        return render(request, 'resultado_alterar.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def resultado_alterar_todos(request, pregao_id, participante_id, situacao):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        participante = get_object_or_404(ParticipantePregao, pk=participante_id)
        title=u'Alterar Participante'
        form = RemoverParticipanteForm(request.POST or None)
        if form.is_valid():
            ids_itens_ganhador = list()
            for item_resultado in ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, participante=participante):
                if item_resultado.ganhador_atual() and item_resultado.ganhador_atual() == participante:
                    ids_itens_ganhador.append(item_resultado.item.id)

            if situacao ==u'1':
                ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, participante=participante).update(situacao=ResultadoItemPregao.INABILITADO, observacoes=form.cleaned_data.get('motivo'))


            elif situacao == u'2':
                ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, participante=participante).update(situacao=ResultadoItemPregao.DESCLASSIFICADO, observacoes=form.cleaned_data.get('motivo'))

            elif situacao == u'3':
                ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, participante=participante).update(situacao=ResultadoItemPregao.CLASSIFICADO, observacoes=form.cleaned_data.get('motivo'))

            historico = HistoricoPregao()
            historico.pregao = pregao
            historico.data = datetime.datetime.now()
            if situacao == u'1':
                historico.obs = u'Inabilitação do participante: %s de todos os itens. Motivo: %s' % (participante, form.cleaned_data.get('motivo'))
            elif situacao == u'2':
                historico.obs = u'Desclassificação do participante: %s de todos os itens. Motivo: %s' % (participante, form.cleaned_data.get('motivo'))
            elif situacao == u'3':
                historico.obs = u'Reintegração do participante: %s em todos os itens. Motivo: %s' % (participante, form.cleaned_data.get('motivo'))

            historico.save()
            if situacao == u'3':
                participante.excluido_dos_itens = False
            else:
                participante.excluido_dos_itens = True
            participante.save()

            if ids_itens_ganhador:
                for ids in ids_itens_ganhador:
                    if ResultadoItemPregao.objects.filter(item__id=ids, situacao=ResultadoItemPregao.CLASSIFICADO).count() >= 2:
                        if ResultadoItemPregao.objects.filter(item__id=ids, situacao=ResultadoItemPregao.CLASSIFICADO)[0].valor == ResultadoItemPregao.objects.filter(item__id=ids, situacao=ResultadoItemPregao.CLASSIFICADO)[1].valor:
                            registro = ResultadoItemPregao.objects.filter(item__id=ids, situacao=ResultadoItemPregao.CLASSIFICADO)[0]
                            registro.empate = True
                            registro.save()

            return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao.id)
        return render(request, 'encerrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def resultado_ajustar_preco(request, resultado_id):
    title=u'Ajustar Preço de Fornecedor'
    resultado = get_object_or_404(ResultadoItemPregao, pk=resultado_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and resultado.item.get_licitacao().solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = ResultadoAjustePrecoForm(request.POST or None, instance=resultado)
        if form.is_valid():
            form.save()
            if resultado.item.get_licitacao().eh_pregao():
                messages.success(request, u'Situação alterada com sucesso.')
                return HttpResponseRedirect(u'/base/resultado_classificacao/%s/' % resultado.item.id)
            else:
                PropostaItemPregao.objects.filter(item=resultado.item, participante=resultado.participante, marca=resultado.marca).update(valor=resultado.valor)
                return HttpResponseRedirect(u'/base/gerar_resultado_licitacao/%s/' % resultado.item.get_licitacao().id)
        return render(request, 'resultado_ajustar_preco.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def desempatar_item(request, item_id):
    title=u'Desempatar Item'
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    eh_credenciamento = item.solicitacao.eh_credenciamento()
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        resultados = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO).order_by('ordem')
        pregao = item.get_licitacao()

        return render(request, 'desempatar_item.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def definir_colocacao(request, resultado_id):
    resultado = get_object_or_404(ResultadoItemPregao, pk=resultado_id)
    title = u'Informar Colocação do Participante'
    if request.user.has_perm('base.pode_cadastrar_pregao') and resultado.item.get_licitacao().solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = DefinirColocacaoForm(request.POST or None, instance=resultado)
        if form.is_valid():
            o = form.save(False)
            o.empate = False
            o.save()
            messages.success(request, u'Colocação registrada com sucesso.')
            return HttpResponseRedirect(u'/base/desempatar_item/%s/' % resultado.item.id)
        return render(request, 'definir_colocacao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def movimentar_solicitacao(request, solicitacao_id, tipo):
    title=u'Enviar Solicitação'
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if tipo ==u'3':
        form = SetorEnvioForm(request.POST or None, devolve=True, setor_atual=solicitacao.setor_atual)
    else:
        form = SetorEnvioForm(request.POST or None, devolve=False, setor_atual=solicitacao.setor_atual)
    if form.is_valid():

        if tipo ==u'3':
            solicitacao.situacao = SolicitacaoLicitacao.DEVOLVIDO
            solicitacao.setor_atual = solicitacao.setor_origem
        else:
            solicitacao.situacao = SolicitacaoLicitacao.ENVIADO
            solicitacao.setor_atual = form.cleaned_data.get('setor')

        solicitacao.prazo_aberto = False
        solicitacao.save()
        nova_movimentacao = MovimentoSolicitacao()
        nova_movimentacao.solicitacao = solicitacao
        nova_movimentacao.setor_origem = request.user.pessoafisica.setor
        nova_movimentacao.data_envio = datetime.datetime.now()
        nova_movimentacao.enviado_por = request.user
        nova_movimentacao.setor_destino = form.cleaned_data.get('setor')
        nova_movimentacao.obs = form.cleaned_data.get('obs')
        nova_movimentacao.save()
        messages.success(request, u'Solicitação enviada com sucesso.')
        return HttpResponseRedirect(u'/base/ver_solicitacoes/')
    return render(request, 'movimentar_solicitacao.html', locals(), RequestContext(request))


@login_required()
def cadastrar_anexo_pregao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Cadastrar Anexo - %s' % pregao
        form = AnexoPregaoForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            o = form.save(False)
            o.pregao = pregao
            o.cadastrado_por = request.user
            o.cadastrado_em = datetime.datetime.now()
            o.save()

            if form.cleaned_data.get('publico') and form.cleaned_data.get('enviar_email'):
                registros = LogDownloadArquivo.objects.filter(arquivo__pregao=pregao).distinct('email')
                config = get_config_geral()
                arquivo_nome = u'\'%s\' - %s' % (o.nome, pregao)
                link = config.url + u'/base/baixar_editais/'

                for registro in registros:
                    texto = u'Olá, %s. O arquivo %s foi adicionado no portal da transparência da %s. Endereço para visualização: %s ' % (registro.responsavel, arquivo_nome, config.nome, link)
                    send_mail('Easy Gestão Pública - Novo Arquivo Cadastrado', texto, settings.EMAIL_HOST_USER, [registro.email], fail_silently=True)

            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#anexos' % pregao.id)

        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastrar_anexo_contrato(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)
    if request.user.has_perm('base.pode_gerenciar_contrato'):
        title=u'Cadastrar Anexo - %s' % contrato
        form = AnexoContratoForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            o = form.save(False)
            o.contrato = contrato
            o.cadastrado_por = request.user
            o.cadastrado_em = datetime.datetime.now()
            o.save()
            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_contrato/%s/#anexos' % contrato.id)

        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def cadastrar_anexo_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)
    if request.user.has_perm('base.pode_gerenciar_contrato'):
        title=u'Cadastrar Anexo - %s' % credenciamento
        form = AnexoCredenciamentoForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            o = form.save(False)
            o.credenciamento = credenciamento
            o.cadastrado_por = request.user
            o.cadastrado_em = datetime.datetime.now()
            o.save()
            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/#anexos' % credenciamento.id)

        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@transaction.atomic()
@login_required()
def cadastrar_ata_registro_preco(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    pregao = solicitacao.get_pregao()
    if request.user.has_perm('base.pode_gerenciar_contrato') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and pregao and pregao.data_homologacao and pregao.eh_ata_registro_preco and not solicitacao.get_ata():

        title=u'Cadastrar Ata de Registro de Preço'

        form = AtaRegistroPrecoForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            o = form.save(False)
            o.solicitacao = solicitacao
            o.pregao = pregao
            o.valor = pregao.get_valor_total()
            o.save()
            if solicitacao.eh_lote():
                itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('item')
                for lote in itens_pregao:
                    for item in lote.get_itens_do_lote():
                        novo_item = ItemAtaRegistroPreco()
                        novo_item.ata = o
                        novo_item.item = item
                        novo_item.marca = item.get_marca_item_lote()
                        novo_item.participante = lote.get_vencedor().participante
                        novo_item.valor = item.get_valor_unitario_final()
                        novo_item.quantidade = item.quantidade
                        novo_item.ordem = o.get_ordem()
                        novo_item.unidade = item.unidade
                        novo_item.material = item.material
                        novo_item.fornecedor = lote.get_vencedor().participante.fornecedor
                        novo_item.save()

            else:
                for resultado in solicitacao.get_resultado():
                    novo_item = ItemAtaRegistroPreco()
                    novo_item.ata = o
                    novo_item.item = resultado.item
                    novo_item.marca = resultado.marca
                    novo_item.participante = resultado.participante
                    novo_item.valor = resultado.valor
                    novo_item.ordem = o.get_ordem()
                    novo_item.quantidade = resultado.item.quantidade
                    novo_item.unidade = resultado.item.unidade
                    novo_item.material = resultado.item.material
                    novo_item.fornecedor = resultado.participante.fornecedor
                    novo_item.save()


            messages.success(request, u'Ata de Registro de Preço cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % o.id)
        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastrar_contrato(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    pregao = solicitacao.get_pregao()

    if request.user.has_perm('base.pode_gerenciar_contrato') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and not solicitacao.contrato_set.exists() and not solicitacao.ataregistropreco_set.exists():

        title=u'Cadastrar Contrato'

        if pregao:
            form = CriarContratoForm(request.POST or None, request.FILES or None, pregao=pregao)
        else:
            form = ContratoForm(request.POST or None, request.FILES or None)


        if form.is_valid():


            if pregao:

                for participante in pregao.get_vencedores():

                    o = Contrato()
                    o.solicitacao = solicitacao
                    o.pregao = pregao
                    o.valor = pregao.get_valor_total(participante)
                    o.numero = form.cleaned_data.get('contrato_%d' % participante.id)
                    o.aplicacao_artigo_57 = form.cleaned_data.get('aplicacao_artigo_57_%d' % participante.id)
                    o.garantia_execucao_objeto = form.cleaned_data.get('garantia_%d' % participante.id)

                    o.data_inicio = form.cleaned_data.get('data_inicial_%d' % participante.id)
                    o.data_fim = form.cleaned_data.get('data_final_%d' % participante.id)
                    o.save()

                    if solicitacao.eh_lote():
                        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('item')
                        for lote in itens_pregao:
                            for item in lote.get_itens_do_lote():
                                if lote.get_vencedor().participante == participante:
                                    novo_item = ItemContrato()
                                    novo_item.contrato = o
                                    novo_item.ordem = o.get_ordem()
                                    novo_item.item = item
                                    novo_item.material = item.material
                                    if item.get_marca_item_lote():
                                        novo_item.marca = item.get_marca_item_lote()
                                    novo_item.participante = lote.get_vencedor().participante
                                    novo_item.fornecedor = lote.get_vencedor().participante.fornecedor
                                    novo_item.valor = item.get_valor_unitario_final()
                                    novo_item.quantidade = item.quantidade
                                    novo_item.unidade = item.unidade
                                    novo_item.save()
                    else:
                        for resultado in solicitacao.get_resultado():
                            if resultado.participante == participante:
                                novo_item = ItemContrato()
                                novo_item.contrato = o
                                novo_item.ordem = o.get_ordem()
                                novo_item.item = resultado.item
                                novo_item.material = resultado.item.material
                                if resultado.marca:
                                    novo_item.marca = resultado.marca
                                novo_item.participante = resultado.participante
                                novo_item.fornecedor = resultado.participante.fornecedor
                                novo_item.valor = resultado.valor
                                novo_item.quantidade = resultado.item.quantidade
                                novo_item.unidade = resultado.item.unidade
                                novo_item.save()
            else:
                o = form.save(False)
                o.solicitacao = solicitacao
                o.valor = solicitacao.get_valor_total()
                o.save()


                if PedidoContrato.objects.filter(solicitacao=solicitacao).exists():
                    for resultado in PedidoContrato.objects.filter(solicitacao=solicitacao):
                        novo_item = ItemContrato()
                        novo_item.contrato = o
                        novo_item.ordem = o.get_ordem()
                        if resultado.item.item:
                            novo_item.item = resultado.item.item
                        novo_item.material = resultado.item.material
                        novo_item.marca = resultado.item.marca
                        if resultado.item.participante:
                            novo_item.participante = resultado.item.participante
                        novo_item.fornecedor = resultado.item.fornecedor
                        novo_item.valor = resultado.item.valor
                        novo_item.quantidade = resultado.quantidade
                        novo_item.unidade = resultado.item.unidade
                        novo_item.save()

                elif PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).exists():
                    for resultado in PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao):
                        novo_item = ItemContrato()
                        novo_item.contrato = o
                        novo_item.ordem = o.get_ordem()
                        if resultado.item.item:
                            novo_item.item = resultado.item.item
                        novo_item.material = resultado.item.material

                        novo_item.marca = resultado.item.marca
                        if resultado.item.participante:
                            novo_item.participante = resultado.item.participante
                        novo_item.fornecedor = resultado.item.fornecedor
                        novo_item.valor = resultado.item.valor
                        novo_item.quantidade = resultado.quantidade
                        novo_item.unidade = resultado.item.unidade
                        novo_item.save()

                elif PedidoCredenciamento.objects.filter(solicitacao=solicitacao).exists():
                    for resultado in PedidoCredenciamento.objects.filter(solicitacao=solicitacao):
                        novo_item = ItemContrato()
                        novo_item.contrato = o
                        novo_item.ordem = o.get_ordem()
                        if resultado.item.item:
                            novo_item.item = resultado.item.item
                        novo_item.material = resultado.item.material

                        novo_item.marca = resultado.item.marca
                        novo_item.fornecedor = resultado.fornecedor
                        novo_item.valor = resultado.item.valor
                        novo_item.quantidade = resultado.quantidade
                        novo_item.unidade = resultado.item.unidade
                        novo_item.save()

                    o.pregao = solicitacao.credenciamento_origem.pregao
                    o.save()
                else:
                    lista = list()
                    dicionario = {}
                    for pesquisa in PesquisaMercadologica.objects.filter(solicitacao=solicitacao):
                        total = ItemPesquisaMercadologica.objects.filter(pesquisa=pesquisa, ativo=True).aggregate(soma=Sum('valor_maximo'))['soma']
                        if total:
                            lista.append([pesquisa.id, total])
                            dicionario[pesquisa.id] = total
                    resultado = sorted(dicionario.items(), key=lambda x: x[1])
                    pesquisa = PesquisaMercadologica.objects.get(id=resultado[0][0])
                    if Fornecedor.objects.filter(cnpj=pesquisa.cnpj):
                        fornecedor = Fornecedor.objects.filter(cnpj=pesquisa.cnpj)[0]
                    else:
                        fornecedor = Fornecedor()
                        fornecedor.cnpj = pesquisa.cnpj
                        fornecedor.razao_social = pesquisa.razao_social
                        fornecedor.endereco = pesquisa.endereco
                        fornecedor.telefones = pesquisa.telefone
                        fornecedor.email = pesquisa.email
                        fornecedor.save()
                    for item in ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao):
                        registro = ItemPesquisaMercadologica.objects.filter(pesquisa=pesquisa, item=item)[0]
                        novo_item = ItemContrato()
                        novo_item.contrato = o
                        novo_item.ordem = o.get_ordem()
                        novo_item.item = item
                        novo_item.material = item.material

                        novo_item.marca = registro.marca
                        novo_item.fornecedor = fornecedor
                        novo_item.valor = registro.valor_maximo
                        novo_item.quantidade = item.quantidade
                        novo_item.unidade = item.unidade
                        novo_item.save()


            messages.success(request, u'Cadastrado realizado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' % o.id)

        return render(request, 'cadastrar_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def cadastrar_credenciamento(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    pregao = solicitacao.get_pregao()
    if request.user.has_perm('base.pode_gerenciar_contrato') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and pregao and pregao.data_homologacao and solicitacao.eh_credenciamento():

        title=u'Cadastrar Credenciamento'

        form = CredenciamentoForm(request.POST or None)
        if form.is_valid():
            o = form.save(False)
            o.solicitacao = solicitacao
            o.pregao = pregao
            o.valor = pregao.get_valor_total()
            o.save()

            for resultado in solicitacao.get_resultado():
                novo_item = ItemCredenciamento()
                novo_item.credenciamento = o
                novo_item.ordem = o.get_ordem()
                novo_item.item = resultado.item
                novo_item.marca = resultado.marca
                novo_item.valor = resultado.valor
                novo_item.quantidade = resultado.item.quantidade
                novo_item.unidade = resultado.item.unidade
                novo_item.material = resultado.item.material
                novo_item.save()


            messages.success(request, u'Credenciamento cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/' % o.id)
        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

def baixar_arquivo(request, arquivo_id):

    arquivo = get_object_or_404(AnexoPregao, pk=arquivo_id)
    title=u'Portal da Transparência'
    subtitulo = u'Baixar Arquivo:  \'%s\' - %s' % (arquivo.nome, arquivo.pregao)
    form = LogDownloadArquivoForm(request.POST or None)
    if form.is_valid():
        o = form.save(False)
        o.arquivo = arquivo
        o.save()
        arquivo_nome = u'\'%s\' - %s' % (arquivo.nome, arquivo.pregao)
        link = settings.SITE_URL + u'/media/%s' % arquivo.arquivo
        texto = u'Olá, %s. Segue o link para download do arquivo: %s. Link: %s ' % (o.nome, arquivo_nome, link)
        send_mail('Easy Gestão Pública - Download do Arquivo', texto, settings.EMAIL_HOST_USER,
             [o.email], fail_silently=False)
        messages.success(request, u'O link para download do arquivo foi enviado para seu email.')
        return HttpResponseRedirect(u'/base/baixar_editais/', )
        return HttpResponseRedirect(u'/media/%s' % arquivo.arquivo )
    return render(request, 'baixar_arquivo.html', locals(), RequestContext(request))


def busca_pessoa(request):

    from django.core import serializers
    if request.method == 'GET':
        pessoa = request.GET.get('pessoa')
        medicos = LogDownloadArquivo.objects.filter(Q(cpf=pessoa) | Q(cnpj=pessoa)).order_by('-id')
        if medicos:
            data = serializers.serialize('json', list(medicos), fields=('nome','cpf', 'municipio', 'estado', 'municipio__estado', 'responsavel', 'cnpj', 'endereco', 'municipio', 'telefone', 'email',))
        else:
            data = []

        return HttpResponse(data, content_type='application/json')
@login_required()
def alterar_valor_lance(request, lance_id):
    lance = get_object_or_404(LanceItemRodadaPregao, pk=lance_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and lance.rodada.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = AlteraLanceForm(request.POST or None, instance=lance)
        if form.is_valid():
            rodada_atual = lance.item.get_rodada_atual()
            participante = lance.participante
            item = lance.item
            rodada_anterior = int(rodada_atual.rodada) - 1
            if int(rodada_atual.rodada) == 1 and form.cleaned_data.get('valor') >= PropostaItemPregao.objects.get(item=item, participante=participante).valor:
                messages.error(request, u'Você não pode dar uma lance maior do que sua proposta.')
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)

            if int(rodada_atual.rodada) > 1 and form.cleaned_data.get('valor') >= LanceItemRodadaPregao.objects.get(item=item, participante=participante, rodada__rodada=rodada_anterior).valor:
                messages.error(request, u'Você não pode dar uma lance maior do que o lance da rodada anterior.')
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)

            if form.cleaned_data.get('valor') >= item.valor_medio:
                messages.error(request, u'Você não pode dar uma lance maior do que o valor máximo do item.')
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)

            if LanceItemRodadaPregao.objects.filter(item=item, valor=form.cleaned_data.get('valor')):
                messages.error(request, u'Este lance já foi dado.')
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)



            form.save()
            messages.success(request, u'Lance alterado com sucesso.')
            return HttpResponseRedirect(u'/base/lances_item/%s/' % lance.item.id)
        return render(request, 'alterar_valor_lance.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def avancar_proximo_item(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        if item.faltou_lance_participante():
            messages.error(request, u'A empresa %s não registrou lance na rodada atual nem declinou em nenhuma das rodadas.' % item.faltou_lance_participante())
            return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
        if item.tem_empate_beneficio():
            messages.error(request, u'Este item tem um empate.')
            return HttpResponseRedirect(u'/base/lances_item/%s/?empate=True' % item.id)
        else:
            if item.eh_ativo():
                item.gerar_resultado()
            #item.ativo=False
            #item.save()
            if request.GET.get('ultimo'):
                return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % item.get_licitacao().id)
            else:
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.tem_proximo_item())
    else:
        raise PermissionDenied


@login_required()
def cancelar_rodada(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        rodada = item.get_rodada_atual()
        LanceItemRodadaPregao.objects.filter(rodada=rodada, item=item).delete()
        rodada_anterior = rodada.get_rodada_anterior()
        if rodada_anterior:
            rodada_anterior.atual=True
            rodada_anterior.save()
        rodada.delete()
        messages.success(request, u'Rodada cancelada.')
        return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
    else:
        raise PermissionDenied

@login_required()
def editar_proposta(request, proposta_id):
    proposta = get_object_or_404(PropostaItemPregao, pk=proposta_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and proposta.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Editar Item da Proposta'
        form = EditarPropostaForm(request.POST or None, instance=proposta)
        if form.is_valid():
            form.save()
            messages.success(request, u'Item da proposta atualizada com sucesso.')
            return HttpResponseRedirect(u'/base/cadastra_proposta_pregao/%s/?participante=%s' % (proposta.pregao.id, proposta.participante.id))

        return render(request, 'editar_proposta.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def encerrar_pregao(request, pregao_id, motivo):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if (request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor)) or request.user.is_superuser:
        title=u'Alterar Situação - %s' % pregao

        deserta = False
        if motivo == u'1':
            deserta = True
        form = EncerrarPregaoForm(request.POST or None, instance=pregao, deserta=deserta)
        if form.is_valid():
            o = form.save(False)
            if motivo == u'1':
                if form.cleaned_data.get('republicar'):
                    o.data_abertura = form.cleaned_data.get('data')
                    o.hora_abertura = form.cleaned_data.get('hora')
                    o.republicado = True
                else:
                    o.situacao = Pregao.DESERTO
            elif motivo == u'2':
                o.situacao = Pregao.FRACASSADO
            elif motivo == u'3':
                o.situacao = Pregao.CADASTRADO
            o.save()
            historico = HistoricoPregao()
            historico.pregao = pregao
            historico.data = datetime.datetime.now()
            if motivo == u'1':
                historico.obs = u'Pregão Deserto. Observações: %s' %  form.cleaned_data.get('obs')
            elif motivo == u'2':
                historico.obs = u'Pregão Fracassado. Observações: %s' %  form.cleaned_data.get('obs')
            elif motivo == u'3':
                historico.obs = u'Pregão Reativado pelo Administrador. Observações: %s' %  form.cleaned_data.get('obs')
            historico.save()
            messages.success(request, u'Pregão atualizado com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#fornecedores' % pregao.id)

        return render(request, 'encerrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def encerrar_itempregao(request, item_id, motivo, origem):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Alterar Situação do Item %s' % item
        form = EncerrarItemPregaoForm(request.POST or None, instance=item)
        if form.is_valid():
            o = form.save(False)
            if motivo == u'1':
                o.situacao = ItemSolicitacaoLicitacao.DESERTO
                historico = HistoricoPregao()
                historico.pregao = item.get_licitacao()
                historico.data = datetime.datetime.now()
                historico.obs = u'Item Deserto: %s. Motivo: %s' % (item, form.cleaned_data.get('obs'))
                historico.save()
            elif motivo == u'2':
                o.situacao = ItemSolicitacaoLicitacao.FRACASSADO
                historico = HistoricoPregao()
                historico.pregao = item.get_licitacao()
                historico.data = datetime.datetime.now()
                historico.obs = u'Item Fracassado: %s. Motivo: %s' % (item, form.cleaned_data.get('obs'))
                historico.save()
            o.save()
            messages.success(request, u'Item atualizado com sucesso.')
            if origem ==u'1':
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            else:
                return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % item.get_licitacao().id)


        return render(request, 'encerrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def suspender_pregao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Suspender - %s' % pregao
        if 'retomar' in request.GET:
            pregao.situacao = Pregao.CADASTRADO
            pregao.save()
            registro = HistoricoPregao()
            registro.data =datetime.datetime.now()
            registro.obs = u'Pregão Retomado'
            registro.pregao = pregao
            registro.save()
            messages.success(request, u'Pregão retomado com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#fornecedores' % pregao.id)

        form = SuspenderPregaoForm(request.POST or None)
        if form.is_valid():
            registro = HistoricoPregao()
            registro.data = datetime.datetime.now()
            registro.obs = form.cleaned_data.get('motivo')
            registro.pregao = pregao
            registro.save()
            pregao.categoria_suspensao = form.cleaned_data.get('categoria_suspensao')
            pregao.situacao = Pregao.SUSPENSO
            pregao.data_suspensao = datetime.datetime.now().date()
            if form.cleaned_data.get('sine_die'):
                pregao.sine_die = True
                pregao.data_retorno = None
            else:
                pregao.sine_die = False
                pregao.data_retorno = form.cleaned_data.get('data_retorno')
                pregao.data_abertura = form.cleaned_data.get('data_retorno')
                pregao.hora_abertura = form.cleaned_data.get('hora_retorno')
            pregao.save()

            messages.success(request, u'Pregão suspenso com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#fornecedores' % pregao.id)

        return render(request, 'encerrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def prazo_pesquisa_mercadologica(request, solicitacao_id):
    title=u'Prazo de recebimento de pesquisas'
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and (solicitacao.prazo_aberto or not solicitacao.eh_dispensa() or not solicitacao.tem_ordem_compra()):

        if solicitacao.prazo_aberto:
            solicitacao.prazo_aberto = False
            solicitacao.liberada_para_pedido = False
            solicitacao.save()
            messages.success(request, u'Período para recebimento de pesquisa fechado com sucesso.')
        else:
            solicitacao.prazo_aberto = True
            solicitacao.save()
            messages.success(request, u'Período para recebimento de pesquisa aberto com sucesso.')

        return HttpResponseRedirect(u'/base/itens_solicitacao/%s' % solicitacao.id)

        return render(request, 'encerrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def modelo_memorando(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    texto = u'teste de geracao de doc<br><br>Solicitação: %s <br>Objetivo: %s<br>Justificativa: %s' % (solicitacao, solicitacao.objetivo, solicitacao.justificativa)
    response = HttpResponse(texto, content_type='application/vnd.ms-word')
    response['Content-Disposition'] = 'attachment; filename=file.doc'
    return response


@login_required()
def receber_solicitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    movimento = MovimentoSolicitacao.objects.filter(solicitacao=solicitacao, data_recebimento__isnull=True)
    if movimento.exists():
        atualiza_registro = MovimentoSolicitacao.objects.get(id=movimento[0].id)
        atualiza_registro.data_recebimento = datetime.datetime.now()
        atualiza_registro.recebido_por = request.user
        atualiza_registro.setor_destino = request.user.pessoafisica.setor
        atualiza_registro.save()
        if atualiza_registro.setor_destino == solicitacao.setor_origem:
            solicitacao.situacao = SolicitacaoLicitacao.CADASTRADO
        else:
            solicitacao.situacao = SolicitacaoLicitacao.RECEBIDO
        solicitacao.save()
    messages.success(request, u'Solicitação recebida com sucesso.')
    return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)

@login_required()
def ver_movimentacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    title=u'Movimentação - %s' % solicitacao
    movimentos = MovimentoSolicitacao.objects.filter(solicitacao=solicitacao).order_by('-data_envio')
    return render(request, 'ver_movimentacao.html', locals(), RequestContext(request))

@login_required()
def cadastrar_minuta(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and solicitacao.eh_apta() and not solicitacao.arquivo_minuta:
        title=u'Cadastrar Minuta - %s' % solicitacao
        form = CadastroMinutaForm(request.POST or None, request.FILES or None, instance=solicitacao)
        if form.is_valid():
            form.save()
            messages.success(request, u'Minuta cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)
        return render(request, 'cadastrar_minuta.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def avalia_minuta(request, solicitacao_id, tipo):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_avaliar_minuta') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and not solicitacao.data_avaliacao_minuta:
        if solicitacao.pode_gerar_ordem() or solicitacao.eh_pedido():

            import tempfile
            import zipfile
            solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)

            municipio = None
            config = get_config_geral()
            if config:
                municipio = get_config_geral().municipio

            num_processo = u''
            if solicitacao.processo:
                num_processo = solicitacao.processo.numero

            lista = list()
            dicionario = {}
            entrou = False
            fornecedor_pedido = None
            if solicitacao.eh_pedido():
                if PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).exists():
                    pedidos = PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao)
                    fornecedor_pedido = pedidos[0].item.fornecedor
                elif PedidoContrato.objects.filter(solicitacao=solicitacao).exists():
                    pedidos = PedidoContrato.objects.filter(solicitacao=solicitacao)
                    fornecedor_pedido = pedidos[0].item.fornecedor
                elif PedidoCredenciamento.objects.filter(solicitacao=solicitacao).exists():
                    pedidos = PedidoCredenciamento.objects.filter(solicitacao=solicitacao)
                    fornecedor_pedido = pedidos[0].fornecedor

                for pedido in pedidos:
                    total = pedidos.aggregate(soma=Sum('valor'))['soma']
                    if total:
                        lista.append([pedido.id, total])
                        dicionario[pedido.id] = total
                        entrou = True

            else:
                for pesquisa in PesquisaMercadologica.objects.filter(solicitacao=solicitacao):
                    total = ItemPesquisaMercadologica.objects.filter(pesquisa=pesquisa, ativo=True).aggregate(soma=Sum('valor_maximo'))['soma']
                    if total:
                        lista.append([pesquisa.id, total])
                        dicionario[pesquisa.id] = total
                        entrou = True
            if entrou:
                resultado = sorted(dicionario.items(), key=lambda x: x[1])
                if solicitacao.eh_pedido():
                    fornecedor = fornecedor_pedido
                    itens = pedidos
                else:
                    fornecedor = PesquisaMercadologica.objects.get(id=resultado[0][0])
                    itens = ItemPesquisaMercadologica.objects.filter(pesquisa=resultado[0][0]).order_by('item')
                total = 0

                for item in itens:
                    total += item.get_total()

                data = datetime.datetime.now().date().strftime('%d/%m/%Y')

                dicionario = {
                    '#PROCESSO#' : num_processo,
                    '#CREDOR#' : fornecedor.razao_social,
                    '#CNPJ#' : fornecedor.cnpj,
                    '#ENDERECO#' : fornecedor.endereco,
                    '#CIDADE#' : municipio or u'-',
                    '#DATA#': data,
                    '#VALOR#': str(total.quantize(Decimal(10) ** -2)).replace('.',','),
                    '#DESCRICAO#': format_numero_extenso(total),
                    '#OBJETO#': solicitacao.objeto,
                    '#PREFEITURA#': config.ordenador_despesa.nome,


                }
                if solicitacao.eh_inexigibilidade() or solicitacao.eh_pedido():
                    template_docx = zipfile.ZipFile(os.path.join(settings.MEDIA_ROOT, 'upload/modelos/termo_inexigibilidade.docx'))
                elif solicitacao.tipo_aquisicao == solicitacao.DISPENSA_LICITACAO_ATE_8MIL:
                    template_docx = zipfile.ZipFile(os.path.join(settings.MEDIA_ROOT, 'upload/modelos/termo_dispensa_ate_8mil.docx'))
                elif solicitacao.tipo_aquisicao == solicitacao.DISPENSA_LICITACAO_ATE_15MIL:
                    template_docx = zipfile.ZipFile(os.path.join(settings.MEDIA_ROOT, 'upload/modelos/termo_dispensa_ate_15mil.docx'))
                elif solicitacao.tipo_aquisicao == solicitacao.TIPO_AQUISICAO_DISPENSA:
                    template_docx = zipfile.ZipFile(os.path.join(settings.MEDIA_ROOT, 'upload/modelos/termo_dispensa_outros.docx'))


                new_docx = zipfile.ZipFile('%s.docx' % tempfile.mktemp(), "a")

                tmp_xml_file = open(template_docx.extract("word/document.xml", tempfile.mkdtemp()))
                tempXmlStr = tmp_xml_file.read()
                tmp_xml_file.close()
                os.unlink(tmp_xml_file.name)

                for key in dicionario.keys():
                    value = unicode(dicionario.get(key)).encode("utf8")
                    tempXmlStr = tempXmlStr.replace(key, value)

                tmp_xml_file =  open(tempfile.mktemp(), "w+")
                tmp_xml_file.write(tempXmlStr)
                tmp_xml_file.close()

                for arquivo in template_docx.filelist:
                    if not arquivo.filename == "word/document.xml":
                        new_docx.writestr(arquivo.filename, template_docx.read(arquivo))

                new_docx.write(tmp_xml_file.name, "word/document.xml")

                template_docx.close()
                new_docx.close()
                os.unlink(tmp_xml_file.name)


                # Caso não seja informado, deverá retornar o caminho para o arquivo DOCX processado.
                caminho_arquivo =  new_docx.filename
                nome_arquivo = caminho_arquivo.split('/')[-1]
                extensao = nome_arquivo.split('.')[-1]
                arquivo = open(caminho_arquivo, "rb")


                content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
                response = HttpResponse(arquivo.read(), content_type=content_type)
                response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
                arquivo.close()
                os.unlink(caminho_arquivo)
                return response
            else:
                messages.error(request, u'Nenhuma pesquisa cadastrada.')
                return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)








        title=u'Avaliar Minuta/Emitir Termo -  %s' % solicitacao
        form = ObsForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            if tipo == u'1':
                solicitacao.minuta_aprovada = True
            elif tipo == u'2':
                solicitacao.minuta_aprovada = True
            solicitacao.data_avaliacao_minuta = datetime.datetime.now()
            solicitacao.minuta_avaliada_por = request.user
            solicitacao.obs_avaliacao_minuta = form.cleaned_data.get('obs')
            solicitacao.arquivo_parecer_minuta = form.cleaned_data.get('arquivo')
            solicitacao.save()
            messages.success(request, u'Minuta avaliada com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)
        return render(request, 'avalia_minuta.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def retomar_lances(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        item.ativo=True
        item.save()
        messages.success(request, u'Lances retomados com sucesso.')
        return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
    else:
        raise PermissionDenied


@login_required()
def informar_quantidades(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if True:
        title=u'Informar quantidade dos itens - %s' % solicitacao
        itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao).order_by('item')

        if request.POST:
            ItemQuantidadeSecretaria.objects.filter(secretaria = request.user.pessoafisica.setor.secretaria, solicitacao = solicitacao).delete()
            for idx, item in enumerate(request.POST.getlist('quantidade'), 1):
                if item and request.POST.getlist('quantidade')[idx-1] > 0:
                    item_do_pregao = ItemSolicitacaoLicitacao.objects.get(solicitacao=solicitacao, id=request.POST.getlist('id_item')[idx-1])
                    novo_preco = ItemQuantidadeSecretaria()
                    novo_preco.solicitacao = solicitacao
                    novo_preco.item = item_do_pregao
                    novo_preco.secretaria = request.user.pessoafisica.setor.secretaria
                    novo_preco.quantidade = Decimal(request.POST.getlist('quantidade')[idx-1].replace('.','').replace(',','.'))
                    novo_preco.save()

            messages.success(request, u'Quantidades cadastradas com sucesso. Faça o upload do memorando de solicitação.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)

        return render(request, 'informar_quantidades.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def ver_pedidos_secretaria(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    solicitacao = item.solicitacao
    title = u'Pedidos das Secretarias - %s' % item
    pedidos = ItemQuantidadeSecretaria.objects.filter(item=item)

    tem_pendente = pedidos.filter(avaliado_em__isnull=True).exists()
    total = pedidos.aggregate(total=Sum('quantidade'))
    pode_avaliar = request.user.groups.filter(name=u'Secretaria').exists() and solicitacao.pode_enviar_para_compra()  and solicitacao.setor_origem == request.user.pessoafisica.setor
    return render(request, 'ver_pedidos_secretaria.html', locals(), RequestContext(request))


@login_required()
def importar_itens(request, solicitacao_id):
    unidades = TipoUnidade.objects.all()
    texto=u''
    for item in unidades:
        texto = texto + item.nome +', '
    texto = texto[:-2]
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao).exists():
        messages.error(request, u'Esta funcionalidade só pode ser usada quando a solicitação não tem nenhum item cadastrado.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)

    if request.user.has_perm('base.pode_cadastrar_solicitacao') and  solicitacao.situacao == solicitacao.CADASTRADO and solicitacao.setor_origem == request.user.pessoafisica.setor and not solicitacao.tem_pregao_cadastrado() and not solicitacao.prazo_aberto:

        title=u'Importar Itens da %s' % solicitacao
        form = ImportarItensForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            arquivo_up = form.cleaned_data.get('arquivo')
            if arquivo_up:
                sheet = None
                try:
                    with transaction.atomic():
                        workbook = xlrd.open_workbook(file_contents=arquivo_up.read())
                        sheet = workbook.sheet_by_index(0)
                        ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao).delete()
                except XLRDError:
                    raise Exception(u'Não foi possível processar a planilha. Verfique se o formato do arquivo é .xls ou .xlsx.')

                for row in range(1, sheet.nrows):


                    especificacao = unicode(sheet.cell_value(row, 1)).strip()
                    unidade = unicode(sheet.cell_value(row, 2)).strip()
                    qtd = unicode(sheet.cell_value(row, 3)).strip()
                    if row == 1:
                        if especificacao != u'ESPECIFICAÇÃO DO PRODUTO' or unidade != u'UNIDADE' or qtd != u'QUANTIDADE':
                            raise Exception(u'Não foi possível processar a planilha. As colunas devem ter Especificação, Unidade e Quantidade.')
                    else:
                        if especificacao and unidade and qtd:
                            try:
                                with transaction.atomic():
                                    Decimal(qtd)
                            except:
                                messages.error(request, u'a quantidade %s %s é inválida.' % (qtd))
                                return HttpResponseRedirect(u'/base/importar_itens/%s/' % solicitacao.id)
                            if TipoUnidade.objects.filter(nome=unidade).exists():
                                un = TipoUnidade.objects.filter(nome=unidade)[0]
                            else:
                                un = TipoUnidade.objects.get_or_create(nome=unidade)[0]
                            novo_item = ItemSolicitacaoLicitacao()
                            novo_item.solicitacao = solicitacao
                            novo_item.item = solicitacao.get_proximo_item()

                            if MaterialConsumo.objects.filter(nome=especificacao).exists():
                                material = MaterialConsumo.objects.filter(nome=especificacao)[0]
                            else:
                                material = MaterialConsumo()
                                if MaterialConsumo.objects.exists():
                                    id = MaterialConsumo.objects.latest('id')
                                    material.id = id.pk+1
                                    material.codigo = id.pk+1
                                else:
                                    material.id = 1
                                    material.codigo = 1
                                material.nome = especificacao
                                material.save()
                            novo_item.material = material
                            novo_item.unidade = un
                            novo_item.quantidade = sheet.cell_value(row, 3)
                            novo_item.save()

                            novo_item_qtd = ItemQuantidadeSecretaria()
                            novo_item_qtd.solicitacao = solicitacao
                            novo_item_qtd.item = novo_item
                            novo_item_qtd.secretaria = request.user.pessoafisica.setor.secretaria
                            novo_item_qtd.quantidade = novo_item.quantidade
                            novo_item_qtd.aprovado = True
                            novo_item_qtd.avaliado_por = request.user
                            novo_item_qtd.avaliado_em = datetime.datetime.now()
                            novo_item_qtd.save()


            messages.success(request, u'Itens cadastrados com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)

        return render(request, 'importar_itens.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def upload_itens_pesquisa_mercadologica(request, pesquisa_id):

    pesquisa = get_object_or_404(PesquisaMercadologica, pk=pesquisa_id)
    title=u'Importar Pesquisa - %s' % pesquisa.solicitacao
    form = ImportarItensForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        arquivo_up = form.cleaned_data.get('arquivo')
        if arquivo_up:
            sheet = None
            try:
                with transaction.atomic():
                    workbook = xlrd.open_workbook(file_contents=arquivo_up.read())
                    sheet = workbook.sheet_by_index(0)
                    ItemPesquisaMercadologica.objects.filter(pesquisa=pesquisa).delete()
            except XLRDError:
                raise Exception(u'Não foi possível processar a planilha. Verfique se o formato do arquivo é .xls ou .xlsx.')

            for row in range(0, sheet.nrows):
                item = unicode(sheet.cell_value(row, 0)).strip()
                especificacao = unicode(sheet.cell_value(row, 1)).strip()
                unidade = unicode(sheet.cell_value(row, 2)).strip()
                marca = unicode(sheet.cell_value(row, 3)).strip()
                valor = unicode(sheet.cell_value(row, 4)).strip()
                if row == 0:
                    if item!= u'ITEM' or especificacao != u'MATERIAL' or unidade != u'UNIDADE' or marca != u'MARCA' or valor!=u'VALOR UNITÁRIO':
                        raise Exception(u'Não foi possível processar a planilha. As colunas devem ter Item, Material, Unidade e Marca e Valor Unitario.')
                else:
                    if item and especificacao and unidade and valor:
                        item_do_pregao = ItemSolicitacaoLicitacao.objects.get(solicitacao=pesquisa.solicitacao, item=int(sheet.cell_value(row, 0)))
                        novo_preco = ItemPesquisaMercadologica()
                        novo_preco.pesquisa = pesquisa
                        novo_preco.item = item_do_pregao
                        try:
                            with transaction.atomic():
                                Decimal(valor)
                        except:
                            messages.error(request, u'o valor %s do %s é inválido.' % (valor, item_do_pregao))
                            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % pesquisa.solicitacao.id)
                        novo_preco.valor_maximo = valor
                        if marca:
                            novo_preco.marca = marca
                        novo_preco.save()


        messages.success(request, u'Itens cadastrados com sucesso.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % pesquisa.solicitacao.id)

    return render(request, 'upload_itens_pesquisa_mercadologica.html', locals(), RequestContext(request))

@login_required()
def relatorio_resultado_final(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    eh_lote = pregao.criterio.id == CriterioPregao.LOTE
    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    eh_maior_desconto = pregao.eh_maior_desconto()

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('item')
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('item')
    total = 0
    for item in itens_pregao:
        if item.get_total_lance_ganhador():
            total = total + item.get_total_lance_ganhador()
    observacao = False
    if not pregao.eh_pregao() and ResultadoItemPregao.objects.filter(item__in=itens_pregao.values_list('id', flat=True), observacoes__isnull=False, ordem=1).exists():
        observacao = ResultadoItemPregao.objects.filter(item__in=itens_pregao.values_list('id', flat=True), observacoes__isnull=False, ordem=1)[0].observacoes
    data = {'eh_lote':eh_lote, 'observacao': observacao,  'eh_maior_desconto': eh_maior_desconto, 'configuracao':configuracao, 'logo':logo, 'itens_pregao': itens_pregao, 'data_emissao':data_emissao, 'pregao':pregao, 'total': total}

    template = get_template('relatorio_resultado_final.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def relatorio_economia(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    eh_lote = pregao.criterio.id == CriterioPregao.LOTE

    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    eh_maior_desconto = pregao.eh_maior_desconto()
    tabela = {}
    total = {}

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(lance = list(), total = 0, total_previsto=0, total_final=0, total_desconto_porcento=0)


    total_previsto_geral = 0
    total_final_geral = 0
    total_desconto_geral = 0
    total_economizado_geral = 0

    for item in itens_pregao.order_by('item'):
        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['lance'].append(item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor



            valor = tabela[chave]['total_previsto']
            valor = valor + (item.valor_medio*item.quantidade)
            tabela[chave]['total_previsto'] = valor


            valor = tabela[chave]['total_final']
            valor = valor + item.get_economizado()
            tabela[chave]['total_final'] = valor



    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        reducao = tabela[chave]['total'] / tabela[chave]['total_previsto']
        ajuste= 1-reducao
        tabela[chave]['total_desconto_porcento'] = u'%s%%' % (ajuste.quantize(TWOPLACES) * 100)




    resultado = collections.OrderedDict(sorted(tabela.items()))

    for result in resultado.items():
        if not result[1]['lance']:
            chave = result[0]
            del resultado[chave]
        else:
            total_previsto_geral = total_previsto_geral + result[1]['total_previsto']
            total_final_geral = total_final_geral + result[1]['total']
            total_economizado_geral = total_economizado_geral + result[1]['total_final']


    reducao = total_final_geral / total_previsto_geral
    ajuste= 1-reducao
    total_desconto_geral = u'%s%%' % (ajuste.quantize(TWOPLACES) * 100)

    mostrou=False
    data = {'mostrou': mostrou, 'total_previsto_geral':total_previsto_geral, 'total_final_geral': total_final_geral, 'total_desconto_geral':total_desconto_geral, 'total_economizado_geral':total_economizado_geral, 'eh_lote':eh_lote, 'eh_maior_desconto':eh_maior_desconto, 'configuracao':configuracao, 'logo':logo, 'itens_pregao': itens_pregao, 'data_emissao':data_emissao, 'pregao':pregao, 'resultado':resultado}

    template = get_template('relatorio_economia.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def relatorio_resultado_final_por_vencedor(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    eh_lote = pregao.criterio.id == CriterioPregao.LOTE

    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    eh_maior_desconto = pregao.eh_maior_desconto()
    tabela = {}
    total = {}

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(lance = list(), total = 0)

    for item in itens_pregao.order_by('item'):
        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['lance'].append(item)
            valor = tabela[chave]['total']
            if eh_lote:
                valor = valor + item.get_valor_total_item_lote()
            else:
                valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor



    resultado = collections.OrderedDict(sorted(tabela.items()))

    data = {'eh_lote':eh_lote, 'eh_maior_desconto':eh_maior_desconto, 'configuracao':configuracao, 'logo':logo, 'itens_pregao': itens_pregao, 'data_emissao':data_emissao, 'pregao':pregao, 'resultado':resultado}

    template = get_template('relatorio_resultado_final_por_vencedor.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def relatorio_lista_participantes(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    participantes = ParticipantePregao.objects.filter(pregao=pregao)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    data = {'participantes': participantes, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'pregao':pregao}

    template = get_template('relatorio_lista_participantes.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def relatorio_lista_visitantes(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    participantes = VisitantePregao.objects.filter(pregao=pregao)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    data = {'participantes': participantes, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'pregao':pregao}

    template = get_template('relatorio_lista_visitantes.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def relatorio_classificacao_por_item(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)

    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    eh_lote = pregao.criterio.id == CriterioPregao.LOTE
    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    eh_maior_desconto = pregao.eh_maior_desconto()
    tabela = {}
    itens = {}
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, item__situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('item__item')
    chaves =  resultado.values('item__item').order_by('item__item')
    for num in sorted(chaves):
        chave = '%s' % num['item__item']
        tabela[chave] = []
        itens[chave] =  []

    for item in resultado.order_by('item','ordem'):
        chave = '%s' % str(item.item.item)
        tabela[chave].append(item)
        itens[chave] = item.item.get_itens_do_lote()

    from blist import sorteddict

    def my_key(dict_key):
           try:
               with transaction.atomic():
                  return int(dict_key)
           except ValueError:
                  return dict_key


    resultado =  sorteddict(my_key, **tabela)


    data = {'itens':itens, 'eh_maior_desconto': eh_maior_desconto, 'configuracao':configuracao, 'logo':logo, 'eh_lote':eh_lote, 'data_emissao':data_emissao, 'pregao':pregao, 'resultado':resultado}

    template = get_template('relatorio_classificacao_por_item.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def relatorio_ocorrencias(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    registros = HistoricoPregao.objects.filter(pregao=pregao)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    data = {'registros': registros, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'pregao':pregao}

    template = get_template('relatorio_ocorrencias.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def relatorio_propostas_pregao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    eh_lote = pregao.criterio.id == CriterioPregao.LOTE
    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    tabela = {}
    itens = {}
    resultado = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False).order_by('item')

    for num in resultado.order_by('item'):
        chave = '%s' % str(num.item)
        tabela[chave] = []


    for item in resultado.order_by('item'):

        chave = '%s' % str(item.item)
        for proposta in PropostaItemPregao.objects.filter(item=item):
            tabela[chave].append(proposta)


    from blist import sorteddict

    def my_key(dict_key):
           try:
               with transaction.atomic():
                  return int(dict_key)
           except ValueError:
                  return dict_key


    resultado =  sorteddict(my_key, **tabela)


    data = {'itens':itens,  'configuracao':configuracao, 'logo':logo, 'eh_lote':eh_lote, 'data_emissao':data_emissao, 'pregao':pregao, 'resultado':resultado}

    template = get_template('relatorio_propostas_pregao.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def relatorio_lances_item(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    eh_lote = pregao.criterio.id == CriterioPregao.LOTE
    destino_arquivo = u'upload/resultados/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    tabela = {}
    itens = {}


    for item in ItemSolicitacaoLicitacao.objects.filter(solicitacao = pregao.solicitacao, lanceitemrodadapregao__isnull=False).order_by('item'):
        lista_rodadas = dict()
        rodadas = RodadaPregao.objects.filter(item=item)
        for rodada in rodadas:
            lista_rodadas[rodada.rodada] = dict(lances=list())
        chave = u'%s' % (item.item)
        tabela[chave] =  lista_rodadas
        itens[chave] =  item.get_itens_do_lote()

        for lance in LanceItemRodadaPregao.objects.filter(item=item).order_by('ordem_lance'):
            lista_rodadas[lance.rodada.rodada]['lances'].append(lance)

        # num_rodadas =  rodadas.values('rodada').order_by('rodada').distinct('rodada')
        # for num in num_rodadas:
        #     chave = '%s' % num['rodada']
        #     tabela[chave] = []
        # for lance in lances.order_by('id'):
        #     chave = '%s' % str(lance.rodada.rodada)
        #     tabela[chave].append(lance)


    from blist import sorteddict

    def my_key(dict_key):
           try:
               with transaction.atomic():
                  return int(dict_key)
           except ValueError:
                  return dict_key


    resultado =  sorteddict(my_key, **tabela)

    #resultado = collections.OrderedDict(sorted(tabela.items()))

    itens = collections.OrderedDict(sorted(itens.items()))



    data = {'eh_lote':eh_lote, 'itens':itens, 'data_emissao':data_emissao, 'pregao':pregao, 'resultado':resultado, 'configuracao': configuracao, 'logo': logo}


    template = get_template('relatorio_lances_item.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def relatorio_ata_registro_preco(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)


    secretaria = pregao.solicitacao.setor_origem.secretaria
    configuracao = get_config(secretaria)
    config_geral = get_config_geral()
    municipio = config_geral.municipio

    if pregao.comissao:
        configuracao_logo = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao_logo = get_config(pregao.solicitacao.setor_origem.secretaria)

    logo = None
    if configuracao_logo.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao_logo.logo.name)

    if secretaria.eh_ordenadora_despesa:
        nome_ordenador = configuracao.nome
        cnpj_ordenador =  configuracao.cnpj
        endereco_ordenador =  configuracao.endereco
        orgao = configuracao.ordenador_despesa.setor.secretaria.nome
        nome_pessoa_ordenadora = configuracao.ordenador_despesa.nome
        cpf_pessoa_ordenadora = configuracao.cpf_ordenador_despesa

    else:

        nome_ordenador = config_geral.nome
        cnpj_ordenador =  config_geral.cnpj
        endereco_ordenador =  config_geral.endereco
        orgao = config_geral.nome
        nome_pessoa_ordenadora = config_geral.ordenador_despesa.nome
        cpf_pessoa_ordenadora = config_geral.cpf_ordenador_despesa

    eh_lote = pregao.criterio.id == CriterioPregao.LOTE
    tabela = {}
    total = {}

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao)
    chaves =  resultado.values('participante').order_by('participante')
    for num in chaves:
        participante = get_object_or_404(ParticipantePregao, pk=num['participante'])
        chave = u'%s' % participante.id
        tabela[chave] = dict(lance = list(), total = 0)

    for item in itens_pregao.order_by('item'):
        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.id
            tabela[chave]['lance'].append(item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor



    resultado = collections.OrderedDict(sorted(tabela.items()))

    texto = u''
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    # table = document.add_table(rows=1, cols=2)
    # hdr_cells = table.rows[0].cells
    #
    #
    # style2 = document.styles['Normal']
    # font = style2.font
    # font.name = 'Arial'
    # font.size = Pt(6)
    #
    # style = document.styles['Normal']
    # font = style.font
    # font.name = 'Arial'
    # font.size = Pt(11)
    #
    #
    # paragraph = hdr_cells[0].paragraphs[0]
    # run = paragraph.add_run()
    # run.add_picture(logo, width=Inches(1.75))
    #
    # paragraph2 = hdr_cells[1].paragraphs[0]
    # paragraph2.style = document.styles['Normal']
    # hdr_cells[1].text =  u'%s' % (configuracao.nome)
    imprimir_cabecalho(document, configuracao, logo, municipio)



    ata = AtaRegistroPreco.objects.get(pregao=pregao)
    document.add_paragraph()
    p = document.add_paragraph(u'ATA DE REGISTRO DE PREÇOS – %s' % ata.numero)

    titulo_pregao = u'sdasd'
    texto = u'''
    No dia %s, o(a) %s, inscrito(a) no CNPJ/MF sob o nº %s, situado(a) no(a)  %s, representado neste ato pelo(a) Sr(a) %s, inscrito no CPF n° %s, nos termos da Lei nº 10.520/2002 e de modo subsidiário, da Lei nº 8.666/93 e Decreto Municipal nº 046/2010, conforme a classificação da proposta apresentada no %s, homologado em %s, resolve registrar o preço oferecido pela empresa, conforme os seguintes termos:
    ''' % (ata.data_inicio.strftime('%d/%m/%y'), nome_ordenador, cnpj_ordenador, endereco_ordenador, nome_pessoa_ordenadora, cpf_pessoa_ordenadora, pregao, pregao.data_homologacao.strftime('%d/%m/%y'))

    #document.add_paragraph(texto)
    p = document.add_paragraph()
    p.alignment = 3

    p.add_run(texto)


    eh_lote = pregao.criterio.id == CriterioPregao.LOTE
    fornecedores = list()
    if eh_lote:
        for item in resultado.items():
            if len(item[1]['lance']) > 0:
                participante = get_object_or_404(ParticipantePregao, id=item[0])
                fornecedor = participante.fornecedor
                fornecedores.append(fornecedor)
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(u'%s' % fornecedor).bold = True
                itens = len(item[1]['lance'])



                table = document.add_table(rows=6, cols=2)

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = u'Empresa: %s' % fornecedor
                a, b = hdr_cells[:2]
                a.merge(b)

                hdr_cells = table.rows[1].cells
                hdr_cells[0].text = u'CNPJ: %s' % fornecedor.cnpj
                hdr_cells[1].text = u'Telefones: %s' % fornecedor.telefones

                hdr_cells = table.rows[2].cells
                hdr_cells[0].text = u'Endereço: %s' % fornecedor.endereco
                a, b = hdr_cells[:2]
                a.merge(b)

                hdr_cells = table.rows[3].cells
                hdr_cells[0].text = u'Representante Legal: %s' % participante.nome_representante
                a, b = hdr_cells[:2]
                a.merge(b)

                hdr_cells = table.rows[4].cells
                hdr_cells[0].text = u'RG: %s' % participante.rg_representante
                hdr_cells[1].text = u'CPF: %s' % participante.cpf_representante


                hdr_cells = table.rows[5].cells
                hdr_cells[0].text = u'Email: %s' % fornecedor.email
                p = document.add_paragraph()

                table = document.add_table(rows=itens+2, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = u'Lote / Item'
                hdr_cells[1].text = u'Objeto'
                hdr_cells[2].text = u'Marca/Modelo'
                hdr_cells[3].text = u'Unidade/Qtd'
                hdr_cells[4].text = u'Preço Unitário (R$)'
                hdr_cells[5].text = u'Preço Total (R$)'
                total_geral = 0
                contador = 1
                for lance in item[1]['lance']:

                    conteudo = u''
                    for componente in lance.get_itens_do_lote():
                        total = lance.get_vencedor().valor * componente.quantidade
                        total_geral += total
                        marca = PropostaItemPregao.objects.filter(item=componente, participante=lance.get_vencedor().participante)[0].marca

                        row_cells = table.rows[contador].cells
                        row_cells[0].text = u'Lote %s / %s' % (lance.item, componente)
                        row_cells[1].text = u'%s' % componente.material.nome
                        row_cells[2].text = u'%s' % (marca)
                        row_cells[3].text = u'%s / %s' % (componente.unidade, format_quantidade(componente.quantidade))
                        row_cells[4].text = u'%s' % format_money(lance.get_vencedor().valor)
                        row_cells[5].text = u'%s' % format_money(total)
                        contador += 1
                row_cells = table.add_row().cells
                row_cells[0].text = u'Total'
                row_cells[5].text = format_money(total_geral)
                p = document.add_paragraph()


    else:
        for item in resultado.items():
            if len(item[1]['lance']) > 0:
                participante = get_object_or_404(ParticipantePregao, id=item[0])
                fornecedor = participante.fornecedor
                fornecedores.append(fornecedor)
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table = document.add_table(rows=6, cols=2)

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = u'Empresa: %s' % fornecedor
                a, b = hdr_cells[:2]
                a.merge(b)

                hdr_cells = table.rows[1].cells
                hdr_cells[0].text = u'CNPJ: %s' % fornecedor.cnpj
                hdr_cells[1].text = u'Telefones: %s' % fornecedor.telefones

                hdr_cells = table.rows[2].cells
                hdr_cells[0].text = u'Endereço: %s' % fornecedor.endereco
                a, b = hdr_cells[:2]
                a.merge(b)

                hdr_cells = table.rows[3].cells
                hdr_cells[0].text = u'Representante Legal: %s' % participante.nome_representante
                a, b = hdr_cells[:2]
                a.merge(b)

                hdr_cells = table.rows[4].cells
                hdr_cells[0].text = u'RG: %s' % participante.rg_representante
                hdr_cells[1].text = u'CPF: %s' % participante.cpf_representante


                hdr_cells = table.rows[5].cells
                hdr_cells[0].text = u'Email: %s' % fornecedor.email
                p = document.add_paragraph()


                itens = len(item[1]['lance'])

                table = document.add_table(rows=itens+1, cols=6)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = u'Item'
                hdr_cells[1].text = u'Objeto'
                hdr_cells[2].text = u'Marca/Modelo'
                hdr_cells[3].text = u'Unidade/Qtd'
                hdr_cells[4].text = u'Preço Unitário (R$)'
                hdr_cells[5].text = u'Preço Total (R$)'
                total_geral = 0
                contador = 1
                for lance in item[1]['lance']:

                    if lance.get_vencedor().participante.fornecedor == fornecedor:
                        total = lance.quantidade * lance.get_vencedor().valor
                        total_geral += total
                        marca = PropostaItemPregao.objects.filter(item=lance, participante=lance.get_vencedor().participante)[0].marca

                        row_cells = table.rows[contador].cells
                        row_cells[0].text = u'%s' % (lance.item)
                        row_cells[1].text = u'%s' % lance.material.nome
                        row_cells[2].text = u'%s' % (marca)
                        row_cells[3].text = u'%s / %s' % (lance.unidade, format_quantidade(lance.quantidade))
                        row_cells[4].text = u'%s' % format_money(lance.get_vencedor().valor)
                        row_cells[5].text = u'%s' % format_money(total)
                        contador += 1
                row_cells = table.add_row().cells
                row_cells[0].text = u'Total'
                row_cells[1].text = u'%s (%s)' % (format_money(total_geral), format_numero_extenso(total_geral))
                row_cells[1].merge(row_cells[5])
                p = document.add_paragraph()


    texto = u'''1 – DO OBJETO

    1.1 – %s, conforme quantidades estimadas e especificações técnicas do Edital do %s supracitado.

    2 – DA VALIDADE DOS PREÇOS

    2.1 – Este Registro de Preços tem validade de até 12 (DOZE) MESES, contados da data da sua assinatura, incluídas eventuais prorrogações, com eficácia legal após a publicação no DIÁRIO OFICIAL e demais meios, conforme exigido na legislação aplicável.

    2.2 – Durante o prazo de validade desta Ata de Registro de Preço, o(a) %s não será obrigado(a) a firmar as contratações que dela poderão advir, facultando-se a realização de licitação específica para a aquisição pretendida, sendo assegurado ao beneficiário do registro preferência no fornecimento em igualdade de condições.

    3 – DA UTILIZAÇÃO DA ATA DE REGISTRO DE PREÇOS POR ÓRGÃO OU ENTIDADES NÃO PARTICIPANTES

    3.1 - A Ata de Registro de Preços, durante sua vigência, poderá ser utilizada por qualquer órgão ou entidade da Administração Pública Municipal, Estadual ou Federal, não-participante do certame licitatório, também denominado carona, mediante prévia consulta junto a CPL, órgão gerenciador da ARP que indicará possíveis fornecedores e respectivos preços, obedecida a ordem de classificação e observadas as seguintes regras:

    I - prévia consulta ao órgão gerenciador da ARP; e

    II - observância da quantidade licitada do objeto constante da Ata e sua compatibilidade com a expectativa de compra, no exercício, pelo órgão carona, para que não ocorra fracionamento.

    § 1º. Caberá ao fornecedor beneficiário da Ata de Registro de Preços, observadas as condições nela estabelecidas, optar pela aceitação ou não do fornecimento, independentemente dos quantitativos registrados em Ata, desde que este fornecimento não prejudique as obrigações anteriormente assumidas.

    § 2º. As aquisições ou contratações adicionais a que se refere este artigo não poderão exceder, por órgão ou entidade, a 100%% (cem por cento) dos quantitativos registrados na Ata de Registro de Preços.

    § 3º.  o quantitativo decorrente das adesões à ata de registro de preços não poderá exceder, na totalidade, ao quíntuplo do quantitativo de cada item registrado na ata de registro de preços para o órgão gerenciador e órgãos participantes, independente do número de órgãos não participantes que aderirem. 

    § 4º. Órgão ou entidade que não participar de todos os lotes do registro de preços, observadas as disposições deste artigo, poderá ser carona nos demais lotes do mesmo registro de preços.

    § 5º. Poderão igualmente utilizar-se da ARP, como carona, mediante prévia consulta ao órgão gerenciador, desde que observadas as condições estabelecidas neste artigo:

    I - outros entes da Administração Pública; e

    II - entidades privadas.

    § 6º Observado o disposto nos §§ 12 e 13 do art. 9º, as contratações dos caronas poderão ser aditadas em quantidades, na forma permitida no art. 65, da Lei Federal nº 8.666, de 1993, se a respectiva Ata não tiver sido aditada.


    4 – DAS DISPOSIÇÕES FINAIS

    4.1 – Integram esta ARP, o edital do Pregão supracitado e seus anexos, e a(s) proposta(s) da(s) empresa(s), classificada(s) no respectivo certame.

    4.2 – Os casos omissos serão resolvidos de acordo com a pelas normas constantes nas Leis n.º 8.666/93 e 10.520/02, no que couber.

    4.3 – Fica eleito o Foro da Comarca Local, para dirimir as dúvidas ou controvérsias resultantes da interpretação deste Contrato, renunciando a qualquer outro por mais privilegiado que seja.

    ''' % (pregao.objeto, pregao.modalidade, nome_ordenador)

    p = document.add_paragraph()
    p.alignment = 3

    p.add_run(texto)


    texto = u'%s/%s, %s' % (municipio.nome, municipio.estado.sigla, ata.data_inicio.strftime('%d/%m/%y'))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.add_run(texto)


    document.add_paragraph()
    texto = u'%s' % (nome_pessoa_ordenadora)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run(texto)


    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(orgao)

    document.add_paragraph()
    for fornecedor in fornecedores:
        nome_responsavel = ParticipantePregao.objects.filter(fornecedor=fornecedor, pregao=pregao)[0].nome_representante
        texto = u'%s' % (nome_responsavel)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p.add_run(texto)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(fornecedor.razao_social)
        document.add_paragraph()

    texto = u'TESTEMUNHAS:'
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.add_run(texto).underline=True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=2, cols=4)

    hdr_cells = table.rows[0].cells
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.5)
    table.columns[2].width = Inches(2.0)
    table.columns[3].width = Inches(5.5)



    hdr_cells[0].text = u'1) '
    hdr_cells[1].text = u'______________________'
    hdr_cells[2].text = u'2) '
    hdr_cells[3].text = u'______________________'

    hdr_cells = table.rows[1].cells

    hdr_cells[0].text = u'CPF/MF: '
    hdr_cells[1].text = u'______________________'
    hdr_cells[2].text = u'CPF/MF: '
    hdr_cells[3].text = u'______________________'







#
# for lance in lances['lance']:
#                 total = lance.get_vencedor().valor * lance.quantidade
#                 if eh_lote:
#                     conteudo = u''
#                     for componente in lance.get_itens_do_lote():
#                         conteudo += '%s<br>' % componente.material.nome
#                     marca = PropostaItemPregao.objects.filter(item=componente, participante=lance.get_vencedor().participante)[0].marca
#
#                     texto = texto + u'<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (lance.item, conteudo, marca, componente.unidade, lance.quantidade, format_money(total))
#                 else:
#                     texto = texto + u'<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (lance.item, lance.material.nome, lance.get_vencedor().marca, lance.unidade, lance.quantidade, format_money(lance.get_vencedor().valor), format_money(total))
#
#             texto = texto + u'</table>'

    #
    # p = document.add_paragraph(u'Na sequência, solicitou dos licitantes presentes a declaração de cumprimento dos requisitos de habilitação e dos documentos para credenciamento dos licitantes presentes:')
    #
    # table = document.add_table(rows=1, cols=2)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Empresa'
    # hdr_cells[1].text = 'Representante'
    #
    #
    #
    # for item in ParticipantePregao.objects.filter(pregao=pregao):
    #     me = u'Não Compareceu'
    #     if item.nome_representante:
    #         me = item.nome_representante
    #
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = u'%s - %s' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
    #     row_cells[1].text = u'%s' % me
    #
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    # p = document.add_paragraph(u'Finalizado o credenciamento foram recebidos os envelopes contendo as propostas de preços e a documentação de habilitação (envelopes nº 01 e 02) das mãos dos representantes credenciados.')
    #
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'DO REGISTRO DO PREGÃO').bold = True
    #
    #
    # p = document.add_paragraph()
    # p.alignment = 3
    #
    # p.add_run(u'Ato contínuo, foram abertos os Envelopes contendo as Propostas e, com a colaboração dos membros da Equipe de Apoio, o Pregoeiro examinou a compatibilidade do objeto, prazos e condições de fornecimento ou de execução, com aqueles definidos no Edital, tendo selecionados todos os licitantes para participarem da Fase de Lances em razão dos preços propostos estarem em conformidade  com as exigências do edital.')
    #
    #
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'DOS LANCES').bold = True
    #
    # p = document.add_paragraph()
    # p.alignment = 3
    #
    #
    # p.add_run(u'O Sr. Pregoeiro, com auxílio da Equipe de Pregão, deu início aos lances verbais, solicitando ao (os) representante (es) da (as) licitante (es) que ofertasse (em) seus lance (es) para o (os) %s em sequência, conforme mapa de lance (es) e classificação anexo.' % tipo)
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'DA HABILITAÇÃO').bold = True
    #
    # p = document.add_paragraph()
    # p.alignment = 3
    #
    # p.add_run(u'Em seguida, foi analisada a aceitabilidade da proposta detentora do menor preço, conforme previsto no edital. Posteriormente, foi analisada a documentação da referida empresa.')
    #
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'DO RESULTADO').bold = True
    #
    # p = document.add_paragraph()
    # p.alignment = 3
    #
    # p.add_run(u'Diante da aceitabilidade da proposta e regularidade frente às exigências de habilitação contidas no instrumento convocatório, foi declarada pelo Pregoeiro e equipe, a vencedora do certame, a empresa: ')
    #
    #
    # p.add_run(resultado_pregao)
    #
    # p = document.add_paragraph()
    # p.alignment = 3
    # p.add_run(u'O valor global do certame, considerando o somatório dos itens licitados, será de R$ %s, respeitado os valores máximos indicados, tendo em vista que o tipo da licitação é o de %s.' % (format_money(total_geral), tipo))
    #
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'DAS OCORRÊNCIAS DA SESSÃO PÚBLICA').bold = True
    #
    #
    # for item in ocorrencias:
    #     p = document.add_paragraph()
    #     p.alignment = 3
    #     p.add_run(item)
    #
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'DO ENCERRAMENTO').bold = True
    #
    # p = document.add_paragraph()
    # p.alignment = 3
    # p.add_run(u'O Pregoeiro, após encerramento desta fase, concedeu aos proponentes vistas ao processo e a todos os documentos. Franqueada a palavra, para observações, questionamentos e/ou interposição de recursos, caso alguém assim desejasse, como nenhum dos proponentes manifestou intenção de recorrer, pelo que renunciam, desde logo, em caráter irrevogável e irretratável, ao direito de interposição de recurso. Nada mais havendo a tratar, o Pregoeiro declarou encerrados os trabalhos, lavrando-se a presente Ata que vai assinada pelos presentes.')
    #
    #
    #
    # for item in membros:
    #
    #     texto = item.split(',')
    #     p = document.add_paragraph()
    #     p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p.add_run(texto[0])
    #     p = document.add_paragraph()
    #     p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p.add_run(u'Matrícula %s' % texto[1])
    #     p = document.add_paragraph()
    #     p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    #     p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     p.add_run(texto[2])
    #
    #
    # table = document.add_table(rows=1, cols=4)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = u'CNPJ/CPF'
    # hdr_cells[1].text = u'Razão Social'
    # hdr_cells[2].text = u'ME/EPP'
    # hdr_cells[3].text = u'Representante'
    #
    # for item in ParticipantePregao.objects.filter(pregao=pregao):
    #     me = u'Não '
    #     if item.me_epp:
    #         me = u'SIm'
    #
    #     repre = u'Não Compareceu'
    #     if item.nome_representante:
    #         repre = item.nome_representante
    #
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = u'%s' % item.fornecedor.cnpj
    #     row_cells[1].text = u'%s' % item.fornecedor.razao_social
    #     row_cells[2].text = u'%s' % me
    #     row_cells[3].text = u'%s' % repre


    document.add_page_break()
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT, 'upload/pregao/atas/ata_sessao_%s.docx' % pregao.id)
    document.save(caminho_arquivo)



    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")

    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response

    # for key, lances in resultado.items():
    #
    #     if lances['lance']:
    #         fornecedor = get_object_or_404(ParticipantePregao, pk=key)
    #
    #         dados = u'''<br><br>
    #           <table>
    #             <tr>
    #                 <td colspan=2>%s</td>
    #             </tr>
    #             <tr>
    #                 <td colspan=2>Endereço: %s</td>
    #             </tr>
    #             <tr>
    #                 <td colspan=2>CPF/CNPJ: %s</td>
    #             </tr>
    #
    #             <tr>
    #                 <td>Representante Legal: %s</td>
    #                 <td>CPF: %s</td>
    #             </tr>
    #
    #         </table>
    #         <br><br>
    #         ''' % (fornecedor.fornecedor.razao_social, fornecedor.fornecedor.endereco, fornecedor.fornecedor.cnpj, fornecedor.nome_representante, fornecedor.cpf_representante)
    #
    #         texto =  texto + unicode(dados)
    #
    #         if eh_lote:
    #             texto = texto + u'''<table><tr><td>Lote</td><td>Itens do Lote</td><td>Marca</td><td>Unidade</td><td>Quantidade</td><td>Valor Total</td></tr>'''
    #         else:
    #             texto = texto + u'''<table><tr><td>Item</td><td>Descrição</td><td>Marca</td><td>Unidade</td><td>Quantidade</td><td>Valor Unit.</td><td>Valor Total</td></tr>'''
    #
    #         for lance in lances['lance']:
    #             total = lance.get_vencedor().valor * lance.quantidade
    #             if eh_lote:
    #                 conteudo = u''
    #                 for componente in lance.get_itens_do_lote():
    #                     conteudo += '%s<br>' % componente.material.nome
    #                 marca = PropostaItemPregao.objects.filter(item=componente, participante=lance.get_vencedor().participante)[0].marca
    #
    #                 texto = texto + u'<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (lance.item, conteudo, marca, componente.unidade, lance.quantidade, format_money(total))
    #             else:
    #                 texto = texto + u'<tr><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td><td>%s</td></tr>' % (lance.item, lance.material.nome, lance.get_vencedor().marca, lance.unidade, lance.quantidade, format_money(lance.get_vencedor().valor), format_money(total))
    #
    #         texto = texto + u'</table>'
    #
    #
    # response = HttpResponse(texto, content_type='application/vnd.ms-word')
    # response['Content-Disposition'] = 'attachment; filename=Ata_registro_preco_pregao_%s.doc' % pregao.id
    # return response



@login_required()
def gerenciar_grupos(request):
    if not request.user.is_superuser:
        messages.success(request, u'Acesso proibido.')
        return HttpResponseRedirect(u'/')

    title=u'Gerenciar Grupos do Usuário'
    form = BuscaPessoaForm(request.POST or None)
    if form.is_valid():
        pessoa = form.cleaned_data.get('pessoa')
        ids_grupos_pessoas = pessoa.user.groups.all().values_list('id', flat=True)
        grupos = Group.objects.all()
        grupos_usuario = list()
        for grupo in grupos:
            if grupo.id in ids_grupos_pessoas:
                grupos_usuario.append((grupo, 1))
            else:
                grupos_usuario.append((grupo, 2))

    return render(request, 'gerenciar_grupos.html', locals(), RequestContext(request))


@login_required()
def gerenciar_grupo_usuario(request, usuario_id, grupo_id, acao):
    if not request.user.is_superuser:
        messages.success(request, u'Acesso proibido.')
        return HttpResponseRedirect(u'/')

    usuario = get_object_or_404(User, pk=usuario_id)
    grupo = get_object_or_404(Group, pk=grupo_id)

    if acao == u'1':
        usuario.groups.add(grupo)
        messages.success(request, u'Usuário adicionado com sucesso.')

    elif acao == u'2':
        usuario.groups.remove(grupo)
        messages.success(request, u'Usuário removido com sucesso.')
    return HttpResponseRedirect(u'/base/gerenciar_grupos/')

@login_required()
def pedido_outro_interessado(request, pedido_id, opcao):
    title=u'Rejeitar Pedido'
    pedido = get_object_or_404(ItemQuantidadeSecretaria, pk=pedido_id)
    if request.user.groups.filter(name=u'Secretaria').exists() and pedido.solicitacao.pode_enviar_para_compra()  and pedido.solicitacao.setor_origem == request.user.pessoafisica.setor:
        if not pedido.avaliado_em:
            if opcao == u'1':
                pedido.aprovado = True
                pedido.avaliado_em = datetime.datetime.now()
                pedido.avaliado_por = request.user
                pedido.save()
                item = pedido.item
                item.quantidade += pedido.quantidade
                item.save()
                messages.success(request, u'Pedido aprovado com sucesso.')
                return HttpResponseRedirect(u'/base/avaliar_pedidos/%s/' % pedido.item.solicitacao.id)
            else:
                form = RemoverParticipanteForm(request.POST or None)
                if form.is_valid():
                    pedido.aprovado = False
                    pedido.justificativa_reprovacao = form.cleaned_data.get('motivo')
                    pedido.avaliado_em = datetime.datetime.now()
                    pedido.avaliado_por = request.user
                    pedido.save()
                    messages.success(request, u'Pedido rejeitado com sucesso.')
                    return HttpResponseRedirect(u'/base/avaliar_pedidos/%s/' % pedido.item.solicitacao.id)

        else:
            return HttpResponseRedirect(u'/base/avaliar_pedidos/%s/' % pedido.item.solicitacao.id)

        return render(request, 'pedido_outro_interessado.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def abrir_processo_para_solicitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    title=u'Abrir Processo - Solicitação N°: %s' % solicitacao.num_memorando
    form = AbrirProcessoForm(request.POST or None, solicitacao=solicitacao)
    if form.is_valid():
        o = form.save(False)
        o.setor_origem = solicitacao.setor_origem
        o.pessoa_cadastro = request.user
        o.save()
        solicitacao.processo = o
        solicitacao.save()
        messages.success(request, u'Processo aberto com sucesso.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)
    return render(request, 'abrir_processo_para_solicitacao.html', locals(), RequestContext(request))


@login_required()
def ver_processo(request, processo_id):
    processo = get_object_or_404(Processo, pk=processo_id)
    solicitacao = SolicitacaoLicitacao.objects.filter(processo=processo)
    if solicitacao.exists():
        solicitacao = solicitacao[0]
    title=u'Processo N°: %s' % processo.numero
    return render(request, 'ver_processo.html', locals(), RequestContext(request))


@login_required()
def imprimir_capa_processo(request, processo_id):
    processo = get_object_or_404(Processo, pk=processo_id)
    solicitacao = get_object_or_404(SolicitacaoLicitacao, processo=processo)

    # destino_arquivo = u'upload/pesquisas/rascunhos/%s.pdf' % processo_id
    # if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'rascunhos')):
    #     os.makedirs(os.path.join(settings.MEDIA_ROOT, 'rascunhos'))
    # caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    # data_emissao = datetime.date.today()
    #
    #
    # data = {'processo': processo, 'data_emissao':data_emissao}
    #
    # template = get_template('imprimir_capa_processo.html')
    #
    # html  = template.render(Context(data))
    #
    # pdf_file = open(caminho_arquivo, "w+b")
    # pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
    #         encoding='utf-8')
    # pdf_file.close()
    # file = open(caminho_arquivo, "r")
    # pdf = file.read()
    # file.close()
    # return HttpResponse(pdf, 'application/pdf')



    response = HttpResponse(content_type = 'application/pdf')
    response['Content-Disposition'] = 'attachment; filename=capa_processo.pdf'

    c = canvas.Canvas(response, pagesize = (LARGURA, ALTURA))


    if Configuracao.objects.exists():
        imagem_logo = os.path.join(settings.MEDIA_ROOT, Configuracao.objects.latest('id').logo.name)
        # if imagem_logo:
        #     c.drawInlineImage(imagem_logo, 30*mm, ALTURA - 50*mm)
        c.setFont('Helvetica-Bold', 20)
        c.drawString(50*mm, ALTURA - 42*mm, u'%s' % Configuracao.objects.latest('id').nome)

    c.setLineWidth(0.2*mm)

    # Nº DO PROTOCOLO E CÓDIGO DE BARRAS
    c.rect(LARGURA - 100*mm, ALTURA - 78*mm, 80*mm, 22*mm)
    c.setFont('Helvetica-Bold', 12)
    c.drawString(LARGURA - 96*mm, ALTURA - 62*mm, u'Protocolo nº %s' % processo.numero)
    codbarra = I2of5(processo.numero, width=100*mm,
                     barHeight=10*mm, barWidth=0.5*mm, bearers = 0, checksum = 0)
    codbarra.drawOn(c, LARGURA - 101*mm, ALTURA - 75*mm)

    from django.template.defaultfilters import truncatechars
    # INFORMAÇÕES SOBRE O DOCUMENTO
    c.rect(30*mm, ALTURA - 129*mm, 160*mm, 47*mm)
    c.setFont('Helvetica', 12)
    c.drawString(32*mm, ALTURA - 88*mm, u'Data: %s' % processo.data_cadastro.strftime('%d/%m/%Y'))
    #c.drawString(110*mm, ALTURA - 88*mm, u'Campus: %s' % processo.uo.setor.sigla)
    #c.drawString(32*mm, ALTURA - 95*mm, u'Interessado: %s' % processo.pessoa_interessada.nome[:55] + (processo.pessoa_interessada.nome[55:] and '...'))
    c.drawString(32*mm, ALTURA - 96*mm, u'Setor de Origem: %s' % truncatechars(processo.setor_origem, 65))
    if solicitacao.credenciamento_origem or solicitacao.eh_inexigibilidade():
        c.drawString(32*mm, ALTURA - 104*mm, u'Tipo: Inexigibilidade')
    else:
        c.drawString(32*mm, ALTURA - 104*mm, u'Tipo: %s' % (solicitacao.tipo_aquisicao))
    #c.drawString(32*mm, ALTURA - 109*mm, u'Destino: %s' % (unicode(processo.tramite_set.all()[0].orgao_recebimento)))

    if solicitacao.tipo == solicitacao.COMPRA:
        if PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).exists():
            origem = solicitacao.arp_origem.solicitacao.get_pregao()
            pedido = PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao)[0]
            fornecedor_pedido = pedido.item.fornecedor
        elif PedidoContrato.objects.filter(solicitacao=solicitacao).exists():
            origem = solicitacao.contrato_origem.solicitacao.get_pregao()
            pedido = PedidoContrato.objects.filter(solicitacao=solicitacao)[0]
            fornecedor_pedido = pedido.item.fornecedor
        elif PedidoCredenciamento.objects.filter(solicitacao=solicitacao).exists():
            origem = solicitacao.credenciamento_origem.solicitacao.get_pregao()
            pedido = PedidoCredenciamento.objects.filter(solicitacao=solicitacao)[0]
            fornecedor_pedido = pedido.fornecedor

        c.drawString(32*mm, ALTURA - 112*mm, u'Origem: %s' % (origem))
        c.drawString(32*mm, ALTURA - 120*mm, u'Interessado: %s' % (fornecedor_pedido))
        L = simpleSplit('Objeto: %s' % truncatechars(processo.objeto, 70),'Helvetica',12,155 * mm)
        y = ALTURA - 126*mm
        for t in L:
            c.drawString(32*mm,y,t)
            y -= 5*mm
    else:
        L = simpleSplit('Objeto: %s' % truncatechars(processo.objeto, 200),'Helvetica',12,155 * mm)
        y = ALTURA - 111*mm
        for t in L:
            c.drawString(32*mm,y,t)
            y -= 5*mm

    # TRAMITAÇÃO

    c.setFont('Helvetica-Bold', 16)
    c.drawCentredString(110*mm, 158*mm, u'TRAMITAÇÃO')

    # Linhas verticais
    for h in range(30, 201, 80):
        c.line(h*mm, 20*mm, h*mm, 152*mm)

    # Linhas horizontais
    for v in range(20, 153, 12):
        c.line(30*mm, v*mm, 190*mm, v*mm)

    # Escrevendo "Data e Destino"
    c.setFont('Helvetica', 11)
    for h in range(30, 130, 80): # horizontal
        for v in range(20, 150, 12): # vertical
            c.drawString(h*mm + 2*mm, v*mm + 4*mm, u'Data: ___/___/______    Destino:')

    c.showPage()
    c.save()
    return response


@atomic
@login_required()
def criar_lote(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    title= u'%s - Cadastrar Novo Lote' % pregao
    form = CriarLoteForm(request.POST or None, pregao=pregao)
    if form.is_valid():
        ids = list()
        valor_medio = Decimal(0.00)
        quantidade = 1
        for item in form.cleaned_data.get('solicitacoes'):
            ids.append(item.id)
            valor_medio = valor_medio +  (item.valor_medio * item.quantidade)


        novo_lote = ItemSolicitacaoLicitacao()
        novo_lote.solicitacao = pregao.solicitacao
        novo_lote.item = pregao.solicitacao.get_proximo_item(eh_lote=True)
        novo_lote.quantidade = quantidade
        novo_lote.valor_medio = valor_medio
        novo_lote.total = valor_medio
        novo_lote.eh_lote = True
        novo_lote.save()

        controle = 1
        for item in ids:
            solicitacao = get_object_or_404(ItemSolicitacaoLicitacao, pk=item)
            novo_item_lote = ItemLote()
            novo_item_lote.lote = novo_lote
            novo_item_lote.item = solicitacao
            novo_item_lote.numero_item = controle
            novo_item_lote.save()
            controle +=1

        messages.success(request, u'Lote criado com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/#lotes' % pregao.id)


    return render(request, 'criar_lote.html', locals(), RequestContext(request))


@login_required()
def extrato_inicial(request, pregao_id):
    # pregao = get_object_or_404(Pregao, pk=pregao_id)
    # if pregao.comissao:
    #     configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    # else:
    #     configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    # logo = None
    # if configuracao.logo:
    #     logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    #
    # destino_arquivo = u'upload/extratos/%s.pdf' % pregao.id
    # if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/extratos')):
    #     os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/extratos'))
    # caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    #
    #
    # data = {'pregao': pregao, 'configuracao': configuracao, 'logo': logo}
    #
    # template = get_template('extrato_inicial.html')
    #
    # html  = template.render(Context(data))
    #
    # pdf_file = open(caminho_arquivo, "w+b")
    # pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
    #         encoding='utf-8')
    # pdf_file.close()
    #
    #
    # file = open(caminho_arquivo, "r")
    # pdf = file.read()
    # file.close()
    # return HttpResponse(pdf, 'application/pdf')
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    municipio = None
    if get_config_geral():
        municipio = get_config_geral().municipio
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    # table = document.add_table(rows=2, cols=2)
    # hdr_cells = table.rows[0].cells
    # hdr_cells2 = table.rows[1].cells
    #
    #
    #
    #
    # style2 = document.styles['Normal']
    # font = style2.font
    # font.name = 'Arial'
    # font.size = Pt(6)
    #
    # style = document.styles['Normal']
    # font = style.font
    # font.name = 'Arial'
    # font.size = Pt(11)
    #
    #
    #
    # paragraph = hdr_cells[0].paragraphs[0]
    # run = paragraph.add_run()
    # run.add_picture(logo, width=Inches(1.75))
    #
    # paragraph2 = hdr_cells[1].paragraphs[0]
    # paragraph2.style = document.styles['Normal']
    # hdr_cells[1].text =  u'%s' % (configuracao.nome)
    #
    #
    # paragraph3 = hdr_cells2[1].paragraphs[0]
    # paragraph3.style2 = document.styles['Normal']
    #
    #
    #
    # #hdr_cells2[0].text =  u'Sistema Orçamentário, Financeiro e Contábil'
    #
    # hdr_cells2[1].text =  u'Endereço: %s, %s' % (configuracao.endereco, municipio)
    # a, b = hdr_cells2[:2]
    # a.merge(b)
    imprimir_cabecalho(document, configuracao, logo, municipio)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(u'Extrato Inicial -  %s' % pregao).bold = True




    texto = u'''
    O Pregoeiro Oficial da %s, objetivando o grau de competitividade preconizado pela administração pública, torna público que estará realizando a(s) licitação(ões) abaixo descrita(s), a saber:

     - %s - Processo Administrativo nº %s
    Originado pelo Memorando n° %s - %s, que objetiva a %s, conforme
    quantidades, condições e especificações constantes no Anexo I – Termo de Referência do Edital, cuja sessão inicial
    está marcada para o %s, às %s (Horário local).

    A(s) referida(s) sessão(ões) será(ão) realizada(s) em: %s.
    O(s) Edital(is) e seus anexos, com as condições e especificações, encontra(m)-se à disposição dos interessados no Setor de Licitações, no endereço acima indicado, das 07:00h às 13:00h, de segunda a sexta-feira, em dias de expediente. O(s) Edital(is) poderão ser requeridos por meio do email %s, através de solicitação contendo o timbrado da requerente e assinado por representante habilitado.
    Quaisquer esclarecimentos poderão ser prestados no endereço indicado ou através dos telefones: %s.


    ''' % (configuracao.nome, pregao, pregao.solicitacao.processo, pregao.solicitacao.num_memorando, pregao.solicitacao.cadastrado_por.pessoafisica.setor.secretaria, pregao.solicitacao.objeto, localize(pregao.data_abertura), pregao.hora_abertura, pregao.local, configuracao.email, configuracao.telefones)

    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(texto)


    texto = u'''
    %s, %s
    ''' % (municipio, localize(pregao.data_inicio))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(texto)

    texto = u'''
    %s

    Pregoeiro
    ''' % (pregao.responsavel)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(texto)




    caminho_arquivo = os.path.join(settings.MEDIA_ROOT, 'upload/pregao/atas/extrato_inicial_%s.docx' % pregao.id)
    document.save(caminho_arquivo)



    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")

    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response


@login_required()
def termo_adjudicacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/extratos/%s.pdf' % pregao.id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/extratos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/extratos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    config_geral = get_config_geral()

    tabela = {}

    eh_lote = pregao.criterio.id == CriterioPregao.LOTE

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(ordem=1, item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(itens = list(), total = 0)
    total_geral = 0
    for item in itens_pregao.order_by('item'):

        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['itens'].append(item.item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor


    for item in tabela:
        total_geral = total_geral + tabela[item]['total']

    fracassados = list()
    for item in itens_pregao.filter(situacao=ItemSolicitacaoLicitacao.FRACASSADO):
        fracassados.append(item.item)

    resultado = collections.OrderedDict(sorted(tabela.items()))


    data = {'pregao': pregao, 'eh_lote': eh_lote, 'configuracao': configuracao, 'logo': logo, 'resultado': resultado, 'total_geral': total_geral, 'fracassados': fracassados, 'config_geral': config_geral}

    template = get_template('termo_adjudicacao.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()


    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def editar_meu_perfil(request, pessoa_id):
    pessoa = get_object_or_404(PessoaFisica, pk=pessoa_id)
    if pessoa.id == request.user.pessoafisica.id:
        title=u'Editar Meu Perfil'
        form = PessoaFisicaForm(request.POST or None, request=request, instance=pessoa, edicao=True)
        if form.is_valid():
            form.save()
            messages.success(request, u'Perfil editado com sucesso.')
            return HttpResponseRedirect(u'/')


    return render(request, 'editar_meu_perfil.html', locals(), RequestContext(request))


@login_required()
def editar_pedido(request, pedido_id):
    title=u'Editar Pedido'
    pedido = get_object_or_404(ItemQuantidadeSecretaria, pk=pedido_id)
    solicitacao = pedido.solicitacao
    form = EditarPedidoForm(request.POST or None, instance=pedido)
    if form.is_valid():
        form.save()
        messages.success(request, u'Perfil editado com sucesso.')
        return HttpResponseRedirect(u'/base/ver_pedidos_secretaria/%s/' % pedido.item.id)


    return render(request, 'editar_pedido.html', locals(), RequestContext(request))



@login_required()
def aprovar_todos_pedidos(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)

    total=0
    for pedido in ItemQuantidadeSecretaria.objects.filter(item=item, avaliado_em__isnull=True):
        total += pedido.quantidade
    item.quantidade += total
    item.save()
    ItemQuantidadeSecretaria.objects.filter(item=item, avaliado_em__isnull=True).update(aprovado=True, avaliado_por=request.user, avaliado_em=datetime.datetime.now())
    messages.success(request, u'Pedidos aprovados com sucesso.')
    return HttpResponseRedirect(u'/base/ver_pedidos_secretaria/%s/' % item.id)


def gestao_contratos_tipo(request):
    title = u'Gestão de Contratos'
    return render(request, 'gestao_contratos_tipo.html', locals(), RequestContext(request))

@login_required()
def gestao_pedidos_tipo(request):
    title = u'Gestão de Pedidos'
    return render(request, 'gestao_pedidos_tipo.html', locals(), RequestContext(request))



@login_required()
def gestao_pedidos(request, tipo_id):
    setor = request.user.pessoafisica.setor

    meus_pedidos = ItemQuantidadeSecretaria.objects.filter(secretaria=setor.secretaria).values_list('solicitacao', flat=True)
    contratos = atas = credenciamentos = sem_registro = nome = None
    if tipo_id == u'1':
        nome = u'Contratos'
        contratos = Contrato.objects.filter(Q(liberada_compra=True), Q(solicitacao__in=meus_pedidos) | Q(solicitacao__setor_origem__secretaria=setor.secretaria))
        if not contratos.exists():
            sem_registro = u'Nenhum contrato disponível para pedidos.'
    elif tipo_id == u'2':
        nome = u'Atas de Registro de Preço'
        atas = AtaRegistroPreco.objects.filter(Q(liberada_compra=True), Q(solicitacao__in=meus_pedidos) | Q(solicitacao__setor_origem__secretaria=setor.secretaria)).exclude(adesao=True)
        if not atas.exists():
            sem_registro = u'Nenhuma ata disponível para pedidos.'
    #contratos = SolicitacaoLicitacao.objects.filter(liberada_compra=True, id__in=contratos_finalizados.values_list('solicitacao', flat=True))
    elif tipo_id == u'3':
        nome = u'Credenciamentos'
        credenciamentos = Credenciamento.objects.filter(Q(liberada_compra=True), Q(solicitacao__in=meus_pedidos) | Q(solicitacao__setor_origem__secretaria=setor.secretaria))
        if not credenciamentos.exists():
            sem_registro = u'Nenhum credenciamento disponível para pedidos.'
    pode_editar = request.user.groups.filter(name=u'Gerente')
    title=u'Gestão de Pedidos - %s' % nome
    return render(request, 'gestao_pedidos.html', locals(), RequestContext(request))

@login_required()
def gestao_contratos(request, tipo_id):
    setor = request.user.pessoafisica.setor
    pode_editar = False
    if request.user.groups.filter(name=u'Gerente'):

        solicitacoes = SolicitacaoLicitacao.objects.filter(setor_origem__secretaria=setor.secretaria)

        nome = None
        if tipo_id == u'1':
            registros = Contrato.objects.all().order_by('id')
            nome = u'Contratos'
            tipo = u'contrato'
            url_relatorio = 'ver_relatorios_gerenciais_contratos'
        elif tipo_id == u'2':
            registros = AtaRegistroPreco.objects.all().order_by('id')
            nome = u'Atas de Registro de Preço'
            url_relatorio = 'ver_relatorios_gerenciais_atas'
            tipo = 'ata_registro_preco'
        elif tipo_id == u'3':
            registros = Credenciamento.objects.all().order_by('id')
            nome = u'Credenciamentos'
            url_relatorio = 'ver_relatorios_gerenciais_credenciamentos'
            tipo = u'credenciamento'
        pode_editar = True
        form = GestaoContratoForm(request.GET or None, tipo=tipo_id)
        if form.is_valid():
            if form.cleaned_data.get('info'):
                registros = registros.filter(numero__icontains=form.cleaned_data.get('info'))

            if form.cleaned_data.get('ano'):
                registros = registros.filter(data_inicio__year=form.cleaned_data.get('ano'))


            if form.cleaned_data.get('secretaria'):
                registros = registros.filter(solicitacao__setor_origem__secretaria=form.cleaned_data.get('secretaria'))

            if tipo_id == u'1' and form.cleaned_data.get('fornecedor'):
                itens = ItemContrato.objects.filter(Q(fornecedor=form.cleaned_data.get('fornecedor')) | Q(participante__fornecedor=form.cleaned_data.get('fornecedor')))
                registros = registros.filter(id__in=itens.values_list('contrato', flat=True))

        title=u'Gestão de %s' % nome

    else:
        return HttpResponseRedirect(u'/')

    return render(request, 'gestao_contratos.html', locals(), RequestContext(request))


@login_required()
def avaliar_pedidos(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    title=u'Avaliar Pedidos - %s' % solicitacao
    tabela = {}
    pode_avaliar = request.user.groups.filter(name=u'Secretaria').exists()  and solicitacao.setor_origem == request.user.pessoafisica.setor
    if request.user.groups.filter(name=u'Secretaria').exists() and solicitacao.pode_enviar_para_compra()  and solicitacao.setor_origem == request.user.pessoafisica.setor:
        pedidos = ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao)

        total = solicitacao.interessados.count()
        informados = pedidos.values_list('secretaria', flat=True).distinct('secretaria').count() - 1

        form = FiltrarSecretariaForm(request.POST or None, pedidos=pedidos)
        if request.GET.get('secretaria'):
            pedidos = pedidos.filter(secretaria=request.GET.get('secretaria'))
        chaves =  pedidos.values_list('secretaria', flat=True).distinct('secretaria')
        for num in chaves:
            secretaria = get_object_or_404(Secretaria, pk=num)
            chave = u'%s' % secretaria.id
            pendente = pedidos.filter(secretaria=secretaria, avaliado_em__isnull=True).exists()
            tabela[chave] = dict(pedido = list(), nome=secretaria, pendente=pendente)

        for item in pedidos.order_by('item'):

            chave = u'%s' % item.secretaria.id
            tabela[chave]['pedido'].append(item)

        resultado = collections.OrderedDict(sorted(tabela.items()))
        return render(request, 'avaliar_pedidos.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def aprovar_todos_pedidos_secretaria(request, solicitacao_id, secretaria_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.groups.filter(name=u'Secretaria').exists()  and solicitacao.setor_origem == request.user.pessoafisica.setor:
        if ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao_id, secretaria=secretaria_id, avaliado_em__isnull=True).exists():
            for pedido in ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao_id, secretaria=secretaria_id, avaliado_em__isnull=True):
                pedido.item.quantidade += pedido.quantidade
                pedido.item.save()
            ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao_id, secretaria=secretaria_id, avaliado_em__isnull=True).update(aprovado=True, avaliado_por=request.user, avaliado_em=datetime.datetime.now())
            messages.success(request, u'Pedidos aprovados com sucesso.')
        return HttpResponseRedirect(u'/base/avaliar_pedidos/%s/' % solicitacao_id)


@login_required()
def novo_pedido_compra_contrato(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)
    title=u'Novo Pedido de Compra - %s' % contrato
    form = NovoPedidoCompraForm(request.POST or None)
    if form.is_valid():
        o = form.save(False)
        o.tipo = SolicitacaoLicitacao.COMPRA
        o.tipo_aquisicao = SolicitacaoLicitacao.TIPO_AQUISICAO_COMPRA
        o.setor_origem = request.user.pessoafisica.setor
        o.setor_atual = request.user.pessoafisica.setor
        o.data_cadastro = datetime.datetime.now()
        o.cadastrado_por = request.user
        o.contrato_origem = contrato
        o.save()

        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_contrato/%s/%s/' % (contrato_id, o.id))
    return render(request, 'novo_pedido_compra.html', locals(), RequestContext(request))

@login_required()
def novo_pedido_compra_arp(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    title=u'Novo Pedido de Compra - %s' % ata
    form = NovoPedidoCompraForm(request.POST or None)
    if form.is_valid():
        o = form.save(False)
        o.tipo = SolicitacaoLicitacao.COMPRA
        o.tipo_aquisicao = SolicitacaoLicitacao.TIPO_AQUISICAO_COMPRA
        o.setor_origem = request.user.pessoafisica.setor
        o.setor_atual = request.user.pessoafisica.setor
        o.data_cadastro = datetime.datetime.now()
        o.cadastrado_por = request.user
        o.arp_origem = ata
        o.save()
        if ata.adesao:
            return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_adesao_arp/%s/%s/' % (ata_id, o.id))

        else:
            return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, o.id))
    return render(request, 'novo_pedido_compra.html', locals(), RequestContext(request))

@login_required()
def novo_pedido_compra_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)
    title=u'Novo Pedido de Compra - %s' % credenciamento
    form = NovoPedidoCompraForm(request.POST or None)
    if form.is_valid():
        o = form.save(False)
        o.tipo = SolicitacaoLicitacao.COMPRA
        o.tipo_aquisicao = SolicitacaoLicitacao.TIPO_AQUISICAO_COMPRA
        o.setor_origem = request.user.pessoafisica.setor
        o.setor_atual = request.user.pessoafisica.setor
        o.data_cadastro = datetime.datetime.now()
        o.cadastrado_por = request.user
        o.credenciamento_origem = credenciamento
        o.save()
        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_credenciamento/%s/%s/' % (credenciamento_id, o.id))
    return render(request, 'novo_pedido_compra.html', locals(), RequestContext(request))

@login_required()
def informar_quantidades_do_pedido_credenciamento(request, credenciamento_id, solicitacao_id):
    setor = request.user.pessoafisica.setor
    solicitacao_atual = get_object_or_404(SolicitacaoLicitacaoTmp, pk=solicitacao_id)
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)
    itens_credenciamento = credenciamento.itemcredenciamento_set.all()
    solicitacao = credenciamento.solicitacao
    title=u'Pedido de Compra - %s' % credenciamento
    participantes = ParticipantePregao.objects.filter(pregao=credenciamento.pregao, desclassificado=False, excluido_dos_itens=False, credenciado=True)
    form = FiltraVencedorPedidoForm(request.POST or None, participantes=participantes)
    eh_lote = False
    origem_pregao = credenciamento.solicitacao.get_pregao()
    resultados = itens_credenciamento
    buscou = False

    if form.is_valid():
        buscou = True


        fornecedor = form.cleaned_data.get('vencedor')


        if 'quantidades' in request.POST:
            if not fornecedor:
                messages.error(request, u'Selecione um fornecedor')
                return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_credenciamento/%s/%s/' % (credenciamento_id, solicitacao_atual.id))

            for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                valor_pedido = int(valor)

                if valor_pedido > 0:
                    if valor_pedido > resultados.get(id=request.POST.getlist('id')[idx]).get_quantidade_disponivel():
                        messages.error(request, u'A quantidade disponível do item "%s" é menor do que a quantidade solicitada.' % resultados[idx].item)
                        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_credenciamento/%s/%s/' % (credenciamento_id, solicitacao_atual.id))


            nova_solicitacao = SolicitacaoLicitacao()
            nova_solicitacao.num_memorando = solicitacao_atual.num_memorando
            nova_solicitacao.objeto = solicitacao_atual.objeto
            nova_solicitacao.objetivo = solicitacao_atual.objetivo
            nova_solicitacao.situacao = solicitacao_atual.situacao
            nova_solicitacao.tipo = solicitacao_atual.tipo
            nova_solicitacao.tipo_aquisicao = solicitacao_atual.tipo_aquisicao
            nova_solicitacao.data_cadastro = solicitacao_atual.data_cadastro
            nova_solicitacao.cadastrado_por = solicitacao_atual.cadastrado_por
            nova_solicitacao.setor_origem = solicitacao_atual.setor_origem
            nova_solicitacao.setor_atual = solicitacao_atual.setor_atual
            nova_solicitacao.arp_origem = solicitacao_atual.arp_origem
            nova_solicitacao.contrato_origem = solicitacao_atual.contrato_origem
            nova_solicitacao.credenciamento_origem = solicitacao_atual.credenciamento_origem
            nova_solicitacao.save()
            solicitacao_atual.delete()

            for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                valor_pedido = int(valor)
                if valor_pedido > 0:

                    novo_pedido = PedidoCredenciamento()
                    novo_pedido.credenciamento = credenciamento
                    novo_pedido.solicitacao = nova_solicitacao
                    novo_pedido.item = resultados.get(id=request.POST.getlist('id')[idx])
                    novo_pedido.valor = resultados.get(id=request.POST.getlist('id')[idx]).valor
                    novo_pedido.fornecedor = fornecedor.fornecedor
                    novo_pedido.quantidade = valor_pedido
                    novo_pedido.setor = setor
                    novo_pedido.pedido_por = request.user
                    novo_pedido.pedido_em = datetime.datetime.now()
                    novo_pedido.save()

            messages.success(request, u'Pedido cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % nova_solicitacao.id)
    return render(request, 'informar_quantidades_do_pedido_credenciamento.html', locals(), RequestContext(request))


@login_required()
def informar_quantidades_do_pedido_arp(request, ata_id, solicitacao_id):
    setor = request.user.pessoafisica.setor
    solicitacao_atual = get_object_or_404(SolicitacaoLicitacaoTmp, pk=solicitacao_id)
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    itens_ata = ata.itemataregistropreco_set.all()
    solicitacao = ata.solicitacao
    title=u'Pedido de Compra - %s' % ata
    participantes = ParticipantePregao.objects.filter(id__in=itens_ata.values_list('participante', flat=True))
    form = FiltraVencedorPedidoForm(request.POST or None, participantes=participantes)
    eh_lote = solicitacao.eh_lote()
    origem_pregao = ata.solicitacao.get_pregao()
    if eh_lote:
        resultados = solicitacao.get_lotes()
    else:
        resultados = itens_ata
    buscou = False

    if form.is_valid():
        buscou = True
        if eh_lote:
            ids = list()
            for item in solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=True):
                registro = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO).order_by('ordem')[0]
                if registro.participante == form.cleaned_data.get('vencedor'):
                    ids.append(registro.item.id)
            resultados = resultados.filter(id__in=ids)
        else:
            resultados = itens_ata.filter(participante=form.cleaned_data.get('vencedor'))


        fornecedor = form.cleaned_data.get('vencedor')


        if 'quantidades' in request.POST:

            fornecedor = request.POST.get('fornecedor')
            participante = ParticipantePregao.objects.get(id=fornecedor)

            if eh_lote and '0' in request.POST.getlist('quantidades'):
                messages.error(request, u'Informe a quantidade solicitada para cada item do lote')
                return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))


            if eh_lote:
                ids = list()
                resultados = solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=False)
                for item in solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=True):
                    registro = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO).order_by('ordem')[0]
                    if registro.participante == participante:
                        for id_do_item in registro.item.get_itens_do_lote():
                            ids.append(id_do_item.id)
                resultados = resultados.filter(id__in=ids)


                for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                    try:
                        with transaction.atomic():
                            int(valor)
                    except:
                        messages.error(request, u'o valor %s é inválido.' % (valor))
                        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))
                    valor_pedido = int(valor)
                    if valor_pedido > 0:
                        if valor_pedido > resultados.get(id=request.POST.getlist('id')[idx]).get_item_arp().get_quantidade_disponivel():
                            messages.error(request, u'A quantidade disponível do item "%s" é menor do que a quantidade solicitada.' % resultados[idx])
                            return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))
            else:
                resultados = itens_ata.filter(participante=participante)
                for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                    try:
                        with transaction.atomic():
                            int(valor)
                    except:
                        messages.error(request, u'o valor %s é inválido.' % (valor))
                        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))
                    valor_pedido = int(valor)

                    if valor_pedido > 0:
                        if valor_pedido > resultados.get(id=request.POST.getlist('id')[idx]).get_quantidade_disponivel():
                            messages.error(request, u'A quantidade disponível do item "%s" é menor do que a quantidade solicitada.' % resultados[idx].item)
                            return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))


            nova_solicitacao = SolicitacaoLicitacao()
            nova_solicitacao.num_memorando = solicitacao_atual.num_memorando
            nova_solicitacao.objeto = solicitacao_atual.objeto
            nova_solicitacao.objetivo = solicitacao_atual.objetivo
            nova_solicitacao.situacao = solicitacao_atual.situacao
            nova_solicitacao.tipo = solicitacao_atual.tipo
            nova_solicitacao.tipo_aquisicao = solicitacao_atual.tipo_aquisicao
            nova_solicitacao.data_cadastro = solicitacao_atual.data_cadastro
            nova_solicitacao.cadastrado_por = solicitacao_atual.cadastrado_por
            nova_solicitacao.setor_origem = solicitacao_atual.setor_origem
            nova_solicitacao.setor_atual = solicitacao_atual.setor_atual
            nova_solicitacao.arp_origem = solicitacao_atual.arp_origem
            nova_solicitacao.contrato_origem = solicitacao_atual.contrato_origem
            nova_solicitacao.save()
            solicitacao_atual.delete()

            for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):

                valor_pedido = int(valor)
                if valor_pedido > 0:
                    novo_pedido = PedidoAtaRegistroPreco()
                    novo_pedido.ata = ata
                    novo_pedido.solicitacao = nova_solicitacao
                    if eh_lote:

                        novo_pedido.item = ItemAtaRegistroPreco.objects.get(item=resultados.get(id=request.POST.getlist('id')[idx]))
                        novo_pedido.valor = ItemAtaRegistroPreco.objects.get(item=resultados.get(id=request.POST.getlist('id')[idx])).valor
                    else:
                        novo_pedido.item = resultados.get(id=request.POST.getlist('id')[idx])
                        novo_pedido.valor = resultados.get(id=request.POST.getlist('id')[idx]).valor

                    novo_pedido.quantidade = valor_pedido
                    novo_pedido.setor = setor
                    novo_pedido.pedido_por = request.user
                    novo_pedido.pedido_em = datetime.datetime.now()
                    novo_pedido.save()

            messages.success(request, u'Pedido cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % nova_solicitacao.id)
    return render(request, 'informar_quantidades_do_pedido_arp.html', locals(), RequestContext(request))


@login_required()
def informar_quantidades_do_pedido_adesao_arp(request, ata_id, solicitacao_id):
    setor = request.user.pessoafisica.setor
    solicitacao_atual = get_object_or_404(SolicitacaoLicitacaoTmp, pk=solicitacao_id)
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    itens_ata = ata.itemataregistropreco_set.all()
    solicitacao = ata.solicitacao
    title=u'Pedido de Compra - %s' % ata
    participantes = Fornecedor.objects.filter(id__in=itens_ata.values_list('fornecedor', flat=True))
    form = FiltraFornecedorPedidoForm(request.POST or None, participantes=participantes)
    resultados = itens_ata
    buscou = False

    if form.is_valid():

        buscou = True
        resultados = itens_ata.filter(fornecedor=form.cleaned_data.get('vencedor'))


        fornecedor = form.cleaned_data.get('vencedor')


        if 'quantidades' in request.POST:

            fornecedor = request.POST.get('fornecedor')
            participante = Fornecedor.objects.get(id=fornecedor)


            resultados = itens_ata.filter(fornecedor=participante)
            for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                try:
                    with transaction.atomic():
                        int(valor)
                except:
                    messages.error(request, u'o valor %s é inválido.' % (valor))
                    return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))
                valor_pedido = int(valor)

                if valor_pedido > 0:
                    if valor_pedido > resultados.get(id=request.POST.getlist('id')[idx]).get_quantidade_disponivel():
                        messages.error(request, u'A quantidade disponível do item "%s" é menor do que a quantidade solicitada.' % resultados[idx].item)
                        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_arp/%s/%s/' % (ata_id, solicitacao_atual.id))


            nova_solicitacao = SolicitacaoLicitacao()
            nova_solicitacao.num_memorando = solicitacao_atual.num_memorando
            nova_solicitacao.objeto = solicitacao_atual.objeto
            nova_solicitacao.objetivo = solicitacao_atual.objetivo
            nova_solicitacao.situacao = solicitacao_atual.situacao
            nova_solicitacao.tipo = solicitacao_atual.tipo
            nova_solicitacao.tipo_aquisicao = solicitacao_atual.tipo_aquisicao
            nova_solicitacao.data_cadastro = solicitacao_atual.data_cadastro
            nova_solicitacao.cadastrado_por = solicitacao_atual.cadastrado_por
            nova_solicitacao.setor_origem = solicitacao_atual.setor_origem
            nova_solicitacao.setor_atual = solicitacao_atual.setor_atual
            nova_solicitacao.arp_origem = solicitacao_atual.arp_origem
            nova_solicitacao.contrato_origem = solicitacao_atual.contrato_origem
            nova_solicitacao.save()
            solicitacao_atual.delete()

            for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                valor_pedido = int(valor)
                if valor_pedido > 0:
                    novo_pedido = PedidoAtaRegistroPreco()
                    novo_pedido.ata = ata
                    novo_pedido.solicitacao = nova_solicitacao
                    novo_pedido.item = resultados.get(id=request.POST.getlist('id')[idx])
                    novo_pedido.valor = resultados.get(id=request.POST.getlist('id')[idx]).valor

                    novo_pedido.quantidade = valor_pedido
                    novo_pedido.setor = setor
                    novo_pedido.pedido_por = request.user
                    novo_pedido.pedido_em = datetime.datetime.now()
                    novo_pedido.save()

            messages.success(request, u'Pedido cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % nova_solicitacao.id)
    return render(request, 'informar_quantidades_do_pedido_arp.html', locals(), RequestContext(request))




@login_required()
def informar_quantidades_do_pedido_contrato(request, contrato_id, solicitacao_id):
    setor = request.user.pessoafisica.setor
    solicitacao_atual = get_object_or_404(SolicitacaoLicitacaoTmp, pk=solicitacao_id)
    contrato = get_object_or_404(Contrato, pk=contrato_id)
    itens_contrato = contrato.itemcontrato_set.all()
    solicitacao = contrato.solicitacao
    title=u'Pedido de Compra - %s' % contrato
    origem_pregao = contrato.solicitacao.get_pregao()
    if origem_pregao:
        participantes = ParticipantePregao.objects.filter(id__in=itens_contrato.values_list('participante', flat=True))
        form = FiltraVencedorPedidoForm(request.POST or None, participantes=participantes)
    else:
        participantes = Fornecedor.objects.filter(id__in=itens_contrato.values_list('fornecedor', flat=True))
        form = FiltraFornecedorPedidoForm(request.POST or None, participantes=participantes)

    eh_lote = False
    resultados = itens_contrato
    buscou = False

    if form.is_valid():
        buscou = True
        if eh_lote:
            ids = list()
            for item in solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=True):
                registro = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO).order_by('ordem')[0]
                if registro.participante == form.cleaned_data.get('vencedor'):
                    ids.append(registro.item.id)
            resultados = resultados.filter(id__in=ids)
        else:
            if origem_pregao:
                resultados = itens_contrato.filter(participante=form.cleaned_data.get('vencedor'))
            else:
                resultados = itens_contrato.filter(fornecedor=form.cleaned_data.get('vencedor'))


        fornecedor = form.cleaned_data.get('vencedor')


        if 'quantidades' in request.POST:

            fornecedor = request.POST.get('fornecedor')
            if origem_pregao:
                participante = ParticipantePregao.objects.get(id=fornecedor)
            else:
                participante = Fornecedor.objects.get(id=fornecedor)


            if eh_lote and '0' in request.POST.getlist('quantidades'):
                messages.error(request, u'Informe a quantidade solicitada para cada item do lote')
                return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_contrato/%s/%s/' % (contrato_id, solicitacao_atual.id))


            if eh_lote:
                ids = list()
                resultados = solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=False)
                for item in solicitacao.itemsolicitacaolicitacao_set.filter(eh_lote=True):
                    registro = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO).order_by('ordem')[0]
                    if registro.participante == participante:
                        for id_do_item in registro.item.get_itens_do_lote():
                            ids.append(id_do_item.id)
                resultados = resultados.filter(id__in=ids)


                for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                    try:
                        with transaction.atomic():
                            int(valor)
                    except:
                        messages.error(request, u'o valor %s é inválido.' % (valor))
                        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_contrato/%s/%s/' % (contrato_id, solicitacao_atual.id))
                    valor_pedido = int(valor)
                    if valor_pedido > 0:
                        if valor_pedido > resultados.get(id=request.POST.getlist('id')[idx]).get_item_contrato().get_quantidade_disponivel():
                            messages.error(request, u'A quantidade disponível do item "%s" é menor do que a quantidade solicitada.' % resultados[idx])
                            return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_contrato/%s/%s/' % (contrato_id, solicitacao_atual.id))
            else:
                if origem_pregao:
                    resultados = itens_contrato.filter(participante=participante)
                else:
                    resultados = itens_contrato.filter(fornecedor=participante)
                for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                    try:
                        with transaction.atomic():
                            int(valor)
                    except:
                        messages.error(request, u'o valor %s é inválido.' % (valor))
                        return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_contrato/%s/%s/' % (contrato_id, solicitacao_atual.id))
                    valor_pedido = int(valor)

                    if valor_pedido > 0:
                        if valor_pedido > resultados.get(id=request.POST.getlist('id')[idx]).get_quantidade_disponivel():
                            messages.error(request, u'A quantidade disponível do item "%s" é menor do que a quantidade solicitada.' % resultados[idx].item)
                            return HttpResponseRedirect(u'/base/informar_quantidades_do_pedido_contrato/%s/%s/' % (contrato_id, solicitacao_atual.id))

            nova_solicitacao = SolicitacaoLicitacao()
            nova_solicitacao.num_memorando = solicitacao_atual.num_memorando
            nova_solicitacao.objeto = solicitacao_atual.objeto
            nova_solicitacao.objetivo = solicitacao_atual.objetivo
            nova_solicitacao.situacao = solicitacao_atual.situacao
            nova_solicitacao.tipo = solicitacao_atual.tipo
            nova_solicitacao.tipo_aquisicao = solicitacao_atual.tipo_aquisicao
            nova_solicitacao.data_cadastro = solicitacao_atual.data_cadastro
            nova_solicitacao.cadastrado_por = solicitacao_atual.cadastrado_por
            nova_solicitacao.setor_origem = solicitacao_atual.setor_origem
            nova_solicitacao.setor_atual = solicitacao_atual.setor_atual
            nova_solicitacao.arp_origem = solicitacao_atual.arp_origem
            nova_solicitacao.contrato_origem = solicitacao_atual.contrato_origem
            nova_solicitacao.save()
            solicitacao_atual.delete()

            for idx, valor in enumerate(request.POST.getlist('quantidades'), 0):
                valor_pedido = int(valor)
                if valor_pedido > 0:
                    novo_pedido = PedidoContrato()
                    novo_pedido.contrato = contrato
                    novo_pedido.solicitacao = nova_solicitacao
                    if eh_lote:

                        novo_pedido.item = ItemContrato.objects.get(item=resultados.get(id=request.POST.getlist('id')[idx]))
                        novo_pedido.valor = ItemContrato.objects.get(item=resultados.get(id=request.POST.getlist('id')[idx])).get_valor_item_contrato()
                    else:
                        novo_pedido.item = resultados.get(id=request.POST.getlist('id')[idx])
                        novo_pedido.valor = resultados.get(id=request.POST.getlist('id')[idx]).get_valor_item_contrato()


                    novo_pedido.quantidade = valor_pedido

                    novo_pedido.setor = setor
                    novo_pedido.pedido_por = request.user
                    novo_pedido.pedido_em = datetime.datetime.now()
                    novo_pedido.save()

            messages.success(request, u'Pedido cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % nova_solicitacao.id)
    return render(request, 'informar_quantidades_do_pedido_contrato.html', locals(), RequestContext(request))


@login_required()
def apagar_anexo_pregao(request, item_id):
    anexo = get_object_or_404(AnexoPregao, pk=item_id)
    pregao = anexo.pregao
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        anexo.delete()
        messages.success(request, u'Anexo removido com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/' % pregao.id)
    else:
        raise PermissionDenied

@login_required()
def apagar_anexo_contrato(request, item_id):
    anexo = get_object_or_404(AnexoContrato, pk=item_id)
    contrato = anexo.contrato
    if request.user.has_perm('base.pode_gerenciar_contrato') and contrato.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        anexo.delete()
        messages.success(request, u'Anexo removido com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_contrato/%s/#anexos' % contrato.id)
    else:
        raise PermissionDenied

@login_required()
def editar_anexo_contrato(request, item_id):
    anexo = get_object_or_404(AnexoContrato, pk=item_id)
    contrato = anexo.contrato
    if request.user.has_perm('base.pode_gerenciar_contrato') and contrato.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Editar Anexo - %s' % contrato
        form = AnexoContratoForm(request.POST or None, request.FILES or None, instance=anexo)
        if form.is_valid():
            form.save()
            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_contrato/%s/#anexos' % contrato.id)

        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied





@login_required()
def apagar_anexo_credenciamento(request, item_id):
    anexo = get_object_or_404(AnexoCredenciamento, pk=item_id)
    credenciamento = anexo.credenciamento
    if request.user.has_perm('base.pode_gerenciar_contrato') and credenciamento.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        anexo.delete()
        messages.success(request, u'Anexo removido com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/#anexos' % credenciamento.id)
    else:
        raise PermissionDenied

@login_required()
def editar_anexo_credenciamento(request, item_id):
    anexo = get_object_or_404(AnexoCredenciamento, pk=item_id)
    credenciamento = anexo.credenciamento
    if request.user.has_perm('base.pode_gerenciar_contrato') and credenciamento.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Editar Anexo - %s' % credenciamento
        form = AnexoContratoForm(request.POST or None, request.FILES or None, instance=anexo)
        if form.is_valid():
            form.save()
            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/#anexos' % credenciamento.id)

        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied




@login_required()
def cadastrar_anexo_arp(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    if request.user.has_perm('base.pode_gerenciar_contrato'):
        title=u'Cadastrar Anexo - %s' % ata
        form = AnexoARPForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            o = form.save(False)
            o.ata = ata
            o.cadastrado_por = request.user
            o.cadastrado_em = datetime.datetime.now()
            o.save()
            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/#anexos' % ata.id)

        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def apagar_anexo_arp(request, item_id):
    anexo = get_object_or_404(AnexoAtaRegistroPreco, pk=item_id)
    ata = anexo.ata
    if request.user.has_perm('base.pode_gerenciar_contrato') and ata.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        anexo.delete()
        messages.success(request, u'Anexo removido com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/#anexos' % ata.id)
    else:
        raise PermissionDenied

@login_required()
def editar_anexo_arp(request, item_id):
    anexo = get_object_or_404(AnexoAtaRegistroPreco, pk=item_id)
    ata = anexo.ata
    if request.user.has_perm('base.pode_gerenciar_contrato') and ata.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Editar Anexo - %s' % ata
        form = AnexoContratoForm(request.POST or None, request.FILES or None, instance=anexo)
        if form.is_valid():
            form.save()
            messages.success(request, u'Anexo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/#anexos' % ata.id)

        return render(request, 'cadastrar_anexo_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied



@login_required()
def gerar_pedido_fornecedores(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)


    configuracao = get_config(solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/pedidos/%s.pdf' % solicitacao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/pedidos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/pedidos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    eh_lote = False
    pedidos = None
    if PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).order_by('item')
    elif PedidoContrato.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoContrato.objects.filter(solicitacao=solicitacao).order_by('item')
    elif PedidoCredenciamento.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoCredenciamento.objects.filter(solicitacao=solicitacao).order_by('item')


    tabela = {}


    for pedido in pedidos:
        if solicitacao.credenciamento_origem:
            chave = u'%s' % pedido.fornecedor
        else:
            if pedido.item.participante:
                chave = u'%s' % pedido.item.participante.fornecedor
            else:
                chave = u'%s' % pedido.item.fornecedor
        tabela[chave] = dict(pedidos = list(), total = 0)

    for pedido in pedidos:
        if solicitacao.credenciamento_origem:
            chave = u'%s' % pedido.fornecedor
        else:
            if pedido.item.participante:
                chave = u'%s' % pedido.item.participante.fornecedor
            else:
                chave = u'%s' % pedido.item.fornecedor
        tabela[chave]['pedidos'].append(pedido)
        valor = tabela[chave]['total']
        valor = valor + (pedido.item.valor*pedido.quantidade)
        tabela[chave]['total'] = valor


    resultado = collections.OrderedDict(sorted(tabela.items()))



    data = {'configuracao': configuracao, 'logo': logo, 'resultado': resultado, 'data_emissao': data_emissao, 'eh_lote': eh_lote}

    template = get_template('gerar_pedido_fornecedores.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def apagar_lote(request, item_id, pregao_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        itens_do_lote = ItemLote.objects.filter(lote=item)
        PropostaItemPregao.objects.filter(item__in=itens_do_lote.values_list('item', flat=True)).delete()
        ItemLote.objects.filter(lote=item).delete()
        item.delete()
        return HttpResponseRedirect(u'/base/pregao/%s/#lotes' % pregao_id)
    else:
        raise PermissionDenied


@login_required()
def informar_valor_final_item_lote(request, item_id, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
        lote =ItemLote.objects.filter(item=item)[0].lote
        ids_itens_do_lote = ItemLote.objects.filter(lote=lote).values_list('item', flat=True)
        vencedor = lote.get_empresa_vencedora()
        title=u'Informar Valor Unitário Final do %s - %s' % (item , lote)
        form = ValorFinalItemLoteForm(request.POST or None)
        if form.is_valid():
            if form.cleaned_data.get('valor') > item.get_valor_total_proposto() or form.cleaned_data.get('valor') > item.valor_medio:
                messages.error(request, u'O valor não pode ser maior do que o valor unitário proposto: %s nem do que o valor máximo do item: %s.' % ( item.get_valor_total_proposto(), item.valor_medio))
                return HttpResponseRedirect(u'/base/informar_valor_final_item_lote/%s/%s/' % (item.id, pregao.id))


            valor = PropostaItemPregao.objects.filter(participante=vencedor, item__in=ids_itens_do_lote).aggregate(total=Sum('valor_item_lote'))['total'] or 0
            if (valor + form.cleaned_data.get('valor')) > lote.get_total_lance_ganhador():
                messages.error(request, u'O valor informado faz o valor total dos itens do lote ultrapassar o valor do lance ganhador.')
                return HttpResponseRedirect(u'/base/informar_valor_final_item_lote/%s/%s/' % (item.id, pregao.id))

            valor_final = form.cleaned_data.get('valor') * item.quantidade
            PropostaItemPregao.objects.filter(participante=vencedor, item=item).update(valor_item_lote=valor_final)

            messages.success(request, u'Valor cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao_id)
        return render(request, 'informar_valor_final_item_lote.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def informar_valor_final_itens_lote(request, lote_id, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        lote =get_object_or_404(ItemSolicitacaoLicitacao, pk=lote_id)
        itens = ItemLote.objects.filter(lote=lote)
        ids_itens_do_lote = ItemLote.objects.filter(lote=lote).values_list('item', flat=True)
        vencedor = lote.get_empresa_vencedora()
        title=u'Informar Valor Unitário Final do %s' % (lote)
        form = ValorFinalItemLoteForm(request.POST or None)
        if request.POST:
            contador = 0
            for id_do_item in request.POST.getlist('id_item'):
                item = ItemSolicitacaoLicitacao.objects.get(pk=id_do_item)
                valor_informado = Decimal(request.POST.getlist('itens')[contador])
                if valor_informado > item.get_valor_total_proposto() or valor_informado > item.valor_medio:
                    messages.error(request, u'O valor não pode ser maior do que o valor unitário proposto (%s) nem do que o valor máximo do item: %s.' % (item.get_valor_total_proposto(), item.valor_medio))
                    return HttpResponseRedirect(u'/base/informar_valor_final_itens_lote/%s/%s/' % (lote.id, pregao.id))

                valor_final = valor_informado * item.quantidade
                valor = PropostaItemPregao.objects.filter(participante=vencedor, item__in=ids_itens_do_lote).exclude(item=item).aggregate(total=Sum('valor_item_lote'))['total'] or 0
                if (valor + valor_final) > lote.get_total_lance_ganhador():
                    messages.error(request, u'O valor informado faz o valor total dos itens do lote ultrapassar o valor do lance ganhador: %s.' % lote.get_total_lance_ganhador())
                    return HttpResponseRedirect(u'/base/informar_valor_final_itens_lote/%s/%s/' % (lote.id, pregao.id))


                PropostaItemPregao.objects.filter(participante=vencedor, item=item).update(valor_item_lote=valor_final)
                contador += 1
            messages.success(request, u'Valor cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao_id)
        return render(request, 'informar_valor_final_itens_lote.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def gerar_ordem_compra(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if solicitacao.tem_valor_acima_permitido():
        messages.error(request, u'O valor desta solicitação está acima do permitido.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao_id)

    if solicitacao.tem_empate_propostas():
        messages.error(request, u'Esta solicitação possui mais de uma proposta empatada com o melhor valor.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao_id)

    id_sessao = "%s_solicitacao" % (request.user.pessoafisica.id)
    request.session[id_sessao] = solicitacao.id
    title = u'Gerar Ordem de Compra/Serviço - %s' % solicitacao
    form = CriarOrdemForm(request.POST or None)
    if form.is_valid():
        o = form.save(False)
        o.solicitacao = solicitacao
        o.save()
        messages.success(request, u'Ordem de Compra/Serviço gerada com sucesso.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao_id)
    return render(request, 'gerar_ordem_compra.html', locals(), RequestContext(request))


@login_required()
def ver_ordem_compra(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)

    configuracao = get_config(solicitacao.setor_origem.secretaria)
    config_geral = get_config_geral()
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/ordens_compra/%s.pdf' % solicitacao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/ordens_compra')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/ordens_compra'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    ordem = OrdemCompra.objects.get(solicitacao=solicitacao)

    eh_lote = solicitacao.eh_lote()

    tabela = {}
    pregao = contrato = ata = None
    pregao = solicitacao.get_pregao()
    credenciamento = None
    if PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).order_by('item')
        ata = get_object_or_404(AtaRegistroPreco, pk=pedidos[0].ata.id)
    elif PedidoContrato.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoContrato.objects.filter(solicitacao=solicitacao).order_by('item')
        contrato = get_object_or_404(Contrato, pk=pedidos[0].contrato.id)
    elif PedidoCredenciamento.objects.filter(solicitacao=solicitacao).exists():
        pedidos = PedidoCredenciamento.objects.filter(solicitacao=solicitacao).order_by('item')
        credenciamento = get_object_or_404(Credenciamento, pk=pedidos[0].credenciamento.id)


    tabela = {}


    for pedido in pedidos:
        if credenciamento:
            chave = u'%s' % pedido.fornecedor
            fornecedor = pedido.fornecedor

        else:
            if pedido.item.participante:
                chave = u'%s' % pedido.item.participante.fornecedor
                fornecedor = pedido.item.participante.fornecedor
            else:
                chave = u'%s' % pedido.item.fornecedor
                fornecedor = pedido.item.fornecedor
        tabela[chave] = dict(pedidos = list(), total = 0)

    for pedido in pedidos:
        if credenciamento:
            chave = u'%s' % pedido.fornecedor
            fornecedor = pedido.fornecedor

        else:

            if pedido.item.participante:
                chave = u'%s' % pedido.item.participante.fornecedor
            else:
                chave = u'%s' % pedido.item.fornecedor
        tabela[chave]['pedidos'].append(pedido)
        valor = tabela[chave]['total']
        valor = valor + (pedido.item.valor*pedido.quantidade)
        tabela[chave]['total'] = valor


    resultado = collections.OrderedDict(sorted(tabela.items()))

    cpf_secretario = configuracao.responsavel.cpf
    cpf_secretario_formatado = "%s.%s.%s-%s" % ( cpf_secretario[0:3], cpf_secretario[3:6], cpf_secretario[6:9], cpf_secretario[9:11] )

    cpf_ordenador = configuracao.ordenador_despesa.cpf
    cpf_ordenador_formatado = "%s.%s.%s-%s" % ( cpf_ordenador[0:3], cpf_ordenador[3:6], cpf_ordenador[6:9], cpf_ordenador[9:11] )
    data = {'config_geral': config_geral, 'cpf_ordenador_formatado': cpf_ordenador_formatado, 'cpf_secretario_formatado': cpf_secretario_formatado, 'solicitacao': solicitacao, 'pregao': pregao, 'ata':ata, 'contrato':contrato, 'credenciamento': credenciamento, 'configuracao': configuracao, 'logo': logo, 'fornecedor': fornecedor, 'resultado': resultado, 'data_emissao': data_emissao, 'eh_lote': eh_lote, 'ordem': ordem}

    template = get_template('ver_ordem_compra.html')

    html  = template.render(Context(data))
    if 'xls' in request.GET:
        nome = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/ordem_compra')
        file_path = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/ordem_compra.xls')
        rb = open_workbook(file_path,formatting_info=True)

        wb = copy(rb) # a writable copy (I can't read values out of this, only write to it)
        w_sheet = wb.get_sheet(0) # the sheet to write to within the writable copy

        sheet = rb.sheet_by_name("Sheet1")

        w_sheet.write(0, 1, u'Ordem de Compra/Serviço - %s' % ordem.numero)
        w_sheet.write(1, 1, configuracao.nome)
        if solicitacao.arp_origem and solicitacao.arp_origem.solicitacao.get_pregao():
            texto = u'Pregão %s' % solicitacao.arp_origem.solicitacao.get_pregao()
        elif solicitacao.contrato_origem and solicitacao.contrato_origem.solicitacao.get_pregao():
            texto = u'Pregão %s' % solicitacao.contrato_origem.solicitacao.get_pregao()
        elif solicitacao.credenciamento_origem and solicitacao.credenciamento_origem.solicitacao.get_pregao():
            texto = u'Pregão %s' % solicitacao.credenciamento_origem.solicitacao.get_pregao()
        else:
            texto = u'Origem %s' % solicitacao


        w_sheet.write(2, 1, texto)
        w_sheet.write(3, 1,  ordem.solicitacao.objeto)
        w_sheet.write(4, 1,  ordem.solicitacao.objetivo)
        w_sheet.write(5, 1,  fornecedor.razao_social)
        w_sheet.write(6, 1,  fornecedor.endereco)
        w_sheet.write(7, 1,  fornecedor.cnpj )
        if hasattr(fornecedor, 'banco') and fornecedor.banco:
            texto = u'%s / Agência: %s / Conta: %s'% (fornecedor.banco, fornecedor.agencia, fornecedor.conta)
        else:
            texto = u'Não Informado.'
        w_sheet.write(8, 1,  texto)
        contador = 12
        conta_item = 1

        for item in resultado.items():

            for pedido in item[1]['pedidos']:

                row_index = contador + 1


                w_sheet.write(row_index, 0, conta_item)
                if pedido.item.item and pedido.item.item.material:
                    w_sheet.write(row_index, 1, pedido.item.item.material.nome )
                else:
                    w_sheet.write(row_index, 1, pedido.item.material.nome )
                if pedido.item.item and pedido.item.item.unidade:
                    w_sheet.write(row_index, 2, pedido.item.item.unidade.nome)
                else:
                    w_sheet.write(row_index, 2, pedido.item.unidade.nome)


                w_sheet.write(row_index, 3, format_quantidade(pedido.quantidade))
                w_sheet.write(row_index, 4, format_money(pedido.valor))
                w_sheet.write(row_index, 5, format_money(pedido.get_total()))
                contador += 1
                conta_item += 1

            w_sheet.write(contador+2, 0, u'Total')
            w_sheet.write(contador+2, 5, format_money(item[1]['total']))


        salvou = nome + u'_%s' % solicitacao.id + '.xls'
        wb.save(salvou)

        arquivo = open(salvou, "rb")


        content_type = 'application/vnd.ms-excel'
        response = HttpResponse(arquivo.read(), content_type=content_type)
        nome_arquivo = salvou.split('/')[-1]
        response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
        arquivo.close()
        os.unlink(salvou)
        return response
    else:
        pdf_file = open(caminho_arquivo, "w+b")
        pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                encoding='utf-8')
        pdf_file.close()
        file = open(caminho_arquivo, "r")
        pdf = file.read()
        file.close()
        return HttpResponse(pdf, 'application/pdf')




@login_required()
def excluir_ordem_compra_dispensa(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if solicitacao.tem_ordem_compra() and request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and solicitacao.recebida_setor(request.user.pessoafisica.setor):
        OrdemCompra.objects.filter(solicitacao=solicitacao).delete()
        messages.success(request, u'Ordem de Compra/Serviço excluída com sucesso.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao_id)

    else:
        raise PermissionDenied



@login_required()
def ver_ordem_compra_dispensa(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    config_geral = get_config_geral()
    configuracao = get_config(solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/ordens_compra/%s.pdf' % solicitacao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/ordens_compra')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/ordens_compra'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    ordem = OrdemCompra.objects.get(solicitacao=solicitacao)

    lista = list()
    dicionario = {}
    for pesquisa in PesquisaMercadologica.objects.filter(solicitacao=solicitacao):
        total = ItemPesquisaMercadologica.objects.filter(pesquisa=pesquisa, ativo=True).aggregate(soma=Sum('valor_maximo'))['soma']
        if total:
            lista.append([pesquisa.id, total])
            dicionario[pesquisa.id] = total
    resultado = sorted(dicionario.items(), key=lambda x: x[1])
    fornecedor = PesquisaMercadologica.objects.get(id=resultado[0][0])
    itens = ItemPesquisaMercadologica.objects.filter(pesquisa=resultado[0][0]).order_by('item')
    total = 0

    for item in itens:
        total += item.get_total()


    cpf_secretario = configuracao.responsavel.cpf
    cpf_secretario_formatado = "%s.%s.%s-%s" % ( cpf_secretario[0:3], cpf_secretario[3:6], cpf_secretario[6:9], cpf_secretario[9:11] )

    cpf_ordenador = configuracao.ordenador_despesa.cpf
    cpf_ordenador_formatado = "%s.%s.%s-%s" % ( cpf_ordenador[0:3], cpf_ordenador[3:6], cpf_ordenador[6:9], cpf_ordenador[9:11] )
    data = {'config_geral': config_geral, 'cpf_ordenador_formatado': cpf_ordenador_formatado, 'cpf_secretario_formatado': cpf_secretario_formatado, 'solicitacao': solicitacao, 'pregao': pregao, 'total':total, 'itens': itens, 'fornecedor': fornecedor, 'configuracao': configuracao, 'logo': logo,  'data_emissao': data_emissao, 'ordem': ordem}

    template = get_template('ver_ordem_compra_dispensa.html')

    html  = template.render(Context(data))
    if 'xls' in request.GET:
        nome = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/ordem_compra')
        file_path = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/ordem_compra.xls')
        rb = open_workbook(file_path,formatting_info=True)

        wb = copy(rb) # a writable copy (I can't read values out of this, only write to it)
        w_sheet = wb.get_sheet(0) # the sheet to write to within the writable copy

        sheet = rb.sheet_by_name("Sheet1")

        w_sheet.write(0, 1, u'Ordem de Compra/Serviço - %s' % ordem.numero)
        w_sheet.write(1, 1, configuracao.nome)
        if solicitacao.arp_origem and solicitacao.arp_origem.solicitacao.get_pregao():
            texto = u'Pregão %s' % solicitacao.arp_origem.solicitacao.get_pregao()
        elif solicitacao.contrato_origem and solicitacao.contrato_origem.solicitacao.get_pregao():
            texto = u'Pregão %s' % solicitacao.contrato_origem.solicitacao.get_pregao()
        elif solicitacao.credenciamento_origem and solicitacao.credenciamento_origem.solicitacao.get_pregao():
            texto = u'Pregão %s' % solicitacao.credenciamento_origem.solicitacao.get_pregao()
        else:
            texto = u'Origem %s' % solicitacao


        w_sheet.write(2, 1, texto)
        w_sheet.write(3, 1,  ordem.solicitacao.objeto)
        w_sheet.write(4, 1,  ordem.solicitacao.objetivo)
        w_sheet.write(5, 1,  fornecedor.razao_social)
        w_sheet.write(6, 1,  fornecedor.endereco)
        w_sheet.write(7, 1,  fornecedor.cnpj )
        if hasattr(fornecedor, 'banco') and fornecedor.banco:
            texto = u'%s / Agência: %s / Conta: %s'% (fornecedor.banco, fornecedor.agencia, fornecedor.conta)
        else:
            texto = u'Não Informado.'
        w_sheet.write(8, 1,  texto)
        contador = 12
        conta_item = 1


        for proposta in itens:

            row_index = contador + 1

            w_sheet.write(row_index, 0,  proposta.item.item)
            w_sheet.write(row_index, 1, u'%s - MARCA: %s' % (proposta.item.material.nome, proposta.marca))
            w_sheet.write(row_index, 2, proposta.item.unidade.nome)
            w_sheet.write(row_index, 3, format_quantidade(proposta.item.quantidade))
            w_sheet.write(row_index, 4, format_money(proposta.valor_maximo))
            w_sheet.write(row_index, 5, format_money(proposta.get_total()))
            contador += 1
            conta_item += 1

        w_sheet.write(contador+2, 0, u'Total')
        w_sheet.write(contador+2, 5, format_money(total))


        salvou = nome + u'_%s' % solicitacao.id + '.xls'
        wb.save(salvou)

        arquivo = open(salvou, "rb")


        content_type = 'application/vnd.ms-excel'
        response = HttpResponse(arquivo.read(), content_type=content_type)
        nome_arquivo = salvou.split('/')[-1]
        response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
        arquivo.close()
        os.unlink(salvou)
        return response
    else:
        pdf_file = open(caminho_arquivo, "w+b")
        pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                encoding='utf-8')
        pdf_file.close()
        file = open(caminho_arquivo, "r")
        pdf = file.read()
        file.close()
        return HttpResponse(pdf, 'application/pdf')

@login_required()
def registrar_adjudicacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Registrar Adjudicação'
        form = RegistrarAdjudicacaoForm(request.POST or None, instance=pregao)
        if form.is_valid():
            o = form.save(False)
            o.situacao = Pregao.ADJUDICADO
            o.save()
            messages.success(request, u'Data de adjudicação registrada com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/#homologacao' % pregao.id)
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def registrar_homologacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.pessoafisica == pregao.solicitacao.setor_origem.secretaria.ordenador_despesa:
        if pregao.eh_credenciamento():
            title=u'Credenciar'
        else:
            title=u'Registrar Homologação'
        form = RegistrarHomologacaoForm(request.POST or None, instance=pregao)
        if form.is_valid():
            o = form.save(False)
            o.ordenador_despesa = request.user.pessoafisica
            o.situacao = Pregao.CONCLUIDO
            o.save()
            messages.success(request, u'Data de homologação registrada com sucesso. Prossiga com o envio do termo assinado')
            return HttpResponseRedirect(u'/base/upload_termo_homologacao/%s/' % pregao.id)
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def termo_homologacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    config_geral = get_config_geral()
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/extratos/%s.pdf' % pregao.id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/extratos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/extratos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)



    tabela = {}

    eh_lote = pregao.criterio.id == CriterioPregao.LOTE

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(ordem=1, item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(itens = list(), total = 0)
    total_geral = 0
    for item in itens_pregao.order_by('item'):

        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['itens'].append(item.item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor


    for item in tabela:
        total_geral = total_geral + tabela[item]['total']

    fracassados = list()
    for item in itens_pregao.filter(situacao=ItemSolicitacaoLicitacao.FRACASSADO):
        fracassados.append(item.item)

    resultado = collections.OrderedDict(sorted(tabela.items()))


    data = {'pregao': pregao, 'eh_lote': eh_lote, 'configuracao': configuracao, 'logo': logo, 'resultado': resultado, 'total_geral': total_geral, 'fracassados': fracassados, 'config_geral': config_geral}
    if pregao.eh_pregao():
        template = get_template('termo_homologacao.html')
    else:
        template = get_template('termo_homologacao_e_adjudicacao.html')


    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()


    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def visualizar_contrato(request, solicitacao_id):
    contrato = get_object_or_404(Contrato, pk=solicitacao_id)
    title = u'Contrato: %s - Fornecedor: %s' % (contrato.numero, contrato.get_fornecedor())
    pedidos = PedidoContrato.objects.filter(contrato=contrato).order_by('item__material', 'setor')
    pode_gerenciar = contrato.solicitacao.recebida_setor(request.user.pessoafisica.setor)
    eh_gerente = request.user.groups.filter(name='Gerente') and pode_gerenciar
    itens = ItemContrato.objects.filter(contrato=contrato).order_by('item__item')

    return render(request, 'visualizar_contrato.html', locals(), RequestContext(request))

@login_required()
def visualizar_ata_registro_preco(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    title = u'Ata de Registro de Preço N° %s' % ata.numero

    pode_gerenciar = ata.solicitacao.recebida_setor(request.user.pessoafisica.setor)
    eh_gerente = request.user.groups.filter(name='Gerente') and pode_gerenciar


    pedidos = PedidoAtaRegistroPreco.objects.filter(ata=ata).order_by('setor__secretaria', 'item__material', 'pedido_em')
    tabela  = {}
    itens = ata.itemataregistropreco_set.all()
    if ata.adesao:
        participantes = Fornecedor.objects.filter(id__in=itens.values_list('fornecedor', flat=True))
    else:
        participantes = ParticipantePregao.objects.filter(id__in=itens.values_list('participante', flat=True))

    materiais  = dict()
    secretarias =  pedidos.values('setor__secretaria__nome').order_by('setor__secretaria__nome').distinct('setor__secretaria__nome')
    # for num in secretarias:
    #     chave = '%s' % num['setor__secretaria__nome']
    #     tabela[chave] = materiais
    #     for item in pedidos.filter(setor__secretaria__nome=chave):
    #         if item.item.fornecedor:
    #             nome = u'Fornecedor: %s' % (item.item.fornecedor.razao_social)
    #         else:
    #             nome = u'Fornecedor: %s' % (item.item.participante.fornecedor.razao_social)
    #         materiais[nome] = dict(pedidos=list())
    #
    # for pedido in pedidos:
    #     if pedido.item.fornecedor:
    #         nome = u'Fornecedor: %s' % (pedido.item.fornecedor.razao_social)
    #     else:
    #         nome = u'Fornecedor: %s' % (pedido.item.participante.fornecedor.razao_social)
    #     materiais[nome]['pedidos'].append(pedido)
    #
    # resultado = collections.OrderedDict(sorted(tabela.items()))


    tem_transferencias = TransferenciaItemARP.objects.filter(item__ata=ata)

    return render(request, 'visualizar_ata_registro_preco.html', locals(), RequestContext(request))

@login_required()
def visualizar_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)
    title = u'Credenciamento N° %s' % credenciamento.numero

    pode_gerenciar = credenciamento.solicitacao.recebida_setor(request.user.pessoafisica.setor)
    eh_gerente = request.user.groups.filter(name='Gerente') and pode_gerenciar

    pedidos = PedidoCredenciamento.objects.filter(credenciamento=credenciamento).order_by('item__material', 'setor')
    tabela  = {}

    materiais  = dict()
    secretarias =  pedidos.values('setor__secretaria__nome').order_by('setor__secretaria__nome').distinct('setor__secretaria__nome')
    for num in secretarias:
        chave = '%s' % num['setor__secretaria__nome']
        tabela[chave] = materiais
        for item in pedidos.filter(setor__secretaria__nome=chave):
            nome = u'Fornecedor: %s' % (item.fornecedor.razao_social)
            materiais[nome] = dict(pedidos=list())

    for pedido in pedidos:
        nome = u'Fornecedor: %s' % (pedido.fornecedor.razao_social)
        materiais[nome]['pedidos'].append(pedido)

    resultado = collections.OrderedDict(sorted(tabela.items()))

    return render(request, 'visualizar_credenciamento.html', locals(), RequestContext(request))

@login_required()
def liberar_solicitacao_contrato(request, solicitacao_id, origem):
    contrato = get_object_or_404(Contrato, pk=solicitacao_id)
    if request.user.has_perm('base.pode_gerenciar_contrato') and contrato.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        if origem == u'1':
            contrato.liberada_compra = True
            contrato.suspenso = False
        elif origem == u'2':
            contrato.liberada_compra = False
            contrato.suspenso = True

        contrato.save()
        return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' % contrato.id)
    else:
        raise PermissionDenied

@login_required()
def liberar_solicitacao_ata(request, ata_id, origem):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    if not ata.adesao and request.user.has_perm('base.pode_gerenciar_contrato') and ata.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        if origem == u'1':
            ata.liberada_compra = True
            ata.suspenso = False
        elif origem == u'2':
            ata.liberada_compra = False
            ata.suspenso = True
        ata.save()
        return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % ata.id)
    else:
        raise PermissionDenied

@login_required()
def liberar_solicitacao_credenciamento(request, credenciamento_id, origem):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)
    if request.user.has_perm('base.pode_gerenciar_contrato') and credenciamento.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        if origem == u'1':
            credenciamento.liberada_compra = True
            credenciamento.suspenso = False
        elif origem == u'2':
            credenciamento.liberada_compra = False
            credenciamento.suspenso = True
        credenciamento.save()
        return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/' % credenciamento.id)
    else:
        raise PermissionDenied

@login_required()
def definir_vigencia_contrato(request, contrato_id):
    title=u'Definir Vigência do Contrato'
    contrato = get_object_or_404(Contrato, pk=contrato_id)
    if request.user.has_perm('base.pode_gerenciar_contrato') and contrato.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = DefinirVigenciaContratoForm(request.POST or None, instance=contrato)
        if form.is_valid():
            form.save()
            messages.success(request, u'Data de Vigência registrada com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' % contrato.id)

        return render(request, 'definir_vigencia_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied



@login_required()
def lista_documentos(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    title = u'Lista de Documentos'
    documentos = DocumentoSolicitacao.objects.filter(solicitacao=solicitacao)
    minha_secretaria = request.user.pessoafisica.setor.secretaria
    ve_tudo = minha_secretaria == solicitacao.setor_origem.secretaria
    listas = None
    if not ve_tudo:
        listas = solicitacao.get_pedidos_secretarias(minha_secretaria)

    return render(request, 'lista_documentos.html', locals(), RequestContext(request))


def libreoffice_new_line(tokens, align_center='', font_size=17):
    if len(tokens)>1:
        if align_center:
            align_center = '<w:jc w:val="center"/>'
        out = ['</w:t></w:r></w:p>']
        for token in tokens:
            out.append('<w:p><w:pPr><w:spacing w:after="0"/>%s</w:pPr><w:r><w:rPr><w:sz w:val="%s"/></w:rPr><w:t>%s' % (align_center, font_size, token))
            out.append('</w:t></w:r></w:p>')
        del(out[-1])
        return ''.join(out)
    return tokens[0]


@login_required()
def memorando(request, solicitacao_id):
    import tempfile
    import zipfile
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)

    municipio = None
    if get_config_geral():
        municipio = get_config_geral().municipio

    itens = []
    quantidades = []
    unidades  = []
    descricoes = []

    for item in solicitacao.itemsolicitacaolicitacao_set.all():
        itens.append(item)
        quantidades.append(item.quantidade)
        unidades.append(item.unidade)
        descricoes.append(item.material.nome)

    dicionario = {
        '#NUM#' : solicitacao.num_memorando,
        '#MUNICIPIO#' : municipio or u'-',
        '#DATA#': datetime.date.today(),
        '#OBJETIVO#': solicitacao.objetivo,
        '#OBJETO#': solicitacao.objeto,
        '#IT#': libreoffice_new_line(itens or '-'),
        '#QUANT#': libreoffice_new_line(quantidades or '-'),
        '#UN#': libreoffice_new_line(unidades or '-'),
        '#DES#': libreoffice_new_line(descricoes or '-'),

    }
    template_docx = zipfile.ZipFile(os.path.join(settings.MEDIA_ROOT, 'upload/modelos/memorando.docx'))
    new_docx = zipfile.ZipFile('%s.docx' % tempfile.mktemp(), "a")

    tmp_xml_file = open(template_docx.extract("word/document.xml", tempfile.mkdtemp()))
    tempXmlStr = tmp_xml_file.read()
    tmp_xml_file.close()
    os.unlink(tmp_xml_file.name)

    for key in dicionario.keys():
        value = unicode(dicionario.get(key)).encode("utf8")
        tempXmlStr = tempXmlStr.replace(key, value)

    tmp_xml_file =  open(tempfile.mktemp(), "w+")
    tmp_xml_file.write(tempXmlStr)
    tmp_xml_file.close()

    for arquivo in template_docx.filelist:
        if not arquivo.filename == "word/document.xml":
            new_docx.writestr(arquivo.filename, template_docx.read(arquivo))

    new_docx.write(tmp_xml_file.name, "word/document.xml")

    template_docx.close()
    new_docx.close()
    os.unlink(tmp_xml_file.name)


    # Caso não seja informado, deverá retornar o caminho para o arquivo DOCX processado.
    caminho_arquivo =  new_docx.filename
    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")


    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response


@login_required()
def termo_referencia(request, solicitacao_id):
    import tempfile
    import zipfile
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)

    municipio = None
    if get_config_geral():
        municipio = get_config_geral().municipio

    itens = []
    quantidades = []
    unidades  = []
    descricoes = []

    for item in solicitacao.itemsolicitacaolicitacao_set.all():
        itens.append(item)
        quantidades.append(item.quantidade)
        unidades.append(item.unidade)
        descricoes.append(item.material.nome)

    dicionario = {
        '#NUM#' : solicitacao.num_memorando,
        '#MUNICIPIO#' : municipio or u'-',
        '#DATA#': datetime.date.today(),
        '#OBJETIVO#': solicitacao.objetivo,
        '#OBJETO#': solicitacao.objeto,
        '#JUST#': solicitacao.justificativa,
        '#IT#': libreoffice_new_line(itens or '-'),
        '#QUANT#': libreoffice_new_line(quantidades or '-'),
        '#UN#': libreoffice_new_line(unidades or '-'),
        '#DES#': libreoffice_new_line(descricoes or '-'),

    }
    template_docx = zipfile.ZipFile(os.path.join(settings.MEDIA_ROOT, 'upload/modelos/termo_referencia.docx'))
    new_docx = zipfile.ZipFile('%s.docx' % tempfile.mktemp(), "a")

    tmp_xml_file = open(template_docx.extract("word/document.xml", tempfile.mkdtemp()))
    tempXmlStr = tmp_xml_file.read()
    tmp_xml_file.close()
    os.unlink(tmp_xml_file.name)

    for key in dicionario.keys():
        value = unicode(dicionario.get(key)).encode("utf8")
        tempXmlStr = tempXmlStr.replace(key, value)

    tmp_xml_file =  open(tempfile.mktemp(), "w+")
    tmp_xml_file.write(tempXmlStr)
    tmp_xml_file.close()

    for arquivo in template_docx.filelist:
        if not arquivo.filename == "word/document.xml":
            new_docx.writestr(arquivo.filename, template_docx.read(arquivo))

    new_docx.write(tmp_xml_file.name, "word/document.xml")

    template_docx.close()
    new_docx.close()
    os.unlink(tmp_xml_file.name)


    # Caso não seja informado, deverá retornar o caminho para o arquivo DOCX processado.
    caminho_arquivo =  new_docx.filename
    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")


    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response

@login_required()
def apagar_documento(request, documento_id):
    documento = get_object_or_404(DocumentoSolicitacao, pk=documento_id)
    if documento.cadastrado_por == request.user:
        solicitacao = documento.solicitacao.id
        documento.delete()
        messages.success(request, u'Documento apagado com sucesso.')
        return HttpResponseRedirect(u'/base/lista_documentos/%s/' % solicitacao)
    else:
        raise PermissionDenied

@login_required()
def editar_fornecedor(request, fornecedor_id):
    fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
    title=u'Editar Fornecedor - %s' % fornecedor
    form = FornecedorForm(request.POST or None, instance=fornecedor)
    if form.is_valid():
        form.save()
        messages.success(request, u'Fornecedor editado com sucesso.')
        return HttpResponseRedirect(u'/base/ver_fornecedores/')

    return render(request, 'editar_fornecedor.html', locals(), RequestContext(request))

@login_required()
def cadastrar_fornecedor(request, opcao):
    title=u'Cadastrar Fornecedor'
    form = FornecedorForm(request.POST or None)
    if form.is_valid():
        form.save()

        messages.success(request, u'Fornecedor cadastrado com sucesso.')
        if opcao == u'0':
            return HttpResponseRedirect(u'/base/ver_fornecedores/')
        else:
            return HttpResponseRedirect(u'/base/cadastra_participante_pregao/%s/' % int(opcao))


    return render(request, 'editar_fornecedor.html', locals(), RequestContext(request))


@login_required()
def remover_participante_pregao(request, participante_id):
    participante = get_object_or_404(ParticipantePregao, pk=participante_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and participante.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        pregao = participante.pregao.id
        participante.delete()
        messages.success(request, u'Participante removido com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/' % pregao)
    else:
        raise PermissionDenied

@login_required()
def editar_pregao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    solicitacao = pregao.solicitacao
    if (request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor)) or request.user.is_superuser:
        title = u'Editar %s' % pregao.modalidade
        form = PregaoForm(request.POST or None, instance=pregao, solicitacao=pregao.solicitacao, request=request)
        if form.is_valid():
            form.save()
            if not solicitacao.processo and form.cleaned_data.get('num_processo'):
                novo_processo = Processo()
                novo_processo.pessoa_cadastro = request.user
                novo_processo.numero = form.cleaned_data.get('num_processo')
                novo_processo.objeto = form.cleaned_data.get('objeto')
                novo_processo.tipo = Processo.TIPO_MEMORANDO
                novo_processo.setor_origem = request.user.pessoafisica.setor
                novo_processo.save()
                solicitacao.processo = novo_processo
                solicitacao.save()
            messages.success(request, u'Licitação editada com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/' % pregao.id)

        return render(request, 'editar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def upload_termo_homologacao(request, pregao_id):
    title=u'Enviar Termo de Homologação'
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if pregao.eh_credenciamento():
        title = u'Enviar Termo de Credenciamento'
    if request.user.pessoafisica == pregao.solicitacao.setor_origem.secretaria.ordenador_despesa:
        form = UploadTermoHomologacaoForm(request.POST or None, request.FILES or None, instance=pregao)
        if form.is_valid():
            form.save()
            messages.success(request, u'Termo de Homologação cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/ver_pregoes/')

        return render(request, 'upload_termo_homologacao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def gerar_resultado_licitacao_toda(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        for item in ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao):
            item.gerar_resultado(apaga=True)
        messages.success(request, u'Resultado gerado com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao.id)
    else:
        raise PermissionDenied



@login_required()
def gerar_resultado_licitacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        tabela = {}
        for proposta in PropostaItemPregao.objects.filter(pregao=pregao).order_by('valor'):
            chave= '%s' %  proposta.participante.id
            tabela[chave] = dict(total = 0)
        for proposta in PropostaItemPregao.objects.filter(pregao=pregao).order_by('valor'):
            chave= '%s' %  proposta.participante.id

            tabela[chave]['total'] += proposta.valor
        resultado = sorted(tabela.items(), key=lambda x:x[1])
        total = len(resultado)
        indice = 0
        # print resultado
        # tem_empate_ficto = False
        # total_global_vencedor = 0
        # ganhador_eh_beneficiario = False
        while indice < total:
            fornecedor = ParticipantePregao.objects.get(id=resultado[indice][0])
            # if indice == 0:
            #     total_global_vencedor = resultado[indice][0]
            #     ganhador_eh_beneficiario = fornecedor.me_epp
            #
            # elif not tem_empate_ficto and not ganhador_eh_beneficiario:
            #     if fornecedor.me_epp:
            #         limite_lance = total_global_vencedor + (total_global_vencedor*10)/100
            #         if resultado[indice][0] < limite_lance:
            #             tem_empate_ficto = True

            itens = PropostaItemPregao.objects.filter(pregao=pregao, participante=fornecedor).order_by('valor')

            for item in itens:
                if ResultadoItemPregao.objects.filter(item=item.item, participante=fornecedor).exists():
                    ResultadoItemPregao.objects.filter(item=item.item, participante=fornecedor).update(valor=item.valor, ordem=indice+1)
                else:
                    novo_resultado = ResultadoItemPregao()
                    novo_resultado.item = item.item
                    novo_resultado.participante = fornecedor
                    novo_resultado.valor = item.valor
                    novo_resultado.marca = item.marca
                    novo_resultado.ordem = indice+1
                    novo_resultado.situacao = ResultadoItemPregao.CLASSIFICADO
                    novo_resultado.save()

            indice += 1


        messages.success(request, u'Resultado gerado com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao.id)
    else:
        raise PermissionDenied


@login_required()
def gerar_resultado_credenciamento(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao)
        for item in itens:
            for participante in ParticipantePregao.objects.filter(pregao=pregao, desclassificado=False, excluido_dos_itens=False):
                novo_resultado = ResultadoItemPregao()
                novo_resultado.item = item
                novo_resultado.participante = participante
                novo_resultado.valor = item.valor_medio
                #novo_resultado.marca = item.marca
                novo_resultado.ordem = 1
                novo_resultado.situacao = ResultadoItemPregao.CLASSIFICADO
                novo_resultado.save()


        messages.success(request, u'Resultado gerado com sucesso.')
        return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao.id)
    else:
        raise PermissionDenied

@login_required()
def lista_materiais(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)

    configuracao = get_config(solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/pesquisas/rascunhos/%s.pdf' % solicitacao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    pode_ver_preco = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, valor_medio__isnull=False).exists()
    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=False)
    total = 0
    if pode_ver_preco:
        for item in itens:
            if item.valor_medio:
                total += item.quantidade * item.valor_medio


    data = {'solicitacao': solicitacao,'itens': itens, 'configuracao': configuracao, 'logo': logo, 'data_emissao':data_emissao, 'pode_ver_preco': pode_ver_preco, 'total': total}

    template = get_template('lista_materiais.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def lista_materiais_por_secretaria(request, solicitacao_id, secretaria_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    secretaria = get_object_or_404(Secretaria, pk=secretaria_id)

    configuracao = get_config(solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/pesquisas/rascunhos/%s.pdf' % solicitacao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    pode_ver_preco = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, valor_medio__isnull=False).exists()
    itens = ItemQuantidadeSecretaria.objects.filter(item__solicitacao=solicitacao, secretaria=secretaria).order_by('item')
    total = 0
    if pode_ver_preco:
        for item in itens:
            if item.item.valor_medio:
                total += item.quantidade * item.item.valor_medio

    data = {'secretaria': secretaria, 'itens': itens, 'solicitacao': solicitacao,'configuracao': configuracao, 'logo': logo, 'data_emissao':data_emissao, 'pode_ver_preco': pode_ver_preco, 'total': total}

    template = get_template('lista_materiais_por_secretaria.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def documentos_atas(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)


    title= u'Documentos - %s' % ata
    return render(request, 'documentos_atas.html', locals(), RequestContext(request))

@login_required()
def documentos_contratos(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)


    title= u'Documentos - %s' % contrato
    return render(request, 'documentos_contratos.html', locals(), RequestContext(request))

@login_required()
def documentos_credenciamentos(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)


    title= u'Documentos - %s' % credenciamento
    return render(request, 'documentos_credenciamentos.html', locals(), RequestContext(request))

@login_required()
def rejeitar_pesquisa(request, item_pesquisa_id):
    item = get_object_or_404(ItemPesquisaMercadologica, pk=item_pesquisa_id)
    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and item.item.solicitacao.setor_atual == request.user.pessoafisica.setor:
        title=u'Rejeitar Proposta'
        form = RejeitarPesquisaForm(request.POST or None, instance=item)
        if form.is_valid():
            o = form.save(False)
            o.ativo = False
            o.rejeitado_em = datetime.datetime.now()
            o.rejeitado_por = request.user
            o.save()
            elemento = item.item
            registros = ItemPesquisaMercadologica.objects.filter(item=elemento, rejeitado_por__isnull=True)
            if registros.exists():
                total_registros = registros.count()
                soma = registros.aggregate(Sum('valor_maximo'))
                if elemento.solicitacao.pode_gerar_ordem():
                    elemento.valor_medio = elemento.get_melhor_proposta()
                else:
                    elemento.valor_medio = soma['valor_maximo__sum']/total_registros
                elemento.total = elemento.valor_medio * elemento.quantidade

            else:
                elemento.valor_medio = 0
                elemento.total = 0
            messages.success(request, u'Proposta rejeitada com sucesso.')
            return HttpResponseRedirect(u'/base/ver_pesquisa_mercadologica/%s/' % item.item.id)
        return render(request, 'rejeitar_pesquisa.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def excluir_item_pesquisa(request, item_pesquisa_id):
    item = get_object_or_404(ItemPesquisaMercadologica, pk=item_pesquisa_id)
    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and item.item.solicitacao.setor_atual == request.user.pessoafisica.setor:
        elemento = item.item
        id = elemento.id
        item.delete()
        registros = ItemPesquisaMercadologica.objects.filter(item=elemento, rejeitado_por__isnull=True)
        if registros.exists():
            total_registros = registros.count()
            soma = registros.aggregate(Sum('valor_maximo'))
            if elemento.solicitacao.pode_gerar_ordem():
                elemento.valor_medio = elemento.get_melhor_proposta()
            else:
                elemento.valor_medio = soma['valor_maximo__sum']/total_registros
            elemento.total = elemento.valor_medio * elemento.quantidade

        else:
            elemento.valor_medio = 0
            elemento.total = 0
        elemento.save()

        messages.success(request, u'Proposta pelo item excluída com sucesso.')
        return HttpResponseRedirect(u'/base/ver_pesquisa_mercadologica/%s/' % id)
    else:
        raise PermissionDenied

@login_required()
def excluir_pesquisa(request, pesquisa_id):
    item = get_object_or_404(PesquisaMercadologica, pk=pesquisa_id)
    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and item.solicitacao.setor_atual == request.user.pessoafisica.setor:
        solicitacao = item.solicitacao
        ids = list()
        for opcao in  ItemPesquisaMercadologica.objects.filter(pesquisa=item):
           ids.append(opcao.item.id)
        ItemPesquisaMercadologica.objects.filter(pesquisa=item).delete()
        item.delete()

        for id_item in ids:
            elemento = get_object_or_404(ItemSolicitacaoLicitacao, pk=id_item)
            registros = ItemPesquisaMercadologica.objects.filter(item=elemento, rejeitado_por__isnull=True)
            if registros.exists():
                total_registros = registros.count()
                soma = registros.aggregate(Sum('valor_maximo'))
                if elemento.solicitacao.pode_gerar_ordem():
                    elemento.valor_medio = elemento.get_melhor_proposta()
                else:
                    elemento.valor_medio = soma['valor_maximo__sum']/total_registros
                elemento.total = elemento.valor_medio * elemento.quantidade

            else:
                elemento.valor_medio = 0
                elemento.total = 0
            elemento.save()
        messages.success(request, u'Proposta excluída com sucesso.')
        return HttpResponseRedirect(u'/base/ver_pesquisas/%s/' % solicitacao.id)
    else:
        raise PermissionDenied

@login_required()
def relatorio_lista_download_licitacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/pesquisas/rascunhos/%s.pdf' % pregao.id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/pesquisas/rascunhos'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    itens = LogDownloadArquivo.objects.filter(arquivo__pregao=pregao)
    total_downloads = itens.count()
    total_empresas = itens.distinct('cnpj').count()


    data = {'itens': itens, 'pregao': pregao,'configuracao': configuracao, 'logo': logo, 'data_emissao':data_emissao, 'total_downloads': total_downloads, 'total_empresas': total_empresas}

    template = get_template('relatorio_lista_download_licitacao.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def apagar_item(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    solicitacao = item.solicitacao
    if request.user.has_perm('base.pode_cadastrar_solicitacao') and solicitacao.setor_atual == request.user.pessoafisica.setor:
        item.delete()
        solicitacao.reorganiza_itens()
        messages.success(request, u'Item apagado com sucesso.')
        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)
    else:
        raise PermissionDenied

@login_required()
def liberar_licitacao_homologacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        pregao.pode_homologar = True
        pregao.save()
        return HttpResponseRedirect(u'/base/pregao/%s/' % pregao.id)
    else:
        raise PermissionDenied

@login_required()
def registrar_ocorrencia_pregao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title=u'Registrar Ocorrência - %s' % pregao
        form = HistoricoPregaoForm(request.POST or None)
        if form.is_valid():
            o = form.save(False)
            o.pregao = pregao
            o.data = datetime.datetime.now()
            o.save()
            messages.success(request, u'Ocorrência registrada com sucesso.')
            return HttpResponseRedirect(u'/pregao/%s/' % pregao.id)
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied



@login_required()
def ata_sessao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)

    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)


    municipio = None
    if get_config_geral():
        municipio = get_config_geral().municipio

    participantes = []
    ocorrencias = []
    comissao  = []
    membros = []
    licitantes = []

    for item in ParticipantePregao.objects.filter(pregao=pregao):
        nome = u'%s' % item.fornecedor
        participantes.append(nome.replace('&',"e"))
        me = u'Não'
        if item.me_epp:
            me = u'Sim'
        texto = u'%s - %s - %s - %s - %s' % (item.fornecedor.cnpj, nome.replace('&',"e"), me, item.nome_representante, item.cpf_representante)
        licitantes.append(texto)



    #     unidades.append(item.unidade)
    #     descricoes.append(item.material.nome)

    for item in HistoricoPregao.objects.filter(pregao=pregao):
        nome = u'%s'% item.obs
        ocorrencias.append(nome.replace('&',"e"))
    portaria = None
    if pregao.comissao:
        for item in MembroComissaoLicitacao.objects.filter(comissao=pregao.comissao).order_by('-funcao'):
            nome = u'%s'% (item.membro.nome)
            if not (item.funcao == MembroComissaoLicitacao.PREGOEIRO):
                comissao.append(nome.replace('&',"e"))

            texto = u'%s, %s,  %s ' % (nome, item.matricula, item.funcao)
            membros.append(texto)

        portaria = pregao.comissao.nome
        tipo = u'%s %s' % (pregao.tipo, pregao.criterio)



    eh_lote = pregao.criterio.id == CriterioPregao.LOTE


    tabela = {}
    total = {}

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(lance = list(), total = 0)

    for item in itens_pregao.order_by('item'):
        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['lance'].append(item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor

    total_geral = Decimal()
    resultado_pregao = u''
    resultado = collections.OrderedDict(sorted(tabela.items()))

    if pregao.criterio.nome == u'Por Item':
        nome_tipo = u'Itens'
    else:
        nome_tipo = u'Lotes'

    for result in resultado.items():
        if result[1]['total'] != 0:
            result[0]
            lista = []
            for item in result[1]['lance']:
                lista.append(item.item)



            resultado_pregao = resultado_pregao + u'%s, quanto aos %s %s, no valor total de R$ %s (%s), ' % (result[0], nome_tipo, lista, format_money(result[1]['total']), format_numero_extenso(result[1]['total']))
            total_geral = total_geral + result[1]['total']


    document = Document()
    imprimir_cabecalho(document, configuracao, logo, municipio)


    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'Ata de %s' % pregao).bold = True



    comissao = u', '.join(comissao)
    texto = u'''
    Às %s do dia %s, no(a) %s, realizou-se  a sessão pública para recebimento e abertura dos envelopes contendo as propostas de preços e as documentações de habilitação, apresentados em razão do certame licitatório na modalidade %s, cujo objeto é %s, conforme especificações mínimas constantes no Termo de Referência (Anexo I) deste Edital. As especificações técnicas dos serviços, objeto deste Pregão, estão contidas no Anexo I do Termo de Referência do Edital. Presentes o Pregoeiro, %s bem como, a Equipe de Apoio constituída pelos servidores: %s - Portaria: %s
    ''' % (pregao.hora_abertura, localize(pregao.data_abertura), pregao.local, pregao, pregao.objeto, pregao.responsavel, comissao, portaria)

    if pregao.comissao and pregao.comissao.data_designacao:
        texto += u'Data de Designação: %s.' % pregao.comissao.data_designacao.strftime('%d/%m/%Y')
    else:
        texto += '.'

    texto += u' O Pregoeiro iniciou a sessão informando os procedimentos da mesma.'
    #document.add_paragraph(texto)
    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(texto)

    if pregao.situacao == pregao.DESERTO:


        texto = u'''
            Aberta a sessão e atendidas todas as prescrições legais, o Sr. Pregoeiro registrou que nenhuma empresa compareceu ou remeteu os envelopes para participação do certame.
            Observa-se que a publicação do Aviso de Licitação pertinente ocorreu no em Diário Oficial Local e de Grande Circulação.
            Diante do ocorrido, o Sr. Pregoeiro decidiu considerar o presente certame como “DESERTO”, tendo em vista o não-comparecimento de empresas interessadas no certame, devendo tal resultado ser publicado no Diário Oficial.
            Após a referida publicação, os autos serão remetidos ao Órgão solicitante para ciência e adoção das providências cabíveis, tendo em vista o não comparecendo interessados na sessão.

            Nada mais havendo a tratar, deu o Sr. Pregoeiro por encerrado os trabalhos da reunião, com a lavratura da presente Ata, a qual depois de lida e aprovada, vai assinada pelo Sr. Pregoeiro e membros da Equipe de Apoio presentes à Sessão.

        '''
        p.alignment = 3
        p.add_run(texto)
    elif pregao.republicado and not pregao.tem_resultado():
        texto = u'''

        Aberta a sessão e atendidas todas as prescrições legais, o Sr. Pregoeiro registrou que nenhuma empresa compareceu ou remeteu os envelopes para participação do certame.
        Observa-se que a publicação do Aviso de Licitação pertinente ocorreu no dia ______________ em Diário Oficial, o Sr. Pregoeiro, após entrar em contato com o titular da pasta solicitante, determinou o reagendamento do presente certame.
        Destarte, estando aprovada a MINUTA do Edital anterior devidamente aprovada pela Assessoria Jurídica através do Parecer anexo, ficou avençado que a sessão inicial do %s (2ª CONVOCAÇÃO), será realizada no DIA %s, PELAS %s, no Setor de Licitações, localizado no térreo do prédio sede da %s, situado na %s. O extrato respectivo, objetivando a ampliação da divulgação, deverá ser publicado no DIÁRIO.
            Na referida data, os termos do EDITAL e seus anexos, com as condições e especificações, devem estar disponibilizados aos interessados no Setor de Licitações, localizado no %s, situado na %s, das __/___h às __/___h, de segunda a sexta-feira, em dias de expediente. O(s) Edital(is) poderão ser requeridos por meio do Portal da Transparência do Município de Guamaré http://www.guamare.rn.gov.br/licitacao/ e do e-mail cpl.guamare@gmail.com e através de solicitação contendo o timbrado da requerente e assinado por representante habilitado. Quaisquer esclarecimentos poderão ser prestados no referido setor ou através dos telefones: %s.
        Nada mais havendo a tratar, deu o Sr. Pregoeiro por encerrado os trabalhos da reunião, com a lavratura da presente Ata, a qual depois de lida e aprovada, vai assinada pelo Sr. Pregoeiro e membros da Equipe de Apoio presentes à Sessão.
        ''' % (pregao, localize(pregao.data_abertura), pregao.hora_abertura, configuracao.nome, pregao.local, configuracao.nome, pregao.local, configuracao.telefones)


        p.alignment = 3
        p.add_run(texto)
    else:

        texto = u'''
        Antes da abertura da sessão, realizou-se o credenciamento do (os) representante (es), feito a partir da apresentação da cédula de identidade ou documento equivalente, e procuração por instrumento público ou particular com firma reconhecida em cartório (documentos do outorgante, poderão ser conferidos na habilitação), atos esses documentados conforme listagem do (os) presente (es), que foram numeradas e juntadas aos autos às fls
        '''
        p.alignment = 3
        p.add_run(texto)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Empresa'
        hdr_cells[1].text = 'Representante'



        for item in ParticipantePregao.objects.filter(pregao=pregao):
            me = u'Não Compareceu'
            if item.nome_representante:
                me = item.nome_representante

            row_cells = table.add_row().cells
            row_cells[0].text = u'%s - %s' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
            row_cells[1].text = u'%s' % me
            texto = u'''

        Aberta a sessão, o Sr. Pregoeiro deu início aos trabalhos, fazendo comunicação ao (os) presente (es) sobre:

        a) Objetivos do Pregão;
        b) Ordenação dos trabalhos;
        c) Forma e ordem em que os licitantes pediriam a palavra;
        d) Vedação a intervenções fora da ordem definida;
        e) Forma como serão feitos os lances;
        f) Aviso sobre empresas coligadas e vedações do art. 90 da lei no 8.666/1993;
        g) Pedido para que não se retirasse (em) antes do término, em face à possibilidade de repregoar;
        h) As penalidades previstas no art. 70 da lei no 10.520/2002 com a correção de redação do texto da lei, conforme exposto no subitem 2.8.2.2;
        i)           Observou o pregoeiro que ele e a comissão de apoio têm interesse em cumprir a lei, respeitar os direitos dos licitantes e a lisura do certame; e
        j)           Após, foram esclarecidas as dúvidas do (os) licitante (es) e informado (os) o (os) nome (es) do (os) licitante (es) que estava (am) credenciado (os) para participar do certame, conforme listagem que foi exibida ao (os) presente (es).

        Dando continuidade passou-se ao procedimento de recebimento dos envelopes, que foram conferidos e apresentado ao (os) presente (es).
        Em seguida passou-se à abertura do (os) envelope (es) da (as) proposta (as) observando-se os seguintes passos:
        Abertura;
        Conferência do conteúdo; e
        Numeração.
        Na oportunidade foi esclarecido que a rubrica por um dos membros da equipe e pelo (os) licitante (es) que convidado (os) aceitar (em) rubricar, seria realizada no final.
        Dando continuidade procedeu-se à análise da (as) proposta (as), quando foi verificado se cada proposta atendia aos requisitos do edital, quanto ao objeto, prazo de entrega, garantia. Como resultado da análise foi (ram) classificada a (as) empresa (as) que atenderam todos os requisitos do edital e seus anexos.
        '''
        p = document.add_paragraph(texto)


        if pregao.tem_resultado():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(u'DOS LANCES').bold = True

            p = document.add_paragraph()
            p.alignment = 3


            if pregao.criterio.nome == u'Por Item':
                nome_tipo = u'item(ns)'
            else:
                nome_tipo = u'lote(s)'


            p.add_run(u'O Sr. Pregoeiro, com auxílio da Equipe de Pregão, deu início aos lances verbais, solicitando ao (os) representante (es) da (as) licitante (es) que ofertasse (em) seus lance (es) para o (os) %s em sequência, conforme mapa de lance (es) e classificação anexo.' % nome_tipo)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'DA HABILITAÇÃO').bold = True

        p = document.add_paragraph()
        p.alignment = 3

        p.add_run(u'Em seguida, foi analisada a aceitabilidade da(s) proposta(s) detentora(s) do menor preço, conforme previsto no edital. Posteriormente, foi analisada a documentação da referida empresa.')

        if pregao.tem_resultado():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(u'DO RESULTADO').bold = True

            p = document.add_paragraph()
            p.alignment = 3

            p.add_run(u'Diante da aceitabilidade da proposta e regularidade frente às exigências de habilitação contidas no instrumento convocatório, o Pregoeiro e equipe declararam como vencedora(s) do certame, a(s) empresa(s): ')


            p.add_run(resultado_pregao)

            p = document.add_paragraph()
            p.alignment = 3
            p.add_run(u'O valor global do certame, considerando o somatório dos itens licitados, será de R$ %s (%s), respeitado os valores máximos indicados, tendo em vista que o tipo da licitação é o de %s.' % (format_money(total_geral), format_numero_extenso(total_geral), tipo))

        if ocorrencias:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(u'DAS OCORRÊNCIAS DA SESSÃO PÚBLICA').bold = True


            for item in ocorrencias:
                p = document.add_paragraph()
                #p.alignment = 3
                p.add_run(item)


        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'DO ENCERRAMENTO').bold = True

        p = document.add_paragraph()
        p.alignment = 3
        if pregao.tem_resultado():
            p.add_run(u'O Pregoeiro, após encerramento desta fase, concedeu aos proponentes vistas ao processo e a todos os documentos. Franqueada a palavra, para observações, questionamentos e/ou interposição de recursos, caso alguém assim desejasse, como nenhum dos proponentes manifestou intenção de recorrer, pelo que renunciam, desde logo, em caráter irrevogável e irretratável, ao direito de interposição de recurso. Nada mais havendo a tratar, o Pregoeiro declarou encerrados os trabalhos, lavrando-se a presente Ata que vai assinada pelos presentes.')
        else:
            p.add_run(u'Nada mais havendo a tratar, o Pregoeiro declarou encerrados os trabalhos, lavrando-se a presente Ata que vai assinada pelos presentes')


    for item in membros:

        texto = item.split(',')
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto[0])
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'Matrícula: %s' % texto[1])
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto[2])



    for item in ParticipantePregao.objects.filter(pregao=pregao):

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        texto = u'%s (%s)' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
        p.add_run(texto)
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item.nome_representante:
            texto = item.nome_representante
            if item.cpf_representante:
                texto += u' (CPF: %s)' % item.cpf_representante
            p.add_run(texto)


    document.add_page_break()
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT, 'upload/pregao/atas/ata_sessao_%s.docx' % pregao.id)
    document.save(caminho_arquivo)



    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")

    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response


@login_required()
def ata_sessao_credenciamento(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)

    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)


    municipio = None
    if get_config_geral():
        municipio = get_config_geral().municipio

    participantes = []
    ocorrencias = []
    comissao  = []
    membros = []
    licitantes = []

    for item in ParticipantePregao.objects.filter(pregao=pregao):
        nome = u'%s' % item.fornecedor
        participantes.append(nome.replace('&',"e"))
        me = u'Não'
        if item.me_epp:
            me = u'Sim'
        texto = u'%s - %s - %s - %s - %s' % (item.fornecedor.cnpj, nome.replace('&',"e"), me, item.nome_representante, item.cpf_representante)
        licitantes.append(texto)



    #     unidades.append(item.unidade)
    #     descricoes.append(item.material.nome)

    for item in HistoricoPregao.objects.filter(pregao=pregao):
        nome = u'%s'% item.obs
        ocorrencias.append(nome.replace('&',"e"))
    portaria = None
    if pregao.comissao:
        for item in MembroComissaoLicitacao.objects.filter(comissao=pregao.comissao).order_by('-funcao'):
            nome = u'%s'% (item.membro.nome)
            if not (item.funcao == MembroComissaoLicitacao.PREGOEIRO):
                comissao.append(nome.replace('&',"e"))

            texto = u'%s, %s,  %s ' % (nome, item.matricula, item.funcao)
            membros.append(texto)

        portaria = pregao.comissao.nome
        tipo = u'%s %s' % (pregao.tipo, pregao.criterio)



    eh_lote = pregao.criterio.id == CriterioPregao.LOTE


    tabela = {}
    total = {}

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(lance = list(), total = 0)

    for item in itens_pregao.order_by('item'):
        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['lance'].append(item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor

    total_geral = Decimal()
    resultado_pregao = u''
    resultado = collections.OrderedDict(sorted(tabela.items()))

    if pregao.criterio.nome == u'Por Item':
        nome_tipo = u'Itens'
    else:
        nome_tipo = u'Lotes'

    for result in resultado.items():
        if result[1]['total'] != 0:
            result[0]
            lista = []
            for item in result[1]['lance']:
                lista.append(item.item)


            resultado_pregao = resultado_pregao + u'%s, quanto aos %s %s, no valor total de R$ %s (%s), ' % (result[0], nome_tipo, lista, format_money(result[1]['total']), format_numero_extenso(result[1]['total']))
            total_geral = total_geral + result[1]['total']

    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()

    imprimir_cabecalho(document, configuracao, logo, municipio)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'%s' % pregao).bold = True



    comissao = u', '.join(comissao)
    texto = u'''
    Às %s do dia %s, no(a) %s, realizou-se  a sessão pública para recebimento e abertura dos envelopes contendo as propostas de preços e as documentações de habilitação, apresentados em razão do certame licitatório na modalidade %s, cujo objeto é %s, conforme especificações mínimas constantes no Termo de Referência (Anexo I) deste Edital..  As especificações técnicas dos serviços, objeto deste Pregão, estão contidas no Anexo I do Termo de Referência do Edital. Presentes o Pregoeiro, %s bem como, a Equipe de Apoio constituída pelos servidores: %s - Portaria: %s. O Pregoeiro iniciou a sessão informando os procedimentos da mesma.
    Aberta a Sessão e atendidas todas as prescrições legais, o Sr. Presidente da CPL/PMG registrou que, de acordo com o prazo estabelecido, foi entregue o Envelope da (s) seguinte (s) licitante (s):

    ''' % (pregao.hora_abertura, localize(pregao.data_abertura), pregao.local, pregao, pregao.objeto, pregao.responsavel, comissao, portaria)

    #document.add_paragraph(texto)
    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(texto)


    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Empresa'
    hdr_cells[1].text = 'Representante'



    for item in ParticipantePregao.objects.filter(pregao=pregao):
        me = u'Não Compareceu'
        if item.nome_representante:
            me = item.nome_representante

        row_cells = table.add_row().cells
        row_cells[0].text = u'%s - %s' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
        row_cells[1].text = u'%s' % me

    texto = u'''

    Verificando que a publicação do Procedimento ocorreu na forma e no prazo previsto na legislação aplicável, inclusive com a disponibilização do Edital no Setor de Licitações da PMG, Portal da Transparência e Quadro de Avisos da PMG, deu o Sr. Presidente da CPL/PMG continuidade do procedimento.

    Em ato contínuo, foi (ram) aberto (s) o (s) Envelope (s) (Documentação de Credenciamento) da (s) referida (s) empresa (s), sendo os documentos devidamente verificados e rubricados pelo (s) representante (s) da (s) empresa (s) presente à sessão, bem como pelos Membros da CPL/PMG.
    '''
    p = document.add_paragraph(texto)



    if pregao.tem_resultado():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'DO RESULTADO').bold = True

        p = document.add_paragraph()
        p.alignment = 3

        p.add_run(u'Diante da regularidade frente às exigências de habilitação contidas no instrumento convocatório, o Presidente da CPL/PMG e Membros da Equipe de Apoio, declararam como APTA (s) do procedimento, a (s) empresa (s):')

        p.add_run(resultado_pregao)

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Empresa'
        hdr_cells[1].text = 'Representante'



        for item in ParticipantePregao.objects.filter(pregao=pregao, desclassificado=False, excluido_dos_itens=False):
            me = u'Não Compareceu'
            if item.nome_representante:
                me = item.nome_representante

            row_cells = table.add_row().cells
            row_cells[0].text = u'%s - %s' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
            row_cells[1].text = u'%s' % me

        texto = u'''
            Após o resultado, o Sr. Presidente da CPL/PMG concedeu a palavra a estes para os eventuais registros quanto a documentação de habilitação apresentada, tendo estes declarado que não há nada a registrar.

            Em seguida, o Presidente da CPL/PMG comunicou aos presentes que encerraria a Sessão, informando que o resultado será publicado no Diário Oficial dos Municípios do Estado do Rio Grande do Norte – FEMURN, e após, o objeto do certame será adjudicado às empresas vencedoras.

            Nada mais havendo a tratar, deu o Sr. Presidente por encerrado os trabalhos da reunião às 15h00min (quinze horas), com a lavratura da presente Ata, a qual depois de lida e aprovada, vai assinada pelo Presidente, Membros da CPL/PMG presentes à Sessão.
        '''
        p = document.add_paragraph(texto)



    for item in membros:

        texto = item.split(',')
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto[0])
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'Matrícula: %s' % texto[1])
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto[2])



    for item in ParticipantePregao.objects.filter(pregao=pregao, desclassificado=False, excluido_dos_itens=False):

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        texto = u'%s (%s)' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
        p.add_run(texto)
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item.nome_representante:
            texto = item.nome_representante
            if item.cpf_representante:
                texto += u' (CPF: %s)' % item.cpf_representante
            p.add_run(texto)


    document.add_page_break()
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT, 'upload/pregao/atas/ata_sessao_%s.docx' % pregao.id)
    document.save(caminho_arquivo)



    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")

    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response


@login_required()
def ata_sessao_outras_modalidades(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)

    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)


    municipio = None
    if get_config_geral():
        municipio = get_config_geral().municipio

    participantes = []
    ocorrencias = []
    comissao  = []
    membros = []
    licitantes = []

    for item in ParticipantePregao.objects.filter(pregao=pregao):
        nome = u'%s' % item.fornecedor
        participantes.append(nome.replace('&',"e"))
        me = u'Não'
        if item.me_epp:
            me = u'Sim'
        texto = u'%s - %s - %s - %s - %s' % (item.fornecedor.cnpj, nome.replace('&',"e"), me, item.nome_representante, item.cpf_representante)
        licitantes.append(texto)



    #     unidades.append(item.unidade)
    #     descricoes.append(item.material.nome)

    for item in HistoricoPregao.objects.filter(pregao=pregao):
        nome = u'%s'% item.obs
        ocorrencias.append(nome.replace('&',"e"))
    portaria = None
    if pregao.comissao:
        for item in MembroComissaoLicitacao.objects.filter(comissao=pregao.comissao).order_by('-funcao'):
            nome = u'%s'% (item.membro.nome)
            if not (item.funcao == MembroComissaoLicitacao.PREGOEIRO):
                comissao.append(nome.replace('&',"e"))

            texto = u'%s, %s,  %s ' % (nome, item.matricula, item.funcao)
            membros.append(texto)

        portaria = pregao.comissao.nome
        tipo = u'%s %s' % (pregao.tipo, pregao.criterio)



    eh_lote = pregao.criterio.id == CriterioPregao.LOTE


    tabela = {}
    total = {}

    if eh_lote:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=True, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    else:
        itens_pregao = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False, situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO])
    resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO)
    chaves =  resultado.values('participante__fornecedor').order_by('participante__fornecedor').distinct('participante__fornecedor')
    for num in chaves:
        fornecedor = get_object_or_404(Fornecedor, pk=num['participante__fornecedor'])
        chave = u'%s' % fornecedor
        tabela[chave] = dict(lance = list(), total = 0)

    for item in itens_pregao.order_by('item'):
        if item.get_vencedor():
            chave = u'%s' % item.get_vencedor().participante.fornecedor
            tabela[chave]['lance'].append(item)
            valor = tabela[chave]['total']
            valor = valor + item.get_total_lance_ganhador()
            tabela[chave]['total'] = valor

    total_geral = Decimal()
    resultado_pregao = u''
    resultado = collections.OrderedDict(sorted(tabela.items()))

    if pregao.criterio.nome == u'Por Item':
        nome_tipo = u'Itens'
    else:
        nome_tipo = u'Lotes'

    for result in resultado.items():
        if result[1]['total'] != 0:
            result[0]
            lista = []
            for item in result[1]['lance']:
                lista.append(item.item)



            resultado_pregao = resultado_pregao + u'%s, quanto aos %s %s, no valor total de R$ %s (%s), ' % (result[0], nome_tipo, lista, format_money(result[1]['total']), format_numero_extenso(result[1]['total']))
            total_geral = total_geral + result[1]['total']


    document = Document()
    imprimir_cabecalho(document, configuracao, logo, municipio)


    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'Ata de %s' % pregao).bold = True



    comissao = u', '.join(comissao)
    texto = u'''
    Às %s do dia %s, no(a) %s, realizou-se  a sessão pública para recebimento e abertura dos envelopes contendo as documentações de habilitação e propostas de preços, apresentados em razão do certame licitatório na modalidade %s, cujo objeto é %s, conforme especificações mínimas constantes no Projeto Base/Termo de Referência. anexo à este Edital. Presentes o Presidente da CPL, %s bem como os Membros da CPL, constituída pelos servidores: %s - Portaria: %s
    ''' % (pregao.hora_abertura, localize(pregao.data_abertura), pregao.local, pregao, pregao.objeto, pregao.responsavel, comissao, portaria)

    if pregao.comissao and pregao.comissao.data_designacao:
        texto += u'Data de Designação: %s.' % pregao.comissao.data_designacao.strftime('%d/%m/%Y')
    else:
        texto += '.'

    texto += u' O Presidente e os Membros da CPL iniciaram a sessão informando os procedimentos da mesma. '
    #document.add_paragraph(texto)
    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(texto)



    texto = u'''
    Antes da abertura da sessão, realizou-se o credenciamento do (os) representante (es), feito a partir da apresentação da cédula de identidade ou documento equivalente, e procuração por instrumento público ou particular com firma reconhecida em cartório (documentos do outorgante, poderão ser conferidos na habilitação), atos esses documentados conforme listagem do (os) presente (es), que foram numeradas e juntadas aos autos às fls.
    '''
    p.alignment = 3
    p.add_run(texto)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Empresa'
    hdr_cells[1].text = 'Representante'



    for item in ParticipantePregao.objects.filter(pregao=pregao):
        me = u'Não Compareceu'
        if item.nome_representante:
            me = item.nome_representante

        row_cells = table.add_row().cells
        row_cells[0].text = u'%s - %s' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
        row_cells[1].text = u'%s' % me
        texto = u'''

    Aberta a sessão, o Sr. Presidente e Membros da CPL, deram início aos trabalhos, fazendo comunicação ao (os) presente (es) sobre:
        a) Objetivos da Licitação;
        b) Ordenação dos trabalhos;
        c) Forma e ordem em que os licitantes pediriam a palavra;
        d) Vedação a intervenções fora da ordem definida;
        e) Aviso sobre empresas coligadas e vedações do art. 90 da lei no 8.666/1993;
        f) Pedido para que não se retirasse (em) antes do término, em face à possibilidade de repregoar;
        g) Observou o Presidente, que ele e a Comissão de Licitação têm interesse em cumprir a lei, respeitar os direitos dos licitantes e a lisura do certame; e
        h)  Após, foram esclarecidas as dúvidas do (os) licitante (es) e informado (os) o (os) nome (es) do (os) licitante (es) que estava (am) credenciado (os) para participar do certame, conforme listagem que foi exibida ao (os) presente (es).
        Dando continuidade passou-se ao procedimento de recebimento dos envelopes, que foram conferidos e apresentado ao (os) presente (es).
        Em seguida passou-se à abertura do (os) envelope (es) de Habilitação, observando-se os seguintes passos:
                Abertura;
                Conferência do conteúdo; e
                Numeração.
        Na oportunidade foi esclarecido que a rubrica por um dos membros da equipe e pelo (os) licitante (es) que convidado (os) aceitar (em) rubricar, seria realizada no final.
        Dando continuidade procedeu-se à análise da (as) Documentações, quando foi verificado se atendia (am) ao (os) requisitos do edital.
    '''
    p = document.add_paragraph(texto)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'DA HABILITAÇÃO').bold = True
    texto += u'''
    Diante da regularidade frente às exigências de habilitação contidas no instrumento convocatório, o Presidente e Membros da CPL, passaram a analisar a Aceitabilidade da(s) proposta(s) detentora(s) do menor preço, conforme previsto no edital.
    Em seguida, foi analisada a aceitabilidade da(s) proposta(s) detentora(s) do menor preço, conforme previsto no edital. Posteriormente, foi analisada a documentação da referida empresa.
    '''
    #document.add_paragraph(texto)
    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(texto)


    if pregao.tem_resultado():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'DO RESULTADO').bold = True

        p = document.add_paragraph()
        p.alignment = 3

        p.add_run(u'Diante da aceitabilidade da proposta e regularidade frente às exigências de habilitação contidas no instrumento convocatório, o Presidente e os Membros da CPL declararam como vencedora(s) do certame, a(s) empresa(s): ')


        p.add_run(resultado_pregao)

        p = document.add_paragraph()
        p.alignment = 3
        p.add_run(u'O valor global do certame, considerando o somatório dos itens licitados, será de R$ %s (%s), respeitado os valores máximos indicados, tendo em vista que o tipo da licitação é o de %s.' % (format_money(total_geral), format_numero_extenso(total_geral), tipo))

    if ocorrencias:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'DAS OCORRÊNCIAS DA SESSÃO PÚBLICA').bold = True


        for item in ocorrencias:
            p = document.add_paragraph()
            #p.alignment = 3
            p.add_run(item)


    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'DO ENCERRAMENTO').bold = True

    p = document.add_paragraph()
    p.alignment = 3
    p.add_run(u'O Presidente e Membros da CPL, após encerramento desta fase, concedeu aos proponentes vistas ao processo e a todos os documentos. Franqueada a palavra, para observações, questionamentos e/ou interposição de recursos, caso alguém assim desejasse, como nenhum dos proponentes manifestou intenção de recorrer, pelo que renunciam, desde logo, em caráter irrevogável e irretratável, ao direito de interposição de recurso. Nada mais havendo a tratar, o Presidente e os Membros da CPL declararam encerrados os trabalhos, lavrando-se a presente Ata que vai assinada pelos presentes.')


    for item in membros:

        texto = item.split(',')
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto[0])
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(u'Matrícula: %s' % texto[1])
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(texto[2])



    for item in ParticipantePregao.objects.filter(pregao=pregao):

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        texto = u'%s (%s)' % (item.fornecedor.razao_social, item.fornecedor.cnpj)
        p.add_run(texto)
        p = document.add_paragraph()
        #p.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item.nome_representante:
            texto = item.nome_representante
            if item.cpf_representante:
                texto += u' (CPF: %s)' % item.cpf_representante
            p.add_run(texto)


    document.add_page_break()
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT, 'upload/pregao/atas/ata_sessao_%s.docx' % pregao.id)
    document.save(caminho_arquivo)



    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")

    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response



@login_required()
def adicionar_membro_comissao(request, comissao_id):
    comissao = get_object_or_404(ComissaoLicitacao, pk=comissao_id)
    title=u'Adicionar Membro da Comissão'
    form = MembroComissaoLicitacaoForm(request.POST or None, comissao=comissao)
    if form.is_valid():
        o = form.save(False)
        o.comissao = comissao
        o.save()
        messages.success(request, u'Membro cadastrado com sucesso.')
        return HttpResponseRedirect(u'/admin/base/comissaolicitacao/')
    return render(request, 'comissao_licitacao.html', locals(), RequestContext(request))

@login_required()
def remover_membro_comissao(request, comissao_id):
    comissao = get_object_or_404(ComissaoLicitacao, pk=comissao_id)
    title=u'Remover Membro da Comissão'
    form = RemoverMembroComissaoLicitacaoForm(request.POST or None, comissao=comissao)
    if form.is_valid():
        MembroComissaoLicitacao.objects.filter(comissao=comissao, membro=form.cleaned_data.get('membro')).delete()
        messages.success(request, u'Membro removido com sucesso.')
        return HttpResponseRedirect(u'/admin/base/comissaolicitacao/')
    return render(request, 'comissao_licitacao.html', locals(), RequestContext(request))

@login_required()
def editar_membro_comissao(request, membro_id):
    membro = get_object_or_404(MembroComissaoLicitacao, pk=membro_id)
    title=u'Editar Membro da Comissão'
    form = MembroComissaoLicitacaoForm(request.POST or None, instance=membro)
    if form.is_valid():
        form.save()

        messages.success(request, u'Membro editado com sucesso.')
        return HttpResponseRedirect(u'/admin/base/comissaolicitacao/')
    return render(request, 'comissao_licitacao.html', locals(), RequestContext(request))

@login_required()
def editar_valor_final(request, item_id, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
        title=u'Editar Valor Unitário Final - %s' % item
        valor = item.get_valor_item_lote() / item.quantidade
        participante_id = request.GET.get("participante")
        form = ValorFinalItemLoteForm(request.POST or None, initial=dict(valor=valor), participante_id=participante_id)
        if form.is_valid():
            lote = ItemLote.objects.filter(item=item)[0].lote
            itens_do_lote = ItemLote.objects.filter(lote=lote).values_list('item', flat=True)

            if form.cleaned_data.get('valor') > item.get_valor_total_proposto() or form.cleaned_data.get('valor') > item.valor_medio:
                messages.error(request, u'O valor não pode ser maior do que o valor unitário proposto: %s nem do que o valor máximo do item: %s.' % (item.get_valor_total_proposto(), item.valor_medio))
                return HttpResponseRedirect(u'/base/editar_valor_final/%s/%s/' % (item.id, pregao.id))

            valor = PropostaItemPregao.objects.filter(item__in=itens_do_lote, participante=lote.get_empresa_vencedora()).exclude(item=item).aggregate(total=Sum('valor_item_lote'))['total'] or 0
            if (valor + form.cleaned_data.get('valor')*item.quantidade) > lote.get_total_lance_ganhador():
                messages.error(request, u'O valor informado faz o valor total dos itens do lote ultrapassar o valor do lance ganhador: %s.' % lote.get_total_lance_ganhador())
                return HttpResponseRedirect(u'/base/editar_valor_final/%s/%s/' % (item_id, pregao.id))

            valor_final = form.cleaned_data.get('valor') * item.quantidade
            PropostaItemPregao.objects.filter(item=item, participante=lote.get_empresa_vencedora()).update(valor_item_lote=valor_final)
            messages.success(request, u'Valor editado com sucesso.')
            if form.cleaned_data.get('participante_id'):
                return HttpResponseRedirect(u'/base/pregao/%s/?participante=%s#classificacao' % (pregao.id, form.cleaned_data.get('participante_id')))
            else:
                return HttpResponseRedirect(u'/base/pregao/%s/#classificacao' % pregao.id)
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def gerar_resultado_item_pregao(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        item.gerar_resultado(apaga=True)
        messages.success(request, u'Resultado do item gerado com sucesso.')
        return HttpResponseRedirect(u'/base/lances_item/%s/' % item_id)
    else:
        raise PermissionDenied



@login_required()
def aderir_arp(request):
    title=u'Aderir à ARP'
    form = AderirARPForm(request.POST or None)
    if form.is_valid():

        nova_solicitacao = SolicitacaoLicitacao()
        nova_solicitacao.num_memorando = form.cleaned_data.get('num_memorando')
        nova_solicitacao.objeto = form.cleaned_data.get('objeto')
        nova_solicitacao.objetivo = form.cleaned_data.get('objetivo')
        nova_solicitacao.justificativa = form.cleaned_data.get('justificativa')
        nova_solicitacao.tipo_aquisicao = SolicitacaoLicitacao.TIPO_AQUISICAO_ADESAO_ARP
        nova_solicitacao.tipo = SolicitacaoLicitacao.ADESAO_ARP
        nova_solicitacao.setor_origem = request.user.pessoafisica.setor
        nova_solicitacao.setor_atual = request.user.pessoafisica.setor
        nova_solicitacao.data_cadastro = datetime.datetime.now()
        nova_solicitacao.cadastrado_por = request.user
        nova_solicitacao.save()

        o = form.save(False)
        o.adesao = True
        o.solicitacao = nova_solicitacao
        o.secretaria = request.user.pessoafisica.setor.secretaria
        o.save()
        messages.success(request, u'Ata cadastrada com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % o.id)
    return render(request, 'aderir_arp.html', locals(), RequestContext(request))


@login_required()
def adicionar_item_adesao_arp(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    if ata.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        title = u'Adicionar Item - %s' % ata
        form = AdicionarItemAtaForm(request.POST or None)
        if form.is_valid():
            o = form.save(False)
            o.ata = ata
            o.fornecedor = ata.fornecedor_adesao_arp
            o.save()
            total = 0
            for item in ItemAtaRegistroPreco.objects.filter(ata=ata):
                total = total + (item.valor * item.quantidade)
            ata.valor = total
            ata.save()
            messages.success(request, u'Item da ata cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % ata.id)
        return render(request, 'adicionar_item_adesao_arp.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def carregar_planilha_itens_adesao_arp(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    title = u'Carregar Itens'
    if ata.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        form = ImportarItensForm(request.POST or None, request.FILES or None)
        if form.is_valid():
            arquivo_up = form.cleaned_data.get('arquivo')
            if arquivo_up:
                sheet = None
                try:
                    with transaction.atomic():
                        workbook = xlrd.open_workbook(file_contents=arquivo_up.read())
                        sheet = workbook.sheet_by_index(0)

                except XLRDError:
                    raise Exception(u'Não foi possível processar a planilha. Verfique se o formato do arquivo é .xls ou .xlsx.')

                for row in range(3, sheet.nrows):

                    #codigo = unicode(sheet.cell_value(row, 0)).strip()
                    especificacao = unicode(sheet.cell_value(row, 0)).strip()
                    marca = unicode(sheet.cell_value(row, 1)).strip()
                    unidade = unicode(sheet.cell_value(row, 2)).strip()
                    qtd = unicode(sheet.cell_value(row, 3)).strip()
                    valor = unicode(sheet.cell_value(row, 4)).strip()
                    if row == 3:
                        if especificacao != u'MATERIAL' or unidade != u'UNIDADE' or qtd != u'QUANTIDADE':
                            raise Exception(u'Não foi possível processar a planilha. As colunas devem ter Especificação, Unidade e Quantidade.')
                    else:
                        if especificacao and unidade and qtd:
                            if TipoUnidade.objects.filter(nome=unidade).exists():
                                un = TipoUnidade.objects.filter(nome=unidade)[0]
                            else:
                                un = TipoUnidade.objects.get_or_create(nome=unidade)[0]


                            if MaterialConsumo.objects.filter(nome=especificacao).exists():
                                material = MaterialConsumo.objects.filter(nome=especificacao)[0]
                            else:
                                material = MaterialConsumo()
                                if MaterialConsumo.objects.exists():
                                    id = MaterialConsumo.objects.latest('id')
                                    material.id = id.pk+1
                                    material.codigo = id.pk+1
                                else:
                                    material.id = 1
                                    material.codigo = 1
                                material.nome = especificacao
                                material.save()

                            novo_item = ItemAtaRegistroPreco()
                            novo_item.material = material
                            novo_item.marca = marca
                            novo_item.unidade = un
                            novo_item.quantidade = qtd
                            novo_item.ordem = ata.get_ordem()
                            try:
                                with transaction.atomic():
                                    Decimal(valor)
                            except:
                                messages.error(request, u'o valor %s é inválido.' % (valor))
                                return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % ata.id)
                            novo_item.valor = valor
                            novo_item.ata = ata
                            novo_item.fornecedor = ata.fornecedor_adesao_arp
                            novo_item.ata = ata
                            novo_item.fornecedor = ata.fornecedor_adesao_arp
                            novo_item.save()


                total = 0
                for item in ItemAtaRegistroPreco.objects.filter(ata=ata):
                    total = total + (item.valor * item.quantidade)
                ata.valor = total
                ata.save()

            messages.success(request, u'Itens cadastrados com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % ata.id)
        return render(request, 'carregar_planilha_itens_adesao_arp.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def cadastrar_material_arp(request, ata_id):
    title = u'Cadastrar Material'
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    if ata.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        form = MaterialConsumoForm(request.POST or None)
        if form.is_valid():
            form.save()
            messages.success(request, u'Item da ata cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/adicionar_item_adesao_arp/%s/' % ata_id)
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def editar_item(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    title=u'Editar Item'
    if request.user.has_perm('base.pode_cadastrar_solicitacao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        pedido = ItemQuantidadeSecretaria.objects.filter(item=item, secretaria=request.user.pessoafisica.setor.secretaria)

        if pedido.exists():
            form = EditarItemSolicitacaoLicitacaoForm(request.POST or None, instance=pedido[0], unidade=item.unidade)
            valor_original = pedido[0].quantidade
            if form.is_valid():
                o = form.save()
                item.unidade = form.cleaned_data.get('unidade')
                item.quantidade = item.quantidade - valor_original + form.cleaned_data.get('quantidade')

                item.save()
                messages.success(request, u'Item editado com sucesso.')
                return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % item.solicitacao.id)

            return render(request, 'cadastrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def solicitar_pedidos_novamente(request, solicitacao_id, secretaria_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_solicitacao') and solicitacao.recebida_setor(request.user.pessoafisica.setor):
        if solicitacao.setor_origem.secretaria.id == int(secretaria_id):
            messages.error(request, u'Edite as quantidades na própria solicitação.')
            return HttpResponseRedirect(u'/base/avaliar_pedidos/%s/' % solicitacao_id)

        else:
            for item in ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao_id, secretaria=secretaria_id):
                if item.aprovado:
                    item.item.quantidade -= item.quantidade
                    item.item.save()

            ItemQuantidadeSecretaria.objects.filter(solicitacao=solicitacao_id, secretaria=secretaria_id).delete()
            messages.success(request, u'Os pedidos serão solicitados novamente às secretarias.')
            return HttpResponseRedirect(u'/base/avaliar_pedidos/%s/' % solicitacao_id)
    else:
        raise PermissionDenied


@login_required()
def revogar_pregao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        title = u'Revogar pregão - %s' % pregao
        form = RevogarPregaoForm(request.POST or None, instance=pregao)
        if form.is_valid():
            o = form.save(False)
            o.situacao = Pregao.REVOGADO
            o.save()
            historico = HistoricoPregao()
            historico.pregao = pregao
            historico.data = datetime.datetime.now()
            historico.obs = u'Pregão revogado.'
            historico.save()

            messages.success(request, u'Pregão revogado com sucesso.')
            return HttpResponseRedirect(u'/base/pregao/%s/' % pregao.id)
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def imprimir_fornecedor(request, fornecedor_id):
    fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
    configuracao = get_config()
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/ordens_compra/%s.pdf' % fornecedor_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/ordens_compra')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/ordens_compra'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    data = {'fornecedor': fornecedor, 'configuracao': configuracao, 'logo': logo, 'data_emissao': data_emissao}

    template = get_template('imprimir_fornecedor.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')

@login_required()
def excluir_solicitacao_pedido(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_solicitacao') and solicitacao.recebida_setor(request.user.pessoafisica.setor):
        PedidoAtaRegistroPreco.objects.filter(solicitacao=solicitacao).delete()
        PedidoContrato.objects.filter(solicitacao=solicitacao).delete()
        ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao).delete()
        solicitacao.delete()

        messages.success(request, u'Solicitação excluída com sucesso.')
        return HttpResponseRedirect(u'/base/ver_solicitacoes/')
    else:
        raise PermissionDenied


@login_required()
def ver_variaveis_configuracao(request):
    if request.user.is_superuser:
        title = u'Listar Variáveis de Configuração'
        config = None
        if Configuracao.objects.exists():
            config = Configuracao.objects.latest('id')
        return render(request, 'ver_variaveis_configuracao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def cadastrar_variaveis_configuracao(request):
    if request.user.is_superuser:
        title = u'Cadastar/Editar Variáveis de Configuração'
        if not Configuracao.objects.exists():
            config = Configuracao()
        else:
            config = Configuracao.objects.latest('id')

        form = ConfiguracaoForm(request.POST or None, request.FILES or None, instance=config)
        if form.is_valid():
            form.save()
            messages.success(request, u'Variáveis de configuração com sucesso.')
            return HttpResponseRedirect(u'/base/ver_variaveis_configuracao/')
        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def ver_pesquisas(request, solicitacao_id):
    title = u'Ver Pesquisas'
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and (solicitacao.prazo_aberto or not solicitacao.eh_dispensa() or not solicitacao.tem_ordem_compra()):
        pesquisas = PesquisaMercadologica.objects.filter(solicitacao=solicitacao)
        return render(request, 'ver_pesquisas.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def editar_item_pedido(request, pedido_id, tipo):
    title = u'Editar Item do Pedido'
    if tipo == u'1':
        pedido = get_object_or_404(PedidoContrato, pk=pedido_id)
        if pedido.solicitacao.recebida_setor(request.user.pessoafisica.setor):
            form = EditarPedidoContratoForm(request.POST or None, instance=pedido)
            if form.is_valid():
                form.save()
                messages.success(request, u'Item editado com sucesso.')
                return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' %  pedido.solicitacao.id)
        else:
            raise PermissionDenied

    else:
        pedido = get_object_or_404(PedidoAtaRegistroPreco, pk=pedido_id)
        if pedido.solicitacao.recebida_setor(request.user.pessoafisica.setor):
            form = EditarPedidoARPForm(request.POST or None, instance=pedido)
            if form.is_valid():
                form.save()
                messages.success(request, u'Item editado com sucesso.')
                return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' %  pedido.solicitacao.id)
        else:
            raise PermissionDenied
    return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))



@login_required()
def apagar_item_pedido(request, pedido_id, tipo):
    if tipo == u'1':
        pedido = get_object_or_404(PedidoContrato, pk=pedido_id)
        solicitacao = pedido.solicitacao
        if solicitacao.recebida_setor(request.user.pessoafisica.setor):
            pedido.delete()
            messages.success(request, u'Item excluído com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' %  solicitacao.id)
        else:
            raise PermissionDenied

    else:
        pedido = get_object_or_404(PedidoAtaRegistroPreco, pk=pedido_id)
        solicitacao = pedido.solicitacao
        if request.user.has_perm('base.pode_cadastrar_solicitacao') and solicitacao.recebida_setor(request.user.pessoafisica.setor):
            pedido.delete()
            messages.success(request, u'Item excluído com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' %  solicitacao.id)
        else:
            raise PermissionDenied


@login_required()
def gerenciar_visitantes(request, pregao_id):
    title = u'Lista de Visitantes do Pregão'
    pregao = get_object_or_404(Pregao, pk=pregao_id)
    visitantes = VisitantePregao.objects.filter(pregao=pregao)

    return render(request, 'gerenciar_visitantes.html', locals(), RequestContext(request))


@login_required()
def editar_visitante(request, visitante_id):
    title = u'Editar Visitante'
    visitante = get_object_or_404(VisitantePregao, pk=visitante_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and visitante.pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):

        form = VisitantePregaoForm(request.POST or None, instance=visitante)
        if form.is_valid():
            form.save()

            messages.success(request, u'Visitante editado com sucesso.')
            return HttpResponseRedirect(u'/base/gerenciar_visitantes/%s/' %  visitante.pregao.id)
        return render(request, 'cadastra_visitante_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def excluir_visitante(request, visitante_id):
    visitante = get_object_or_404(VisitantePregao, pk=visitante_id)
    pregao = visitante.pregao
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        visitante.delete()
        messages.success(request, u'Visitante excluído com sucesso.')
        return HttpResponseRedirect(u'/base/gerenciar_visitantes/%s/' % pregao.id)
    else:
        raise PermissionDenied

def localizar_processo(request):
    title = u'Localizar Processo'
    form = LocalizarProcessoForm(request.POST or None)
    if form.is_valid():
        if Processo.objects.filter(numero__icontains=form.cleaned_data.get('numero')).exists():
            processo = Processo.objects.filter(numero__icontains=form.cleaned_data.get('numero'))[0]
            solicitacao = SolicitacaoLicitacao.objects.filter(processo=processo)
            if solicitacao.exists():
                solicitacao = solicitacao[0]
                movimentos = MovimentoSolicitacao.objects.filter(solicitacao=solicitacao).order_by('-data_envio')


    return render(request, 'localizar_processo.html', locals(), RequestContext(request))


@login_required()
def ver_relatorios_gerenciais_licitacao(request):
    title=u'Relatórios Gerenciais das Licitações'

    eh_ordenador_despesa = False
    if get_config():
        eh_ordenador_despesa = request.user.pessoafisica == get_config().ordenador_despesa

    form = RelatoriosGerenciaisForm(request.POST or None)

    if form.is_valid():
        pregoes = Pregao.objects.all().order_by('num_pregao')
        relatorio =  form.cleaned_data.get('relatorio')
        modalidade = form.cleaned_data.get('modalidade')
        situacao = form.cleaned_data.get('situacao')
        visualizar = form.cleaned_data.get('visualizar')
        secretaria = form.cleaned_data.get('secretaria')
        ano = form.cleaned_data.get('ano')

        total = 0
        if ano:
            pregoes = pregoes.filter(data_abertura__year=ano)

        if situacao:
            pregoes = pregoes.filter(situacao=situacao)

        if modalidade:
            pregoes = pregoes.filter(modalidade=modalidade)

        if secretaria:
            pregoes = pregoes.filter(solicitacao__setor_origem__secretaria=secretaria)

        if relatorio == u'Relatório de Economia':
            total_previsto = 0
            total_final = 0
            total_desconto = 0
            total_economizado = 0
            for pregao in pregoes:
                total_previsto += pregao.get_total_previsto()
                total_final += pregao.get_total_final()
                total_economizado += pregao.get_total_economizado()

            if total_previsto:
                reducao = total_final / total_previsto
                ajuste= 1-reducao
                total_desconto = u'%s%%' % (ajuste.quantize(TWOPLACES) * 100)
        else:
            for pregao in pregoes:
                total += pregao.get_valor_total()

        if visualizar == u'2':
            destino_arquivo = u'upload/resultados/relatorio_gerencial_%s.pdf' %  request.user.pessoafisica.id
            if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
                os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
            caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
            data_emissao = datetime.date.today()

            configuracao = get_config(request.user.pessoafisica.setor.secretaria)
            logo = None
            if configuracao.logo:
                logo = os.path.join(settings.MEDIA_ROOT, configuracao.logo.name)


            if relatorio == u'Relatório de Economia':
                data = {'pregoes': pregoes, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'total_previsto': total_previsto, 'total_final': total_final, 'total_desconto': total_desconto,  'total_economizado': total_economizado}
                template = get_template('relatorio_gerencial_economia.html')
            else:
                data = {'pregoes': pregoes, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'total': total }
                template = get_template('relatorio_gerencial_situacao.html')


            html  = template.render(Context(data))

            pdf_file = open(caminho_arquivo, "w+b")
            pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                    encoding='utf-8')
            pdf_file.close()
            file = open(caminho_arquivo, "r")
            pdf = file.read()
            file.close()
            return HttpResponse(pdf, 'application/pdf')


    return render(request, 'ver_relatorios_gerenciais_licitacao.html', locals(), RequestContext(request))



@login_required()
def ver_relatorios_gerenciais_contratos(request):
    title=u'Relatórios Gerenciais dos Contratos'
    eh_gerente = request.user.groups.filter(name='Gerente')
    exibe_fornecedor = True
    if eh_gerente:

        form = RelatoriosGerenciaisContratosForm(request.POST or None, fornecedor=exibe_fornecedor)

        if form.is_valid():
            contratos = Contrato.objects.all().order_by('numero')
            relatorio =  form.cleaned_data.get('relatorio')
            situacao = form.cleaned_data.get('situacao')
            visualizar = form.cleaned_data.get('visualizar')
            secretaria = form.cleaned_data.get('secretaria')
            fornecedor = form.cleaned_data.get('fornecedor')
            ano = form.cleaned_data.get('ano')
            hoje = datetime.date.today()
            total = 0
            if ano:
                contratos = contratos.filter(data_inicio__year=ano)

            if situacao:
                descricao_situacao = u'Todos'
                if situacao == u'2':
                    contratos = contratos.filter(concluido=False, suspenso=False, cancelado=False, data_inicio__lte=hoje, data_fim__gte=hoje)
                    descricao_situacao = u'Vigentes'
                elif situacao == u'3':
                    contratos = contratos.filter(Q(concluido=True), Q(suspenso=True), Q(cancelado=True))
                    descricao_situacao = u'Concluídos'

            if secretaria:
                contratos = contratos.filter(solicitacao__setor_origem__secretaria=secretaria)

            if fornecedor:
                itens = ItemContrato.objects.filter(Q(fornecedor=form.cleaned_data.get('fornecedor')) | Q(participante__fornecedor=form.cleaned_data.get('fornecedor')))
                contratos = contratos.filter(id__in=itens.values_list('contrato', flat=True))

            if visualizar == u'2':
                destino_arquivo = u'upload/resultados/relatorio_gerencial_contratos_%s.pdf' %  request.user.pessoafisica.id
                if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
                    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
                caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
                data_emissao = datetime.date.today()

                configuracao = get_config(request.user.pessoafisica.setor.secretaria)
                logo = None
                if configuracao.logo:
                    logo = os.path.join(settings.MEDIA_ROOT, configuracao.logo.name)



                data = {'contratos': contratos, 'titulo': 'Contratos', 'situacao': descricao_situacao, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'total': total }
                template = get_template('relatorio_gerencial_situacao_contratos.html')


                html  = template.render(Context(data))

                pdf_file = open(caminho_arquivo, "w+b")
                pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                        encoding='utf-8')
                pdf_file.close()
                file = open(caminho_arquivo, "r")
                pdf = file.read()
                file.close()
                return HttpResponse(pdf, 'application/pdf')
    else:
        raise PermissionDenied


    return render(request, 'ver_relatorios_gerenciais_contratos.html', locals(), RequestContext(request))

@login_required()
def ver_relatorios_gerenciais_atas(request):
    title=u'Relatórios Gerenciais das Atas'
    eh_gerente = request.user.groups.filter(name='Gerente')
    exibe_fornecedor = False
    if eh_gerente:

        form = RelatoriosGerenciaisContratosForm(request.POST or None, fornecedor=exibe_fornecedor)

        if form.is_valid():
            contratos = AtaRegistroPreco.objects.all().order_by('numero')
            relatorio =  form.cleaned_data.get('relatorio')
            situacao = form.cleaned_data.get('situacao')
            visualizar = form.cleaned_data.get('visualizar')
            secretaria = form.cleaned_data.get('secretaria')
            ano = form.cleaned_data.get('ano')
            hoje = datetime.date.today()
            total = 0
            if ano:
                contratos = contratos.filter(data_inicio__year=ano)

            if situacao:
                descricao_situacao = u'Todos'
                if situacao == u'2':
                    contratos = contratos.filter(concluido=False, suspenso=False, cancelado=False, data_inicio__lte=hoje, data_fim__gte=hoje)
                    descricao_situacao = u'Vigentes'
                elif situacao == u'3':
                    contratos = contratos.filter(Q(concluido=True), Q(suspenso=True), Q(cancelado=True))
                    descricao_situacao = u'Concluídos'

            if secretaria:
                contratos = contratos.filter(solicitacao__setor_origem__secretaria=secretaria)


            if visualizar == u'2':
                destino_arquivo = u'upload/resultados/relatorio_gerencial_contratos_%s.pdf' %  request.user.pessoafisica.id
                if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
                    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
                caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
                data_emissao = datetime.date.today()

                configuracao = get_config(request.user.pessoafisica.setor.secretaria)
                logo = None
                if configuracao.logo:
                    logo = os.path.join(settings.MEDIA_ROOT, configuracao.logo.name)



                data = {'contratos': contratos, 'titulo': 'Atas', 'situacao': descricao_situacao, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'total': total }
                template = get_template('relatorio_gerencial_situacao_contratos.html')


                html  = template.render(Context(data))

                pdf_file = open(caminho_arquivo, "w+b")
                pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                        encoding='utf-8')
                pdf_file.close()
                file = open(caminho_arquivo, "r")
                pdf = file.read()
                file.close()
                return HttpResponse(pdf, 'application/pdf')
    else:
        raise PermissionDenied


    return render(request, 'ver_relatorios_gerenciais_contratos.html', locals(), RequestContext(request))

@login_required()
def ver_relatorios_gerenciais_credenciamentos(request):
    title=u'Relatórios Gerenciais dos Credenciamentos'
    eh_gerente = request.user.groups.filter(name='Gerente')
    exibe_fornecedor = False
    if eh_gerente:

        form = RelatoriosGerenciaisContratosForm(request.POST or None, fornecedor=exibe_fornecedor)

        if form.is_valid():
            contratos = Credenciamento.objects.all().order_by('numero')
            relatorio =  form.cleaned_data.get('relatorio')
            situacao = form.cleaned_data.get('situacao')
            visualizar = form.cleaned_data.get('visualizar')
            secretaria = form.cleaned_data.get('secretaria')
            ano = form.cleaned_data.get('ano')
            hoje = datetime.date.today()
            total = 0
            if ano:
                contratos = contratos.filter(data_inicio__year=ano)

            if situacao:
                descricao_situacao = u'Todos'
                if situacao == u'2':
                    contratos = contratos.filter(concluido=False, suspenso=False, cancelado=False, data_inicio__lte=hoje, data_fim__gte=hoje)
                    descricao_situacao = u'Vigentes'
                elif situacao == u'3':
                    contratos = contratos.filter(Q(concluido=True), Q(suspenso=True), Q(cancelado=True))
                    descricao_situacao = u'Concluídos'

            if secretaria:
                contratos = contratos.filter(solicitacao__setor_origem__secretaria=secretaria)


            if visualizar == u'2':
                destino_arquivo = u'upload/resultados/relatorio_gerencial_contratos_%s.pdf' %  request.user.pessoafisica.id
                if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
                    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
                caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
                data_emissao = datetime.date.today()

                configuracao = get_config(request.user.pessoafisica.setor.secretaria)
                logo = None
                if configuracao.logo:
                    logo = os.path.join(settings.MEDIA_ROOT, configuracao.logo.name)



                data = {'contratos': contratos, 'titulo': 'Credenciamentos', 'situacao': descricao_situacao, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'total': total }
                template = get_template('relatorio_gerencial_situacao_contratos.html')


                html  = template.render(Context(data))

                pdf_file = open(caminho_arquivo, "w+b")
                pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                        encoding='utf-8')
                pdf_file.close()
                file = open(caminho_arquivo, "r")
                pdf = file.read()
                file.close()
                return HttpResponse(pdf, 'application/pdf')
    else:
        raise PermissionDenied


    return render(request, 'ver_relatorios_gerenciais_contratos.html', locals(), RequestContext(request))


@login_required()
def liberar_pedidos_solicitacao(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_cadastrar_solicitacao') and solicitacao.recebida_setor(request.user.pessoafisica.setor):
        if solicitacao.liberada_para_pedido:
            solicitacao.liberada_para_pedido = False
            messages.success(request, u'Liberação encerrada com sucesso.')
        else:
            solicitacao.liberada_para_pedido = True
            messages.success(request, u'Liberação realizada com sucesso.')
        solicitacao.save()

        return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' %  solicitacao.id)
    else:
        raise PermissionDenied


@login_required()
def relatorio_dados_licitacao(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    destino_arquivo = u'upload/dados_licitacao_procedimento/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/dados_licitacao_procedimento')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/dados_licitacao_procedimento'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    data = {'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'pregao':pregao}

    template = get_template('relatorio_dados_licitacao.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def criar_contrato_adesao_ata(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)
    title = u'Criar Contrato - %s' % ata
    form = CriarContratoAdesaoAtaForm(request.POST or None, request.FILES or None, ata=ata)
    if form.is_valid():
        for participante in ata.get_fornecedores():

            o = Contrato()
            o.solicitacao = ata.solicitacao
            o.valor = ata.get_valor_total(participante)
            o.numero = form.cleaned_data.get('contrato_%d' % participante.id)
            o.aplicacao_artigo_57 = form.cleaned_data.get('aplicacao_artigo_57_%d' % participante.id)
            o.garantia_execucao_objeto = form.cleaned_data.get('garantia_%d' % participante.id)

            o.data_inicio = form.cleaned_data.get('data_inicial_%d' % participante.id)
            o.data_fim = form.cleaned_data.get('data_final_%d' % participante.id)
            o.save()
            for resultado in ItemAtaRegistroPreco.objects.filter(ata=ata, fornecedor=participante):
                novo_item = ItemContrato()
                novo_item.contrato = o
                novo_item.item = resultado.item
                novo_item.material = resultado.material
                novo_item.marca = resultado.marca
                novo_item.fornecedor = resultado.fornecedor
                novo_item.valor = resultado.valor
                novo_item.quantidade = resultado.quantidade
                novo_item.unidade = resultado.unidade
                novo_item.ordem = o.get_ordem()
                novo_item.save()


        messages.success(request, u'Contrato(s) criado(s) com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' %  o.id)
    return render(request, 'criar_contrato_adesao_ata.html', locals(), RequestContext(request))

def erro_500(request):
    return render(request, '500.html', locals(), RequestContext(request))

def erro_404(request):
    return render(request, '404.html', locals(), RequestContext(request))

def erro_403(request):
    return render(request, '403.html', locals(), RequestContext(request))

@login_required()
def cadastrar_empresa_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)
    pode_gerenciar = credenciamento.solicitacao.recebida_setor(request.user.pessoafisica.setor)
    eh_gerente = request.user.groups.filter(name='Gerente') and pode_gerenciar
    if pode_gerenciar:
        title = u'Cadastrar Empresa'
        form = EmpresaCredenciamentoForm(request.POST or None)
        if form.is_valid():
            if ParticipantePregao.objects.filter(pregao=credenciamento.pregao, fornecedor=form.cleaned_data.get('fornecedor')).exists():
                messages.error(request, u'Esta empresa já é fornecedora.')
                return HttpResponseRedirect(u'/base/cadastrar_empresa_credenciamento/%s/' %  credenciamento.id)

            novo_participante = ParticipantePregao()
            novo_participante.fornecedor = form.cleaned_data.get('fornecedor')
            novo_participante.pregao = credenciamento.pregao
            novo_participante.me_epp = form.cleaned_data.get('me_epp')
            novo_participante.nome_representante = ''
            novo_participante.save()

            for item in credenciamento.solicitacao.itemsolicitacaolicitacao_set.all():
                ordem_atual = ResultadoItemPregao.objects.filter(item=item).order_by('-ordem')[0].ordem
                resultado = ResultadoItemPregao()
                resultado.item = item
                resultado.participante = novo_participante
                resultado.valor = item.valor_medio
                resultado.ordem = ordem_atual + 1
                resultado.situacao = ResultadoItemPregao.CLASSIFICADO
                resultado.save()
            messages.success(request, u'Empresa cadastrada com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/' %  credenciamento.id)

    else:
        raise PermissionDenied
    return render(request, 'cadastra_visitante_pregao.html', locals(), RequestContext(request))

@login_required()
def cadastrar_termo_inexigibilidade(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    if request.user.has_perm('base.pode_avaliar_minuta') and solicitacao.recebida_setor(request.user.pessoafisica.setor) and not solicitacao.termo_inexigibilidade:
        title=u'Cadastrar Termo de Inexigibilidade - %s' % solicitacao
        form = CadastroTermoInexigibilidadeForm(request.POST or None, request.FILES or None, instance=solicitacao)
        if form.is_valid():
            form.save()
            messages.success(request, u'Termo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/itens_solicitacao/%s/' % solicitacao.id)
        return render(request, 'cadastrar_minuta.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def relatorio_propostas(request, solicitacao_id):
    solicitacao = get_object_or_404(SolicitacaoLicitacao, pk=solicitacao_id)
    configuracao = get_config(solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)
    destino_arquivo = u'upload/resultados/%s.pdf' % solicitacao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()

    tabela = {}

    resultado = ItemSolicitacaoLicitacao.objects.filter(solicitacao=solicitacao, eh_lote=False).order_by('item')

    for num in resultado.order_by('item'):
        chave = '%s' % str(num.item)
        tabela[chave] = dict(lance = list(), total = 0)


    for item in resultado.order_by('item'):

        chave = '%s' % str(item.item)
        valor = 0
        total = ItemPesquisaMercadologica.objects.filter(item=item).count()
        for proposta in ItemPesquisaMercadologica.objects.filter(item=item):
            tabela[chave]['lance'].append(proposta)
            valor = valor + proposta.valor_maximo

        tabela[chave]['total'] = valor / total

    from blist import sorteddict

    def my_key(dict_key):
           try:
               with transaction.atomic():
                  return int(dict_key)
           except ValueError:
                  return dict_key


    resultado =  sorteddict(my_key, **tabela)


    data = { 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'solicitacao':solicitacao, 'resultado':resultado}

    template = get_template('relatorio_propostas.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def ver_crc(request, fornecedor_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
        title = u'Certificado de Registro Cadastral - %s' % fornecedor
        if FornecedorCRC.objects.filter(fornecedor=fornecedor).exists():
            registro = FornecedorCRC.objects.filter(fornecedor=fornecedor)[0]
            cnaes = CnaeSecundario.objects.filter(crc=registro)
            socios = SocioCRC.objects.filter(crc=registro)
            certidoes = CertidaoCRC.objects.filter(crc=registro)
        return render(request, 'ver_crc.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def cadastrar_crc(request, fornecedor_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
        title = u'Cadastrar CRC - %s' % fornecedor
        if FornecedorCRC.objects.filter(fornecedor=fornecedor).exists():
            registro = FornecedorCRC.objects.filter(fornecedor=fornecedor)[0]
        else:
            registro = FornecedorCRC()
            registro.fornecedor = fornecedor
            registro.validade = datetime.date.today() + timedelta(days=365)
            registro.numero = registro.get_proximo_numero(datetime.date.today().year)
            registro.ano = datetime.date.today().year

        form = CRCForm(request.POST or None, instance=registro)
        if form.is_valid():
            o = form.save(False)

            o.save()
            messages.success(request, u'CRC cadastrado/editado com sucesso.')
            return HttpResponseRedirect(u'/base/ver_crc/%s/' % fornecedor.id)

        return render(request, 'cadastrar_crc.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def cadastrar_cnaes_secundario(request, crc_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        crc = get_object_or_404(FornecedorCRC, pk=crc_id)
        title = u'Cadastrar CNAES Secundário - %s' % crc.fornecedor

        registro = CnaeSecundario()
        registro.crc = crc

        form = CNAESForm(request.POST or None, instance=registro)
        if form.is_valid():
            form.save()
            messages.success(request, u'CNAES Secundário cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/ver_crc/%s/' % crc.fornecedor.id)

        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def editar_cnaes_secundario(request, item_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        item = get_object_or_404(CnaeSecundario, pk=item_id)
        title = u'Editar CNAES Secundário - %s' % item.crc.fornecedor

        form = CNAESForm(request.POST or None, instance=item)
        if form.is_valid():
            form.save()
            messages.success(request, u'CNAES Secundário editado com sucesso.')
            return HttpResponseRedirect(u'/base/ver_crc/%s/' % item.crc.fornecedor.id)

        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def excluir_cnaes_secundario(request, item_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        item = get_object_or_404(CnaeSecundario, pk=item_id)
        crc = item.crc
        item.delete()
        messages.success(request, u'CNAES Secundário excluído com sucesso.')
        return HttpResponseRedirect(u'/base/ver_crc/%s/' % crc.fornecedor.id)


    else:
        raise PermissionDenied

@login_required()
def editar_socio(request, item_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        item = get_object_or_404(SocioCRC, pk=item_id)
        title = u'Editar Sócio - %s' % item.crc.fornecedor

        form = SocioForm(request.POST or None, instance=item)
        if form.is_valid():
            form.save()
            messages.success(request, u'Sócio editado com sucesso.')
            return HttpResponseRedirect(u'/base/ver_crc/%s/' % item.crc.fornecedor.id)

        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def excluir_socio(request, item_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        item = get_object_or_404(SocioCRC, pk=item_id)
        crc = item.crc
        item.delete()
        messages.success(request, u'Sócio excluído com sucesso.')
        return HttpResponseRedirect(u'/base/ver_crc/%s/' % crc.fornecedor.id)


    else:
        raise PermissionDenied


@login_required()
def cadastrar_socio(request, crc_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        crc = get_object_or_404(FornecedorCRC, pk=crc_id)
        title = u'Cadastrar Sócio - %s' % crc.fornecedor

        registro = SocioCRC()
        registro.crc = crc
        form = SocioForm(request.POST or None, instance=registro)
        if form.is_valid():
            form.save()
            messages.success(request, u'Sócio com sucesso.')
            return HttpResponseRedirect(u'/base/ver_crc/%s/' % crc.fornecedor.id)

        return render(request, 'cadastrar_anexo_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def imprimir_crc(request, fornecedor_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
        registro = get_object_or_404(FornecedorCRC, fornecedor=fornecedor)
        configuracao = get_config_geral()
        logo = None
        if configuracao.logo:
            logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

        destino_arquivo = u'upload/extratos/%s.pdf' % fornecedor.id
        if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/extratos')):
            os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/extratos'))
        caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)


        data = {'registro': registro, 'configuracao': configuracao, 'logo': logo}

        template = get_template('imprimir_crc.html')

        html  = template.render(Context(data))

        pdf_file = open(caminho_arquivo, "w+b")
        pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                encoding='utf-8')
        pdf_file.close()


        file = open(caminho_arquivo, "r")
        pdf = file.read()
        file.close()
        return HttpResponse(pdf, 'application/pdf')
    else:
        raise PermissionDenied


@login_required()
def renovar_crc(request, fornecedor_id):
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
        registro = get_object_or_404(FornecedorCRC, fornecedor=fornecedor)
        title = u'Renovar Validade do CRC'
        form = DataRenovaCRCForm(request.POST or None)
        if form.is_valid():
            registro.validade = form.cleaned_data.get('data') + timedelta(days=364)
            registro.save()
            messages.success(request, u'CRC renovado com sucesso.')
            return HttpResponseRedirect(u'/base/ver_crc/%s/' % fornecedor.id)
        return render(request, 'renovar_crc.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def aditivar_contrato(request, contrato_id):
    if request.user.has_perm('base.pode_gerenciar_contrato'):
        contrato = get_object_or_404(Contrato, pk=contrato_id)
        title = u'Aditivar Contrato: %s' % contrato
        form = AditivarContratoForm(request.POST or None, contrato=contrato)
        itens = ItemContrato.objects.filter(contrato=contrato).order_by('item__item')
        if form.is_valid():
            aditivo = Aditivo()
            aditivo.contrato = contrato
            aditivo.ordem = contrato.get_ordem()

            if form.cleaned_data.get('data_final'):
                aditivo.de_prazo = True
                aditivo.data_inicio = form.cleaned_data.get('data_inicial')
                aditivo.data_fim = form.cleaned_data.get('data_final')

            aditivo.tipo = form.cleaned_data.get('opcoes')
            if form.cleaned_data.get('opcoes'):

                aditivo.indice = form.cleaned_data.get('indice_reajuste')

                if form.cleaned_data.get('opcoes') == Aditivo.REAJUSTE_FINANCEIRO:
                    aditivo.de_valor = True
                    aditivo.valor = ((form.cleaned_data.get('indice_reajuste')/100) * contrato.get_valor_aditivado())
                    for item in ItemContrato.objects.filter(contrato=contrato):
                        item.valor = item.valor + ((form.cleaned_data.get('indice_reajuste')/100) * item.valor)
                        item.save()
                else:

                    total_ajuste = 0
                    qtd_ajuste = 0
                    if form.cleaned_data.get('opcoes') == Aditivo.ACRESCIMO_QUANTITATIVOS:
                        aditivo.de_valor = True

                        for idx, indice_informado in enumerate(request.POST.getlist('quantidade_soma'), 1):

                            if indice_informado and int(indice_informado) > 0:
                                item = ItemContrato.objects.get(contrato=contrato, id=request.POST.getlist('id_item')[idx-1])
                                indice_ajuste = (Decimal(request.POST.getlist('quantidade_soma')[idx-1].replace('.','').replace(',','.')) * 100)  / item.quantidade


                                item.save()
                                aditivo_item = AditivoItemContrato()
                                aditivo_item.item = item
                                aditivo_item.indice = indice_ajuste
                                aditivo_item.tipo = form.cleaned_data.get('opcoes')
                                aditivo_item.valor = Decimal(request.POST.getlist('quantidade_soma')[idx-1].replace('.','').replace(',','.'))
                                aditivo_item.save()

                                aditivo.indice = indice_ajuste
                                aditivo.save()

                                total_ajuste +=  Decimal(request.POST.getlist('quantidade_soma')[idx-1].replace('.','').replace(',','.'))
                                qtd_ajuste += 1

                    elif form.cleaned_data.get('opcoes') == Aditivo.SUPRESSAO_QUANTITATIVO:
                        aditivo.de_valor = True
                        for idx, indice_informado in enumerate(request.POST.getlist('quantidade_subtrai'), 1):
                            if indice_informado and int(indice_informado) > 0:
                                item = ItemContrato.objects.get(contrato=contrato, id=request.POST.getlist('id_item')[idx-1])
                                indice_ajuste = (Decimal(request.POST.getlist('quantidade_subtrai')[idx-1].replace('.','').replace(',','.')) * 100) / item.quantidade


                                item.save()

                                aditivo_item = AditivoItemContrato()
                                aditivo_item.item = item
                                aditivo_item.indice = indice_ajuste
                                aditivo_item.tipo = form.cleaned_data.get('opcoes')
                                aditivo_item.valor = Decimal(request.POST.getlist('quantidade_subtrai')[idx-1].replace('.','').replace(',','.'))
                                aditivo_item.save()
                                aditivo.indice = indice_ajuste
                                aditivo.save()


                                total_ajuste +=  Decimal(request.POST.getlist('quantidade_subtrai')[idx-1].replace('.','').replace(',','.'))
                                qtd_ajuste += 1

                    elif form.cleaned_data.get('opcoes') == Aditivo.ACRESCIMO_VALOR:
                        aditivo.de_valor = True

                        for idx, indice_informado in enumerate(request.POST.getlist('valor_soma'), 1):
                            if indice_informado and int(indice_informado) > 0:

                                item = ItemContrato.objects.get(contrato=contrato, id=request.POST.getlist('id_item')[idx-1])
                                #item.valor = ((Decimal(request.POST.getlist('valor_soma')[idx-1].replace('.','').replace(',','.'))/100) * item.valor) + item.valor
                                #item.save()
                                aditivo_item = AditivoItemContrato()
                                aditivo_item.item = item
                                indice_ajuste = Decimal(request.POST.getlist('valor_soma')[idx-1].replace('.','').replace(',','.'))
                                aditivo_item.indice = indice_ajuste
                                aditivo_item.tipo = form.cleaned_data.get('opcoes')
                                aditivo_item.valor = ((Decimal(request.POST.getlist('valor_soma')[idx-1].replace('.','').replace(',','.'))/100) * item.valor)
                                aditivo_item.save()

                                aditivo.indice = indice_ajuste
                                aditivo.save()

                                total_ajuste +=  Decimal(request.POST.getlist('valor_soma')[idx-1].replace('.','').replace(',','.'))
                                qtd_ajuste += 1

                    elif form.cleaned_data.get('opcoes') == Aditivo.SUPRESSAO_VALOR:
                        aditivo.de_valor = True

                        for idx, indice_informado in enumerate(request.POST.getlist('valor_subtrai'), 1):
                            if indice_informado and int(indice_informado) > 0:
                                item = ItemContrato.objects.get(contrato=contrato, id=request.POST.getlist('id_item')[idx-1])
                                #item.valor = item.valor - ((Decimal(request.POST.getlist('valor_subtrai')[idx-1].replace('.','').replace(',','.'))/100) * item.valor)
                                #item.save()

                                aditivo_item = AditivoItemContrato()
                                aditivo_item.item = item
                                indice_ajuste = Decimal(request.POST.getlist('valor_subtrai')[idx-1].replace('.','').replace(',','.'))
                                aditivo_item.indice = indice_ajuste
                                aditivo_item.tipo = form.cleaned_data.get('opcoes')
                                aditivo_item.valor = ((Decimal(request.POST.getlist('valor_subtrai')[idx-1].replace('.','').replace(',','.'))/100) * item.valor)
                                aditivo_item.save()

                                aditivo.indice = indice_ajuste
                                aditivo.save()

                                total_ajuste +=  Decimal(request.POST.getlist('valor_subtrai')[idx-1].replace('.','').replace(',','.'))
                                qtd_ajuste += 1

                    aditivo.valor =   total_ajuste /   qtd_ajuste

            valor_final = Decimal(0.00)
            for item in ItemContrato.objects.filter(contrato=contrato):
                quantidade_aditivo = 0
                aditivos = AditivoItemContrato.objects.filter(item=item)
                if aditivos.exists():
                    for adit in aditivos:
                        if adit.tipo == Aditivo.ACRESCIMO_QUANTITATIVOS:
                            if adit.valor:
                                quantidade_aditivo += adit.valor
                        elif adit.tipo == Aditivo.SUPRESSAO_QUANTITATIVO:
                            if adit.valor:
                                quantidade_aditivo -= adit.valor
                valor_final += (item.get_valor_item_contrato(numero=True) * (item.quantidade + quantidade_aditivo)).quantize(Decimal(10) ** -2)

            aditivo.valor_atual = valor_final
            if form.cleaned_data.get('opcoes') == Aditivo.REAJUSTE_FINANCEIRO:
                aditivo.indice_total_contrato = form.cleaned_data.get('indice_reajuste')
            else:
                reducao = (valor_final - contrato.get_valor_aditivado()) / (contrato.get_valor_aditivado()/100)
                if Aditivo.objects.filter(contrato=contrato, ordem=aditivo.ordem-1).exists():
                    reducao = reducao - Aditivo.objects.filter(contrato=contrato, ordem=aditivo.ordem-1)[0].indice_total_contrato
                aditivo.indice_total_contrato = reducao
            aditivo.save()
            messages.success(request, u'Aditivo cadastrado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' % contrato.id)

        return render(request, 'aditivar_contrato.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied


@login_required()
def rescindir_contrato(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)
    title = u'Rescindir Contrato'
    if request.user.has_perm('base.pode_gerenciar_contrato'):
        form = RescindirContratoForm(request.POST or None)
        if form.is_valid():
            if form.cleaned_data.get('opcao'):
                if form.cleaned_data.get('opcao') == u'Arquivar Contrato':
                    contrato.liberada_compra = False
                    contrato.cancelado = True
                    contrato.save()
                    messages.success(request, u'Contrato arquivado com sucesso.')
                    return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' % contrato.id)

                elif form.cleaned_data.get('opcao') == u'Contratar Remanescentes':
                    return HttpResponseRedirect(u'/base/contratar_remanescentes/%s/' % contrato.id)

        return render(request, 'rescindir_contrato.html', locals(), RequestContext(request))


    else:
        raise PermissionDenied

@login_required()
def contratar_remanescentes(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)
    title = u'Contratar Remanescentes - %s' % contrato
    if request.user.has_perm('base.pode_gerenciar_contrato'):
        itens = ItemContrato.objects.filter(contrato=contrato, inserido_outro_contrato=False)
        form = ContratoRemanescenteForm(request.POST or None, pregao=contrato.pregao, contrato=contrato)
        if form.is_valid():
            o = form.save(False)
            o.valor = 1.00
            o.solicitacao = contrato.solicitacao
            o.pregao = contrato.pregao
            o.save()
            valor_total = 0
            itens_selecionados = request.POST.getlist('registros')
            for item in itens_selecionados:
                copia_item_atual = get_object_or_404(ItemContrato, pk=int(item))
                novo_item = copia_item_atual
                novo_item.id = None
                novo_item.contrato = o
                novo_item.fornecedor = form.cleaned_data.get('fornecedor').fornecedor
                novo_item.participante = form.cleaned_data.get('fornecedor')
                item_atual = get_object_or_404(ItemContrato, pk=int(item))
                novo_item.quantidade = item_atual.get_quantidade_disponivel()
                novo_item.inserido_outro_contrato = False
                novo_item.save()
                valor_total += novo_item.valor * novo_item.quantidade
                pega_item = get_object_or_404(ItemContrato, pk=int(item))
                pega_item.inserido_outro_contrato = True
                pega_item.save()
            o.valor = valor_total
            o.save()
            contrato.cancelado = True
            contrato.liberada_compra = False
            contrato.save()
            messages.success(request, u'Contrato gerado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_contrato/%s/' % o.id)


        return render(request, 'contratar_remanescentes.html', locals(), RequestContext(request))

    else:
        raise PermissionDenied


@login_required()
def salvar_sorteio_item_pregao(request, item_id):

    item = get_object_or_404(ItemSolicitacaoLicitacao, pk= item_id)
    pregao = item.solicitacao.get_pregao()
    if request.user.has_perm('base.pode_cadastrar_pregao') and pregao.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        ordenar = request.POST.getlist('ordens')
        if ordenar:
            if len(ordenar) != len(set(ordenar)):
                messages.error(request, u'Não é permitido ter duas ou mais empresas com a mesma colocação.')
                return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
            lista = list()
            for proposta in PropostaItemPregao.objects.filter(item=item, concorre=True).order_by('-concorre', 'desclassificado','desistencia', 'valor'):
                lista.append(proposta)
            for idx, ordem_informada in enumerate(ordenar, 0):
                if ordem_informada:
                    plano = ResultadoItemPregao.objects.get(participante=lista[idx].participante, item=lista[idx].item)
                    plano.ordem = int(ordem_informada)
                    plano.save()
            ResultadoItemPregao.objects.filter(item=item).update(empate=False)
            messages.success(request, u'Ordem do sorteio salva com sucesso.')
            return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
    else:
        raise PermissionDenied


@login_required()
def mudar_credenciamento_fornecedor(request, credenciamento_id, fornecedor_id, opcao):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)

    pode_gerenciar = credenciamento.solicitacao.recebida_setor(request.user.pessoafisica.setor)
    eh_gerente = request.user.groups.filter(name='Gerente') and pode_gerenciar

    if eh_gerente:
        participante = get_object_or_404(ParticipantePregao, pk=fornecedor_id)
        if opcao == u'1':
            participante.credenciado = True
        elif opcao == u'2':
            participante.credenciado = False
        participante.save()
        messages.success(request, u'Credenciamento do fornecedor alterado com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_credenciamento/%s/' % credenciamento.id)

    else:
        raise PermissionDenied

@login_required()
def relatorio_info_contrato(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)

    total = 0
    itens = contrato.get_itens()
    pedidos = PedidoContrato.objects.filter(contrato=contrato).order_by('pedido_em')
    for item in itens:
        total += item.get_valor_total()
    configuracao = get_config(contrato.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % contrato_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'contrato':contrato, 'pedidos': pedidos, 'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_info_contrato.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_qtd_disponivel_contrato(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)

    total = 0
    itens = contrato.get_itens()

    for item in itens:
        total += item.get_valor_total_disponivel()
    configuracao = get_config(contrato.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % contrato_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'contrato':contrato,  'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_qtd_disponivel_contrato.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_qtd_consumida_contrato(request, contrato_id):
    contrato = get_object_or_404(Contrato, pk=contrato_id)


    total = 0
    itens = contrato.get_itens()
    pedidos = PedidoContrato.objects.filter(contrato=contrato).order_by('pedido_em')
    for item in itens:
        total += item.get_valor_total_consumido()
    configuracao = get_config(contrato.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % contrato_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'contrato':contrato, 'pedidos': pedidos, 'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_qtd_consumida_contrato.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_info_arp(request, ata_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)

    total = 0
    itens = ItemAtaRegistroPreco.objects.filter(ata=ata)

    pedidos = PedidoAtaRegistroPreco.objects.filter(ata=ata).order_by('pedido_em')
    for item in itens:
        total += item.valor * item.quantidade
    configuracao = get_config(ata.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % ata_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'ata':ata, 'pedidos': pedidos, 'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_info_arp.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_qtd_disponivel_ata(request, ata_id, fornecedor_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)

    if fornecedor_id == '0':
        itens = ItemAtaRegistroPreco.objects.filter(ata=ata)
        fornecedor = None
    else:
        if ata.adesao:
            fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
            itens = ItemAtaRegistroPreco.objects.filter(ata=ata, fornecedor=fornecedor)
        else:
            part_pregao = get_object_or_404(ParticipantePregao, pk=fornecedor_id)
            fornecedor = part_pregao.fornecedor
            itens = ItemAtaRegistroPreco.objects.filter(ata=ata, participante=part_pregao)
    total = 0
    for item in itens:
        total += item.get_valor_total_disponivel()
    configuracao = get_config(ata.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % ata.id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'ata':ata, 'fornecedor': fornecedor,  'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_qtd_disponivel_ata.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_qtd_consumida_ata(request, ata_id, fornecedor_id):
    ata = get_object_or_404(AtaRegistroPreco, pk=ata_id)

    if fornecedor_id == '0':
        itens = ItemAtaRegistroPreco.objects.filter(ata=ata)
        fornecedor = None
    else:
        if ata.adesao:
            fornecedor = get_object_or_404(Fornecedor, pk=fornecedor_id)
            itens = ItemAtaRegistroPreco.objects.filter(ata=ata, fornecedor=fornecedor)
        else:
            part_pregao = get_object_or_404(ParticipantePregao, pk=fornecedor_id)
            fornecedor = part_pregao.fornecedor
            itens = ItemAtaRegistroPreco.objects.filter(ata=ata, participante=part_pregao)

    total = 0
    itens = ItemAtaRegistroPreco.objects.filter(ata=ata)
    pedidos = PedidoAtaRegistroPreco.objects.filter(ata=ata).order_by('pedido_em')
    for item in itens:
        total += item.get_valor_total_consumido()
    configuracao = get_config(ata.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % ata.id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'ata':ata, 'fornecedor':fornecedor, 'pedidos': pedidos, 'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_qtd_consumida_ata.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_info_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)


    total = 0
    itens = credenciamento.itemcredenciamento_set.all()
    pedidos = PedidoCredenciamento.objects.filter(credenciamento=credenciamento).order_by('pedido_em')
    for item in itens:
        total += item.valor * item.quantidade
    configuracao = get_config(credenciamento.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % credenciamento_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'credenciamento':credenciamento, 'pedidos': pedidos, 'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_info_credenciamento.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_qtd_disponivel_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)


    total = 0
    itens = credenciamento.itemcredenciamento_set.all()

    for item in itens:
        total += item.get_quantidade_disponivel() * item.valor
    configuracao = get_config(credenciamento.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % credenciamento_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'credenciamento':credenciamento,  'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_qtd_disponivel_credenciamento.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')



@login_required()
def relatorio_qtd_consumida_credenciamento(request, credenciamento_id):
    credenciamento = get_object_or_404(Credenciamento, pk=credenciamento_id)

    total = 0
    itens = credenciamento.itemcredenciamento_set.all()
    pedidos = PedidoCredenciamento.objects.filter(credenciamento=credenciamento).order_by('pedido_em')
    for item in itens:
        total += item.get_valor_total_consumido()
    configuracao = get_config(credenciamento.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    destino_arquivo = u'upload/resultados/%s.pdf' % credenciamento_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()


    data = {'credenciamento':credenciamento, 'pedidos': pedidos, 'itens':itens, 'total': total, 'configuracao':configuracao, 'logo':logo,  'data_emissao':data_emissao}

    template = get_template('relatorio_qtd_consumida_credenciamento.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def notificacoes(request):
    title = u'Módulo de Notificações'
    hoje = datetime.date.today()

    tem_notificacao = False
    if request.user.has_perm('base.pode_cadastrar_pregao'):
        prazo_15dias = hoje + timedelta(days=15)
        licitacoes_agendadas = Pregao.objects.filter(data_abertura__gte=hoje, data_abertura__lte=prazo_15dias).order_by('data_abertura')
        if licitacoes_agendadas.exists():
            tem_notificacao = True

    if request.user.has_perm('base.pode_gerenciar_contrato'):
        ids_a_vencer = list()
        contratos_a_vencer = Contrato.objects.filter(suspenso=False, cancelado=False, concluido=False)
        hoje = datetime.datetime.now().date()
        for contrato in contratos_a_vencer:
            vencimento = contrato.get_data_fim()
            if vencimento > hoje and vencimento < (hoje + timedelta(days=30)):
                ids_a_vencer.append(contrato.id)
        contratos_a_vencer = contratos_a_vencer.filter(id__in=ids_a_vencer)


        ids_atas_a_vencer = list()
        atas_a_vencer = AtaRegistroPreco.objects.filter(suspenso=False, cancelado=False, concluido=False)
        hoje = datetime.datetime.now().date()
        for ata in atas_a_vencer:
            vencimento = ata.data_fim
            if vencimento > hoje and vencimento < (hoje + timedelta(days=30)):
                ids_atas_a_vencer.append(ata.id)
        atas_a_vencer = atas_a_vencer.filter(id__in=ids_atas_a_vencer)

        contratos_sem_vigencia = Contrato.objects.filter(suspenso=False, cancelado=False, concluido=False, data_inicio__lte=hoje)
        ids_vencidos = list()
        for item in contratos_sem_vigencia:
            if item.get_data_fim() < hoje:
                if item.eh_gerente(request.user):
                    ids_vencidos.append(item.id)
                    tem_notificacao = True
        contratos_sem_vigencia = contratos_sem_vigencia.filter(id__in=ids_vencidos)

    return render(request, 'notificacoes.html', locals(), RequestContext(request))

@login_required()
def editar_item_arp(request, item_id):
    item = get_object_or_404(ItemAtaRegistroPreco, pk=item_id)
    ata = item.ata
    title=u'Editar Item'
    if ata.adesao and request.user == ata.solicitacao.cadastrado_por and ata.solicitacao.setor_atual == request.user.pessoafisica.setor and not ata.liberada_compra:
        form = EditarItemARPForm(request.POST or None, instance=item)
        if form.is_valid():
            o = form.save()
            messages.success(request, u'Item editado com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % item.ata.id)

        return render(request, 'cadastrar_pregao.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def apagar_item_arp(request, item_id):
    item = get_object_or_404(ItemAtaRegistroPreco, pk=item_id)
    ata = item.ata
    if ata.adesao and request.user == ata.solicitacao.cadastrado_por and ata.solicitacao.setor_atual == request.user.pessoafisica.setor and not ata.liberada_compra:
        item.delete()
        messages.success(request, u'Item removido com sucesso.')
        return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % ata.id)
    else:
        raise PermissionDenied

def anexo_38(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)


    import openpyxl
    from openpyxl.utils import get_column_letter
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=anexo_38_tce_licitacao_%s.xlsx' % pregao.id
    wb = openpyxl.Workbook()
    ws = wb.get_active_sheet()
    ws.title = "MyModel"

    row_num = 0

    columns = [
        (u"NumeroLote", 12),
        (u"DescricaoLote", 30),
        (u"OrdemClassificatoria", 18),
        (u"ValorProposta", 15),
        (u"NomeParticipante", 30),
        (u"TipoDocumento", 15),
        (u"NumeroDocumento", 20),
    ]

    for col_num in xrange(len(columns)):
        c = ws.cell(row=row_num + 1, column=col_num + 1)
        c.value = columns[col_num][0]
        #c.style.font.bold = True
        # set column width
        ws.column_dimensions[get_column_letter(col_num+1)].width = columns[col_num][1]


    if not pregao.solicitacao.eh_lote():
        itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False).order_by('item')
        contador_total = 0
        for idx, item in enumerate(itens, 0):

            resultado = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO, participante__excluido_dos_itens=False, participante__desclassificado=False, item__solicitacao=pregao.solicitacao, item__situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('ordem')
            contador = 1
            for result in resultado:
                if contador < 6:
                    row_index = contador_total + 1
                    # style = xlwt.XFStyle()
                    # style.alignment.wrap = 1
                    # w_sheet.write(row_index, 0, item.item)
                    # w_sheet.write(row_index, 1, item.material.nome[:100])
                    # w_sheet.write(row_index, 2, contador)
                    # w_sheet.write(row_index, 3, format_money(result.valor))
                    # w_sheet.write(row_index, 4, result.participante.fornecedor.razao_social)
                    # w_sheet.write(row_index, 5, u'CNPJ')
                    # w_sheet.write(row_index, 6, str(result.participante.fornecedor.cnpj).replace('.', '').replace('-', '').replace('/', ''))

                    row = [
                        item.item,
                        item.material.nome[:100],
                        contador,
                        format_money(result.valor),
                        result.participante.fornecedor.razao_social,
                        u'CNPJ',
                        str(result.participante.fornecedor.cnpj).replace('.', '').replace('-', '').replace('/', ''),


                    ]
                    for col_num in xrange(len(row)):
                        c = ws.cell(row=row_index + 1, column=col_num + 1)
                        c.value = row[col_num]
                        #c.style.alignment.wrap_text = True
                    contador += 1
                    contador_total += 1

    else:
        contador_total = 0
        contador = 1
        tabela = {}
        itens = {}
        resultado = ResultadoItemPregao.objects.filter(item__solicitacao=pregao.solicitacao, situacao=ResultadoItemPregao.CLASSIFICADO, participante__excluido_dos_itens=False, participante__desclassificado=False, item__situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('item__item')
        chaves =  resultado.values('item__item').order_by('item__item')
        for num in sorted(chaves):
            chave = '%s' % num['item__item']
            tabela[chave] = []
            itens[chave] =  []

        for item in resultado.order_by('item','ordem'):
            chave = '%s' % str(item.item.item)
            tabela[chave].append(item)
            itens[chave] = item.item.get_itens_do_lote()

        from blist import sorteddict

        def my_key(dict_key):
               try:
                   with transaction.atomic():
                      return int(dict_key)
               except ValueError:
                      return dict_key


        resultado =  sorteddict(my_key, **tabela)

        for result in resultado.items():

            if result[1]:
                for registro in result[1]:


                    if contador < 6:
                        row_index = contador_total + 1
                        # style = xlwt.XFStyle()
                        # style.alignment.wrap = 1
                        # w_sheet.write(row_index, 0, item.item)
                        # w_sheet.write(row_index, 1, item.material.nome[:100])
                        # w_sheet.write(row_index, 2, contador)
                        # w_sheet.write(row_index, 3, format_money(result.valor))
                        # w_sheet.write(row_index, 4, result.participante.fornecedor.razao_social)
                        # w_sheet.write(row_index, 5, u'CNPJ')
                        # w_sheet.write(row_index, 6, str(result.participante.fornecedor.cnpj).replace('.', '').replace('-', '').replace('/', ''))

                        row = [
                            result[0],
                            u'Lote %s' % result[0],
                            registro.ordem,
                            format_money(registro.valor),
                            registro.participante.fornecedor.razao_social,
                            u'CNPJ',
                            str(registro.participante.fornecedor.cnpj).replace('.', '').replace('-', '').replace('/', ''),


                        ]
                        for col_num in xrange(len(row)):
                            c = ws.cell(row=row_index + 1, column=col_num + 1)
                            c.value = row[col_num]
                            #c.style.alignment.wrap_text = True
                        contador += 1
                        contador_total += 1


    wb.save(response)
    return response















   #-------------------------------------------------------

    #
    # pregao = get_object_or_404(Pregao, pk=pregao_id)
    #
    # nome = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/modelo_anexo38')
    # file_path = os.path.join(settings.MEDIA_ROOT, 'upload/modelos/modelo_anexo38.xlsx')
    # rb = open_workbook(file_path)
    #
    # wb = copy(rb) # a writable copy (I can't read values out of this, only write to it)
    # w_sheet = wb.get_sheet(0) # the sheet to write to within the writable copy
    #
    # sheet = rb.sheet_by_name(u"Página1")
    # itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao, eh_lote=False).order_by('item')
    # contador_total = 0
    # for idx, item in enumerate(itens, 0):
    #
    #     resultado = ResultadoItemPregao.objects.filter(item=item, situacao=ResultadoItemPregao.CLASSIFICADO, participante__excluido_dos_itens=False, participante__desclassificado=False, item__solicitacao=pregao.solicitacao, item__situacao__in=[ItemSolicitacaoLicitacao.CADASTRADO, ItemSolicitacaoLicitacao.CONCLUIDO]).order_by('ordem')
    #     contador = 1
    #     for result in resultado:
    #         if contador < 6:
    #             row_index = contador_total + 1
    #             # style = xlwt.XFStyle()
    #             # style.alignment.wrap = 1
    #             w_sheet.write(row_index, 0, item.item)
    #             w_sheet.write(row_index, 1, item.material.nome[:100])
    #             w_sheet.write(row_index, 2, contador)
    #             w_sheet.write(row_index, 3, format_money(result.valor))
    #             w_sheet.write(row_index, 4, result.participante.fornecedor.razao_social)
    #             w_sheet.write(row_index, 5, u'CNPJ')
    #             w_sheet.write(row_index, 6, str(result.participante.fornecedor.cnpj).replace('.', '').replace('-', '').replace('/', ''))
    #
    #
    #             contador += 1
    #             contador_total += 1
    #         # w_sheet.write(row_index, 1, item.material.nome, style)
    #         # w_sheet.write(row_index, 2, item.unidade.nome)
    #
    #
    # salvou = nome + u'_%s' % pregao.id + '.xlsx'
    # wb.save(salvou)
    #
    # arquivo = open(salvou, "rb")
    #
    #
    # content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    # response = HttpResponse(arquivo.read(), content_type=content_type)
    # nome_arquivo = salvou.split('/')[-1]
    # response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    # arquivo.close()
    # os.unlink(salvou)
    # return response


@login_required()
def ativar_item_pregao(request, item_id):
    item = get_object_or_404(ItemSolicitacaoLicitacao, pk=item_id)
    if request.user.has_perm('base.pode_cadastrar_pregao') and item.solicitacao.recebida_setor(request.user.pessoafisica.setor):
        item.situacao = ItemSolicitacaoLicitacao.CADASTRADO
        item.save()
        historico = HistoricoPregao()
        historico.pregao = item.get_licitacao()
        historico.data = datetime.datetime.now()
        historico.obs = u'Item Reativado: %s.' % item
        historico.save()
        messages.success(request, u'Item reativado com sucesso.')
        return HttpResponseRedirect(u'/base/lances_item/%s/' % item.id)
    else:
        raise PermissionDenied


@login_required()
def relatorio_itens_desertos(request, pregao_id):
    pregao = get_object_or_404(Pregao, pk=pregao_id)

    destino_arquivo = u'upload/dados_licitacao_procedimento/%s.pdf' % pregao_id
    if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/dados_licitacao_procedimento')):
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/dados_licitacao_procedimento'))
    caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
    data_emissao = datetime.date.today()
    if pregao.comissao:
        configuracao = get_config(pregao.comissao.secretaria.ordenador_despesa.setor.secretaria)
    else:
        configuracao = get_config(pregao.solicitacao.setor_origem.secretaria)
    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)

    itens = ItemSolicitacaoLicitacao.objects.filter(solicitacao=pregao.solicitacao)
    itens = itens.filter(Q(situacao=ItemSolicitacaoLicitacao.FRACASSADO) | Q(situacao=ItemSolicitacaoLicitacao.DESERTO))
    data = {'configuracao':configuracao, 'logo':logo, 'itens': itens, 'data_emissao':data_emissao, 'pregao':pregao}

    template = get_template('relatorio_itens_desertos.html')

    html  = template.render(Context(data))

    pdf_file = open(caminho_arquivo, "w+b")
    pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
            encoding='utf-8')
    pdf_file.close()
    file = open(caminho_arquivo, "r")
    pdf = file.read()
    file.close()
    return HttpResponse(pdf, 'application/pdf')


@login_required()
def modelos_atas(request, pregao_id):
    title = u'Modelos de Atas'
    atas = ModeloAta.objects.all().order_by('-id')


    form = BuscarModeloAtaForm(request.GET or None)

    if form.is_valid():
        if form.cleaned_data.get('nome'):
            atas = atas.filter(nome=form.cleaned_data.get('nome'))

        if form.cleaned_data.get('palavra'):
            atas = atas.filter(palavras_chaves__icontains=form.cleaned_data.get('palavra'))

        if form.cleaned_data.get('tipo'):
            atas = atas.filter(tipo=form.cleaned_data.get('tipo'))



    return render(request, 'modelos_atas.html', locals(), RequestContext(request))

@login_required()
def cadastrar_modelo_ata(request, pregao_id):
    title = u'Cadastrar Modelo de Ata'
    form = ModeloAtaForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        o = form.save(False)
        o.cadastrado_em = datetime.datetime.now()
        o.cadastrado_por = request.user.pessoafisica
        o.save()
        messages.success(request, u'Modelo cadastrado com sucesso.')
        return HttpResponseRedirect(u'/base/modelos_atas/%s/' % pregao_id)

    return render(request, 'cadastrar_modelo_ata.html', locals(), RequestContext(request))

@login_required()
def editar_modelo_ata(request, ata_id, pregao_id):
    title = u'Editar Modelo de Ata'
    ata = get_object_or_404(ModeloAta, pk=ata_id)
    if request.user.pessoafisica == ata.cadastrado_por:
        form = ModeloAtaForm(request.POST or None, request.FILES or None, instance=ata)
        if form.is_valid():
            form.save()
            messages.success(request, u'Modelo editado com sucesso.')
            return HttpResponseRedirect(u'/base/modelos_atas/%s/' % pregao_id)

        return render(request, 'cadastrar_modelo_ata.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def deletar_modelo_ata(request, ata_id, pregao_id):
    ata = get_object_or_404(ModeloAta, pk=ata_id)
    if request.user.pessoafisica == ata.cadastrado_por:
        ata.delete()
        messages.success(request, u'Modelo excluído com sucesso.')
        return HttpResponseRedirect(u'/base/modelos_atas/%s/' % pregao_id)

    else:
        raise PermissionDenied


@login_required()
def ver_relatorios_gerenciais_compras(request):
    title=u'Relatórios Gerenciais de Compras'

    if request.user.has_perm('base.pode_cadastrar_pesquisa_mercadologica'):

        form = RelatoriosGerenciaisComprasForm(request.POST or None)

        if form.is_valid():
            ordens = OrdemCompra.objects.all().order_by('numero')
            relatorio =  form.cleaned_data.get('relatorio')
            tipo = form.cleaned_data.get('tipo_ordem')
            visualizar = form.cleaned_data.get('visualizar')
            secretaria = form.cleaned_data.get('secretaria')
            ano = form.cleaned_data.get('ano')
            hoje = datetime.date.today()
            total = 0
            if ano:
                ordens = ordens.filter(data__year=ano)

            if ordens:
                descricao_situacao = u'Todos os Tipos'
                if tipo == u'2':
                    ordens = ordens.filter(tipo=u'Compras')
                    descricao_situacao = u'Tipo: Compras'
                elif tipo == u'3':
                    ordens = ordens.filter(tipo=u'Serviços')
                    descricao_situacao = u'Tipo: Serviços'

            if secretaria:
                ordens = ordens.filter(solicitacao__setor_origem__secretaria=secretaria)


            total = Decimal(0.00)
            for ordem in ordens:
                total += ordem.get_valor_global()

            if visualizar == u'2':
                destino_arquivo = u'upload/resultados/relatorio_gerencial_compras_%s.pdf' %  request.user.pessoafisica.id
                if not os.path.exists(os.path.join(settings.MEDIA_ROOT, 'upload/resultados')):
                    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'upload/resultados'))
                caminho_arquivo = os.path.join(settings.MEDIA_ROOT,destino_arquivo)
                data_emissao = datetime.date.today()

                configuracao = get_config(request.user.pessoafisica.setor.secretaria)
                logo = None
                if configuracao.logo:
                    logo = os.path.join(settings.MEDIA_ROOT, configuracao.logo.name)



                data = {'ordens': ordens, 'total':total, 'titulo': 'Ordens de Compra', 'situacao': descricao_situacao, 'configuracao':configuracao, 'logo':logo, 'data_emissao':data_emissao, 'total': total }
                template = get_template('relatorio_gerencial_compras.html')


                html  = template.render(Context(data))

                pdf_file = open(caminho_arquivo, "w+b")
                pisaStatus = pisa.CreatePDF(html.encode('utf-8'), dest=pdf_file,
                        encoding='utf-8')
                pdf_file.close()
                file = open(caminho_arquivo, "r")
                pdf = file.read()
                file.close()
                return HttpResponse(pdf, 'application/pdf')
    else:
        raise PermissionDenied


    return render(request, 'ver_relatorios_gerenciais_compras.html', locals(), RequestContext(request))

@login_required()
def imprimir_aditivo(request, aditivo_id):
    aditivo = get_object_or_404(Aditivo, pk=aditivo_id)
    configuracao = get_config(aditivo.contrato.solicitacao.setor_origem.secretaria)
    contrato = aditivo.contrato

    logo = None
    if configuracao.logo:
        logo = os.path.join(settings.MEDIA_ROOT,configuracao.logo.name)


    municipio = None
    config_geral = get_config_geral()
    if get_config_geral():
        municipio = get_config_geral().municipio


    document = Document()

    # table = document.add_table(rows=2, cols=2)
    # hdr_cells = table.rows[0].cells
    # hdr_cells2 = table.rows[1].cells
    # style2 = document.styles['Normal']
    # font = style2.font
    # font.name = 'Arial'
    # font.size = Pt(6)
    #
    # style = document.styles['Normal']
    # font = style.font
    # font.name = 'Arial'
    # font.size = Pt(11)
    #
    # paragraph = hdr_cells[0].paragraphs[0]
    # run = paragraph.add_run()
    # run.add_picture(logo, width=Inches(1.75))
    #
    # paragraph2 = hdr_cells[1].paragraphs[0]
    # paragraph2.style = document.styles['Normal']
    # hdr_cells[1].text =  u'%s' % (configuracao.nome)
    #
    #
    # paragraph3 = hdr_cells2[1].paragraphs[0]
    # paragraph3.style2 = document.styles['Normal']
    #
    # #hdr_cells2[0].text =  u'Sistema Orçamentário, Financeiro e Contábil'
    # hdr_cells2[1].text =  u'Endereço: %s, %s' % (configuracao.endereco, municipio)
    # a, b = hdr_cells2[:2]
    # a.merge(b)
    imprimir_cabecalho(document, configuracao, logo, municipio)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'')

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    p.add_run(u'MINUTA %sº TERMO ADITIVO AO CONTRATO Nº %s' % (aditivo.ordem, aditivo.contrato.numero)).bold = True


    fornecedor = aditivo.contrato.get_fornecedor()
    if hasattr(fornecedor, 'fornecedor'):
        fornecedor = aditivo.contrato.get_fornecedor().fornecedor

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells

    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.add_run()

    paragraph2 = hdr_cells[1].paragraphs[0]
    hdr_cells[1].text =  u'Minuta %s° Termo Aditivo ao contrato nº %s firmado entre a %s e o MUNICIPIO DE %s.' % (aditivo.ordem, aditivo.contrato.numero, fornecedor.razao_social, config_geral.nome)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'')

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fornecedor_nome = fornecedor_cpf = fornecedor_endereco = fornecedor_nome_representante = fornecedor_cnpj = u''
    registro_fornecedor = ItemContrato.objects.filter(contrato=aditivo.contrato)[0]
    if registro_fornecedor.participante:
        fornecedor_nome_representante = registro_fornecedor.participante.nome_representante
        fornecedor_cpf = registro_fornecedor.participante.cpf_representante
        fornecedor_endereco = registro_fornecedor.participante.fornecedor.endereco
        fornecedor_cnpj = registro_fornecedor.participante.fornecedor.cnpj
        fornecedor_nome = registro_fornecedor.participante.fornecedor.razao_social
    elif registro_fornecedor.fornecedor:
        if FornecedorCRC.objects.filter(fornecedor=fornecedor).exists():
            crc = FornecedorCRC.objects.get(fornecedor=fornecedor)
            fornecedor_nome = crc.razao_social
            fornecedor_cpf = crc.cpf
            fornecedor_endereco = crc.fornecedor.endereco
            fornecedor_cnpj = crc.fornecedor.cnpj
            fornecedor_nome_representante = crc.nome

    p.add_run(u'CONTRATADA: empresa %s, inscrita no CNPJ/MF nº %s, sediada a Rua %s, neste ato, representada por %s - CPF: %s.' % (fornecedor_nome, fornecedor_cnpj, fornecedor_endereco, fornecedor_nome_representante, fornecedor_cpf))


    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(u'CONTRATANTE: O Município %s, inscrito no CNPJ sob o nº. %s, sediado na %s, neste ato representado por %s - CPF: %s.' % (config_geral.nome, config_geral.cnpj, config_geral.endereco, config_geral.ordenador_despesa.nome, config_geral.cpf_ordenador_despesa ))

    if aditivo.de_prazo and aditivo.de_valor:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'OBJETO').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA PRIMEIRA: ').bold=True
        p = document.add_paragraph()
        dias = (aditivo.data_fim - aditivo.contrato.data_fim).days
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        num_processo = u' '
        if aditivo.contrato.solicitacao.processo:
            num_processo = aditivo.contrato.solicitacao.processo.numero
        texto = u'''
        O objeto do presente aditivo é PRORROGAR em %s dias a vigência e ACRESCER o valor do contrato %s em %s%% (%s por cento)
        do valor contratado originariamente  nos autos do Processo Administrativo nº %s, referente a %s, cujo objeto trata de %s
        ''' % (dias, aditivo.contrato.numero, aditivo.indice_total_contrato, num_processo, aditivo.contrato.pregao, aditivo.contrato.solicitacao.objeto)
        p.add_run(texto)



        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA SEGUNDA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        valor_final = aditivo.contrato.valor + aditivo.valor
        texto = u'''A alteração do valor passando de R$ %s para R$ %s ''' % (aditivo.contrato.valor, valor_final)
        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'DA DOTAÇÃO ORÇAMENTÁRIA').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA TERCEIRA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        texto = u'''
        A despesa com a prestação do serviço objeto deste termo aditivo ao Contrato, no valor de R$ (__________), que será mediante a emissão da nota de empenho, em conformidade com as dotações orçamentárias listadas abaixo.
        '''
        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'ÓRGÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'UNIDADE: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'FUNÇÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'SUBFUNÇÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'PROGRAMA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'PROJETO/ATIVIDADE: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'ELEMENTO DE DESPESA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'DA VIGÊNCIA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA QUARTA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        texto = u'''
        O presente Aditivo ao contrato %s terá vigência de ATÉ %s (%s) DIAS CORRIDOS, contatos da data de assinatura.''' % (aditivo.contrato.numero, dias, format_numero(dias))

        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'DAS DEMAIS CLÁUSULAS DO CONTRATO ORIGINAL ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA QUINTA: ').bold=True
        p = document.add_paragraph()
        texto = u'''
        Ficam mantidas integralmente todas as demais cláusulas do contrato original que não se conflitarem com o presente termo.


        E por estarem assim justos e contratados, assinam o presente termo em três vias de igual teor e forma para que produzam os efeitos da Lei.
        '''
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(texto)

    elif aditivo.de_prazo:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'OBJETO').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA PRIMEIRA: ').bold=True
        p = document.add_paragraph()
        dias = (aditivo.data_fim - aditivo.contrato.data_fim).days
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        num_processo = u' '
        if aditivo.contrato.solicitacao.processo:
            num_processo = aditivo.contrato.solicitacao.processo.numero

        texto = u'''

        O objeto do presente aditivo é PRORROGAR em %s dias a vigência do contrato %s do Processo Administrativo nº %s referente a %s,
        cujo objeto trata de %s
        ''' % (dias, aditivo.contrato.numero,  num_processo, aditivo.contrato.pregao, aditivo.contrato.solicitacao.objeto)
        p.add_run(texto)



        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA SEGUNDA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        texto = u'''
        O presente Aditivo ao contrato %s terá vigência de ATÉ %s (%s) DIAS CORRIDOS, contatos da data de assinatura.''' % (aditivo.contrato.numero, dias, format_numero(dias))

        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'DA DOTAÇÃO ORÇAMENTÁRIA').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA TERCEIRA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        texto = u'''
        A despesa com a prestação do serviço objeto deste termo aditivo ao Contrato, no valor de R$ (__________), que será mediante a emissão da nota de empenho, em conformidade com as dotações orçamentárias listadas abaixo.
        '''
        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'ÓRGÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'UNIDADE: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'FUNÇÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'SUBFUNÇÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'PROGRAMA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'PROJETO/ATIVIDADE: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'ELEMENTO DE DESPESA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'DA VIGÊNCIA: ').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'DAS DEMAIS CLÁUSULAS DO CONTRATO ORIGINAL ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA QUARTA: ').bold=True
        p = document.add_paragraph()
        texto = u'''
        Ficam mantidas integralmente todas as demais cláusulas do contrato original que não se conflitarem com o presente termo.


        E por estarem assim justos e contratados, assinam o presente termo em três vias de igual teor e forma para que produzam os efeitos da Lei.
        '''
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(texto)

    elif aditivo.de_valor:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'OBJETO').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA PRIMEIRA: ').bold=True
        p = document.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        num_processo = u' '
        if aditivo.contrato.solicitacao.processo:
            num_processo = aditivo.contrato.solicitacao.processo.numero

        texto = u'''
        O objeto do presente aditivo é ACRESCER o valor do contrato %s em %s%% do valor contratado originariamente  nos autos do Processo Administrativo nº %s, referente a %s, cujo objeto trata de %s
        ''' % (aditivo.contrato.numero, aditivo.indice_total_contrato, num_processo, aditivo.contrato.pregao, aditivo.contrato.solicitacao.objeto)
        p.add_run(texto)


        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA SEGUNDA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        valor_atual = 0
        if aditivo.valor_atual:
            valor_atual = aditivo.valor_atual

        valor_antes_do_aditivo = aditivo.get_valor_atual_aditivo_anterior()
        texto = u'''A alteração do valor passando de R$ %s (%s) para R$ %s (%s)''' % (format_money(valor_antes_do_aditivo), format_numero_extenso(valor_antes_do_aditivo), format_money(valor_atual), format_numero_extenso(valor_atual))
        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'DA DOTAÇÃO ORÇAMENTÁRIA').bold=True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA TERCEIRA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        valor = valor_atual - valor_antes_do_aditivo
        texto = u'''
        A despesa com a prestação do serviço objeto deste termo aditivo ao Contrato, no valor de R$ %s (%s), que será mediante a emissão da nota de empenho, em conformidade com as dotações orçamentárias listadas abaixo.
        ''' % (format_money(valor), format_numero_extenso(valor))
        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'ÓRGÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'UNIDADE: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'FUNÇÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'SUBFUNÇÃO: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'PROGRAMA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'PROJETO/ATIVIDADE: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'ELEMENTO DE DESPESA: ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'DA VIGÊNCIA: ').bold=True

        p.add_run(u'CLÁUSULA QUARTA: ').bold=True
        p = document.add_paragraph()
        texto = u'''
        O presente Aditivo ao contrato %s terá validade, contada da data de assinatura.''' % (aditivo.contrato.numero)

        p.add_run(texto)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(u'DAS DEMAIS CLÁUSULAS DO CONTRATO ORIGINAL ').bold=True
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(u'CLÁUSULA QUINTA: ').bold=True
        p = document.add_paragraph()
        texto = u'''
        Ficam mantidas integralmente todas as demais cláusulas do contrato original que não se conflitarem com o presente termo.


        E por estarem assim justos e contratados, assinam o presente termo em três vias de igual teor e forma para que produzam os efeitos da Lei.
        '''
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(texto)





    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(u'%s, %s' % (config_geral.municipio, datetime.datetime.now().date().strftime('%d/%m/%Y')))

    p = document.add_paragraph()


    orgao = config_geral.nome
    nome_pessoa_ordenadora = config_geral.ordenador_despesa.nome


    document.add_paragraph()
    texto = u'%s' % (nome_pessoa_ordenadora)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run(texto)


    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(orgao)

    document.add_paragraph()

    texto = u'%s' % (fornecedor_nome_representante)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run(texto)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(fornecedor_nome)
    document.add_paragraph()


    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'_________________________').bold=True
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'Contratante ').bold=True
    #
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'_________________________').bold=True
    # p = document.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run(u'Contratada ').bold=True




    texto = u'TESTEMUNHAS:'
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.add_run(texto).underline=True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=2, cols=4)

    hdr_cells = table.rows[0].cells
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(2.5)
    table.columns[2].width = Inches(0.5)
    table.columns[3].width = Inches(2.5)



    hdr_cells[0].text = u'1) '
    hdr_cells[1].text = u'______________________'
    hdr_cells[2].text = u'2) '
    hdr_cells[3].text = u'______________________'

    hdr_cells = table.rows[1].cells

    hdr_cells[0].text = u'CPF: '
    hdr_cells[1].text = u'______________________'
    hdr_cells[2].text = u'CPF: '
    hdr_cells[3].text = u'______________________'





    caminho_arquivo = os.path.join(settings.MEDIA_ROOT, 'upload/pregao/atas/termo_aditivo_%s.docx' % aditivo.id)
    document.save(caminho_arquivo)



    nome_arquivo = caminho_arquivo.split('/')[-1]
    extensao = nome_arquivo.split('.')[-1]
    arquivo = open(caminho_arquivo, "rb")

    content_type = caminho_arquivo.endswith('.pdf') and 'application/pdf' or 'application/vnd.ms-word'
    response = HttpResponse(arquivo.read(), content_type=content_type)
    response['Content-Disposition'] = 'attachment; filename=%s' % nome_arquivo
    arquivo.close()
    os.unlink(caminho_arquivo)
    return response

@login_required()
def transferir_quantidade_item_arp(request, itemarp_id):
    item = get_object_or_404(ItemAtaRegistroPreco, pk=itemarp_id)
    title = u'Transferir Item entre Secretarias'
    ata = item.ata
    if ata.solicitacao.setor_atual == request.user.pessoafisica.setor or True:
        form = TransfereItemARPForm(request.POST or None, item=item)
        if form.is_valid():
            o = form.save(False)
            o.cadastrado_em = datetime.datetime.now()
            o.cadastrado_por = request.user.pessoafisica
            o.item = item
            o.save()
            messages.success(request, u'Transferência realizada com sucesso.')
            return HttpResponseRedirect(u'/base/visualizar_ata_registro_preco/%s/' % ata.id)
        return render(request, 'transferir_quantidade_item_arp.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

def busca_saldo_atual(request):
    from django.core import serializers
    if request.method == 'GET':
        item = request.GET.get('item')
        secretaria = request.GET.get('secretaria')
        if secretaria and item:
            data = ItemAtaRegistroPreco.objects.get(id=item).get_saldo_atual_secretaria(Secretaria.objects.get(id=secretaria))
        else:
            data = []
        return HttpResponse(data, content_type='application/json')


@login_required()
def modelos_documentos(request):
    title = u'Modelos de Documentos'
    documentos = ModeloDocumento.objects.all().order_by('-id')

    form = BuscarModeloDocumentoForm(request.GET or None)

    if form.is_valid():
        if form.cleaned_data.get('nome'):
            documentos = documentos.filter(nome=form.cleaned_data.get('nome'))

        if form.cleaned_data.get('palavra'):
            documentos = documentos.filter(palavras_chaves__icontains=form.cleaned_data.get('palavra'))

        if form.cleaned_data.get('tipo'):
            documentos = documentos.filter(tipo=form.cleaned_data.get('tipo'))

        if form.cleaned_data.get('tipo_objeto'):
            documentos = documentos.filter(tipo_objeto=form.cleaned_data.get('tipo_objeto'))


    return render(request, 'modelos_documentos.html', locals(), RequestContext(request))



@login_required()
def cadastrar_modelo_documento(request):
    title = u'Cadastrar Modelo de Documento'
    form = ModeloDocumentoForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        o = form.save(False)
        o.cadastrado_em = datetime.datetime.now()
        o.cadastrado_por = request.user.pessoafisica
        o.save()
        messages.success(request, u'Modelo cadastrado com sucesso.')
        return HttpResponseRedirect(u'/base/modelos_documentos/')


    return render(request, 'cadastrar_modelo_ata.html', locals(), RequestContext(request))

@login_required()
def editar_modelo_documento(request, documento_id):
    title = u'Editar Modelo de Documento'
    documento = get_object_or_404(ModeloDocumento, pk=documento_id)
    if request.user.pessoafisica == documento.cadastrado_por:
        form = ModeloDocumentoForm(request.POST or None, request.FILES or None, instance=documento)
        if form.is_valid():
            form.save()
            messages.success(request, u'Modelo editado com sucesso.')
            return HttpResponseRedirect(u'/base/modelos_documentos/')

        return render(request, 'cadastrar_modelo_ata.html', locals(), RequestContext(request))
    else:
        raise PermissionDenied

@login_required()
def deletar_modelo_documento(request, documento_id):
    documento = get_object_or_404(ModeloDocumento, pk=documento_id)
    if request.user.pessoafisica == documento.cadastrado_por:
        documento.delete()
        messages.success(request, u'Modelo excluído com sucesso.')
        return HttpResponseRedirect(u'/base/modelos_documentos/')

    else:
        raise PermissionDenied

@login_required()
def cadastrar_certidao_crc(request, crc_id):
    crc = get_object_or_404(FornecedorCRC, pk=crc_id)
    title = u'Cadastrar Certidão - %s' % crc.fornecedor
    form = CertidaoCRCForm(request.POST or None, request.FILES or None)
    if form.is_valid():
        o = form.save(False)
        o.crc = crc
        o.save()
        messages.success(request, u'Certidão cadastrada com sucesso.')
        return HttpResponseRedirect(u'/base/ver_crc/{}/'.format(crc.fornecedor.id))

    return render(request, 'cadastrar_certidao_crc.html', locals(), RequestContext(request))
